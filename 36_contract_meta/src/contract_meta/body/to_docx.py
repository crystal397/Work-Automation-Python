"""마크다운(GFM 일부) → docx 변환.

처리 요소
- ATX 헤딩 (#, ##, ###, ####) → Heading 1~4
- 표 (| ... | ... |) → docx table (가운데 정렬 헤더)
- 인용 (>) → 'Quote' 스타일 단락 (없으면 italic)
- 빈 줄·구분선(---) → 단락 구분
- 줄바꿈 표시 `<br>` → 표 셀 내 줄바꿈

화려한 스타일·페이지 머리말은 v0.2 에서.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt


_TABLE_LINE_RE = "|"


def md_to_docx(md_text: str, out_path: str | Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            i = _consume_table(doc, lines, i)
            continue

        para = doc.add_paragraph()
        for token, bold in _parse_inline(line):
            r = para.add_run(token)
            r.bold = bold
        i += 1

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inside = s.strip("|").strip()
    cells = [c.strip() for c in inside.split("|")]
    return all(set(c) <= set("-: ") and ("-" in c) for c in cells)


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _consume_table(doc: Document, lines: list[str], i: int) -> int:
    header = _split_row(lines[i])
    i += 2  # skip separator
    body_rows: list[list[str]] = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        body_rows.append(_split_row(lines[i]))
        i += 1
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header[:n_cols]):
        hdr_cells[j].text = ""
        p = hdr_cells[j].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(body_rows, start=1):
        cells = table.rows[r_idx].cells
        for j in range(n_cols):
            txt = row[j] if j < len(row) else ""
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            for sub_idx, sub in enumerate(txt.split("<br>")):
                if sub_idx > 0:
                    p = cells[j].add_paragraph()
                for token, bold in _parse_inline(sub):
                    r = p.add_run(token)
                    r.bold = bold
    return i


def _parse_inline(text: str) -> list[tuple[str, bool]]:
    """**bold** 만 처리 (간단)."""
    out: list[tuple[str, bool]] = []
    bold = False
    buf = ""
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            if buf:
                out.append((buf, bold))
                buf = ""
            bold = not bold
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out
