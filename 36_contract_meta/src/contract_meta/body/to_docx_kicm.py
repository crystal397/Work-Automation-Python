"""KICM 표준 양식 docx 렌더러 (to_docx.py 의 보강 버전).

표준 양식 적용
- 한글 폰트: 맑은 고딕 (Malgun Gothic), 본문 10pt / 제목 14~18pt
- 페이지: A4 세로, 여백 2cm
- 페이지 머리말: 공사명
- 페이지 바닥글: '- N -' 형태 페이지 번호 (가운데)
- 인용 박스: 회색 테두리 표 (1×1 셀)
- 표: 그리드 + 헤더 굵게·회색 배경

v0.1 — 핵심 양식만. KICM 로고는 별도 이미지 파일 경로로 받음 (없으면 생략).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


KOREAN_FONT = "맑은 고딕"


def md_to_docx_kicm(
    md_text: str,
    out_path: str | Path,
    *,
    project_name: str = "",
    logo_path: str | Path | None = None,
) -> Path:
    doc = Document()
    _setup_page(doc, project_name=project_name)
    _setup_styles(doc)

    if logo_path and Path(logo_path).exists():
        p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(logo_path), width=Cm(3.5))

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip() or line.strip() == "---":
            i += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            h = doc.add_heading(text, level=min(level, 4))
            _set_font(h, KOREAN_FONT, size=18 - 2 * (level - 1), bold=True)
            i += 1
            continue
        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            _add_quote_box(doc, text)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            i = _consume_table(doc, lines, i)
            continue
        para = doc.add_paragraph()
        for token, bold in _parse_inline(line):
            r = para.add_run(token)
            r.bold = bold
            r.font.name = KOREAN_FONT
            r.font.size = Pt(10)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
        i += 1

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path


def _setup_page(doc: Document, *, project_name: str) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    if project_name:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(project_name)
        run.font.name = KOREAN_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "- "
    _add_page_field(fp)
    run_dash = fp.add_run(" -")
    run_dash.font.name = KOREAN_FONT
    run_dash.font.size = Pt(10)


def _add_page_field(paragraph) -> None:
    """Word 필드: PAGE 번호."""
    from docx.oxml import OxmlElement
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._element.append(fld_char1)
    run._element.append(instr_text)
    run._element.append(fld_char2)


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)


def _set_font(paragraph, font_name: str, *, size: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_quote_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    for token, bold in _parse_inline(text):
        run = p.add_run(token)
        run.bold = bold
        run.font.name = KOREAN_FONT
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inside = s.strip("|").strip()
    cells = [c.strip() for c in inside.split("|")]
    return all(set(c) <= set("-: ") and ("-" in c) for c in cells)


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _consume_table(doc: Document, lines: list[str], i: int) -> int:
    header = _split_row(lines[i])
    i += 2
    body_rows: list[list[str]] = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        body_rows.append(_split_row(lines[i]))
        i += 1
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header[:n_cols]):
        hdr_cells[j].text = ""
        p = hdr_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = KOREAN_FONT
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
        _shade_cell(hdr_cells[j], "DDDDDD")
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(body_rows, start=1):
        cells = table.rows[r_idx].cells
        for j in range(n_cols):
            txt = row[j] if j < len(row) else ""
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            for sub_idx, sub in enumerate(txt.split("<br>")):
                if sub_idx > 0:
                    p = cells[j].add_paragraph()
                for token, bold in _parse_inline(sub):
                    run = p.add_run(token)
                    run.bold = bold
                    run.font.name = KOREAN_FONT
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    return i


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _parse_inline(text: str) -> list[tuple[str, bool]]:
    out: list[tuple[str, bool]] = []
    bold = False
    buf = ""
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            if buf:
                out.append((buf, bold))
                buf = ""
            bold = not bold
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out
