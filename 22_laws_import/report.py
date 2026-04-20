"""Word(.docx) 리포트 생성기

생성 구조:
  표지 — 입찰공고일, 작성일시
  1. 매칭 결과 요약 테이블
  2. 법령별 상세 (기본정보 / 부칙 경고 / 공기연장 관련 조문 원문)
  3. 검토 필요 항목 목록
"""
import logging
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from engine import LawVersion, MatchResult

logger = logging.getLogger(__name__)

_RED = RGBColor(0xC0, 0x00, 0x00)
_ORANGE = RGBColor(0xFF, 0x8C, 0x00)
_GRAY = RGBColor(0x60, 0x60, 0x60)


def _set_font(run, size: int = 10, bold: bool = False, color: RGBColor = None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 한글 폰트 설정
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "맑은 고딕"
        rpr = run._r.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _source_url(version: LawVersion) -> str:
    return f"https://www.law.go.kr/lsInfoP.do?lsiSeq={version.law_id}"


class WordReportGenerator:
    """공기연장 법령 매칭 결과 Word 리포트"""

    def __init__(self):
        self.doc = Document()
        self._bid_date: date = date.today()  # generate() 호출 시 덮어씀
        self._setup_page()

    # ── 초기화 ────────────────────────────────────────────────────────────────

    def _setup_page(self) -> None:
        sec = self.doc.sections[0]
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

        style = self.doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(10)

    # ── 섹션 구성 ─────────────────────────────────────────────────────────────

    def _add_cover(self, bid_date: date) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("입찰공고일 기준 적용 법령 매칭 결과")
        _set_font(run, size=18, bold=True)

        self.doc.add_paragraph()

        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = info_para.add_run(f"입찰공고일: {bid_date}")
        _set_font(r1, size=12)
        info_para.add_run("　　")
        r2 = info_para.add_run(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        _set_font(r2, size=12, color=_GRAY)

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "※ 본 문서의 법령 조문은 법제처 국가법령정보 API에서 수신한 원문입니다.\n"
            "   AI가 조문을 생성·요약·수정하지 않았습니다."
        ).runs[0].font.size = Pt(9)
        self.doc.add_page_break()

    def _add_summary_table(self, results: list[MatchResult]) -> None:
        _heading(self.doc, "1. 매칭 결과 요약", 1)

        headers = ["법령명", "공포번호", "시행일", "부칙\n경과규정", "비고"]
        col_widths = [Cm(5.5), Cm(3.0), Cm(2.8), Cm(2.2), Cm(4.5)]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 헤더
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            cell = table.rows[0].cells[i]
            cell.width = w
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(h)
            _set_font(run, size=9, bold=True)

        # 데이터 행
        for r in results:
            row = table.add_row()
            cells = row.cells

            if r.selected:
                v = r.selected
                trans_cell = (
                    (f"있음 ⚠ ({r.transitional_type}형)" if r.transitional_type else "있음 ⚠")
                    if r.transitional_flag
                    else "없음"
                )
                note = r.consistency_warning or r.warning or (
                    "사용자 확인 필요" if r.needs_user_review else "정상"
                )
                data = [
                    r.display_name,
                    v.announce_num,
                    str(v.enforce_date),
                    trans_cell,
                    note,
                ]
            else:
                data = [r.display_name, "-", "-", "-", r.warning or "조회 실패"]

            admrul_no_history = (
                r.selected
                and r.selected.target == "admrul"
                and "연혁 조회 불가" in (r.warning or "")
            )
            for i, text in enumerate(data):
                para = cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(text)
                _set_font(run, size=9)
                if r.transitional_flag and i == 3:
                    run.font.color.rgb = _RED
                elif admrul_no_history:
                    run.font.color.rgb = _ORANGE  # 행정규칙 연혁 확보 실패
                elif not r.selected:
                    run.font.color.rgb = _ORANGE

        self.doc.add_paragraph()

    def _add_detail_sections(self, results: list[MatchResult]) -> None:
        _heading(self.doc, "2. 법령별 상세 내용", 1)

        for idx, r in enumerate(results, 1):
            _heading(self.doc, f"2-{idx}. {r.display_name}", 2)

            if not r.selected:
                p = self.doc.add_paragraph()
                run = p.add_run(f"⚠ {r.warning or '조회 실패'}")
                _set_font(run, color=_RED)
                continue

            v = r.selected

            # 기본 정보 — 인용 형식
            info = self.doc.add_paragraph()
            citation = v.citation if (v.law_type or v.revision_type) else None
            if citation:
                run_c = info.add_run(citation)
                _set_font(run_c, bold=True)
            else:
                # 법령 본문 미로드 시 간략 형식으로 fallback
                for label, value in [
                    ("법령명", v.name),
                    ("공포번호", v.announce_num),
                    ("공포일", str(v.announce_date)),
                    ("시행일", str(v.enforce_date)),
                ]:
                    run_l = info.add_run(f"{label}: ")
                    _set_font(run_l, bold=True)
                    run_v = info.add_run(f"{value}　")
                    _set_font(run_v)

            url_p = self.doc.add_paragraph()
            run_label = url_p.add_run("출처 URL: ")
            _set_font(run_label, bold=True)
            run_url = url_p.add_run(_source_url(v))
            _set_font(run_url, color=_GRAY)

            # 버전 선정 근거
            basis_p = self.doc.add_paragraph()
            if v.target == "admrul":
                # 연혁 조회 불가 = warning에 "연혁 조회 불가" 문구가 있는 경우만
                # (warning이 빈 문자열이면 API 연혁 성공, 다른 문구면 scraper/search 성공)
                history_found = "연혁 조회 불가" not in (r.warning or "")
                if history_found:
                    run_b = basis_p.add_run(
                        f"[선정 근거] 행정규칙 — 입찰공고일({self._bid_date}) 기준 "
                        f"시행일 {v.enforce_date} 이하 중 최근 버전 선택 "
                        f"(연혁 확보 성공 — 실제 시행 버전 추가 확인 권장)"
                    )
                    _set_font(run_b, size=9, color=_GRAY)
                else:
                    run_b = basis_p.add_run(
                        "[선정 근거] 행정규칙 — 연혁 조회 불가, 현행 버전 기준 표시 (수동 확인 필요)"
                    )
                    _set_font(run_b, size=9, color=_ORANGE)
            else:
                run_b = basis_p.add_run(
                    f"[선정 근거] 입찰공고일({self._bid_date}) 기준 — 시행일 {v.enforce_date} 이하 중 최근 버전 선택"
                )
                _set_font(run_b, size=9, color=_GRAY)

            # 5단계 정합성 경고
            if r.consistency_warning:
                con_box = self.doc.add_paragraph()
                run_c = con_box.add_run(r.consistency_warning)
                _set_font(run_c, size=9, bold=True, color=_ORANGE)

            # 부칙 경과규정 경고 (유형 A/B 구분)
            if r.transitional_flag:
                type_label = {
                    "A": "유형 A (법령 전체 경과규정)",
                    "B": "유형 B (특정 조문 단위 경과규정)",
                }.get(r.transitional_type, "유형 미확인")

                warn_box = self.doc.add_paragraph()
                run_w1 = warn_box.add_run(
                    f"⚠ 부칙 경과규정 탐지 [{type_label}] — 사용자 확인 필요\n"
                )
                _set_font(run_w1, bold=True, color=_RED)
                run_w2 = warn_box.add_run(f"발견 문장: {r.transitional_text}")
                _set_font(run_w2, size=9, color=_RED)

                # 유형 B: 영향받는 조번호 표시
                if r.transitional_type == "B" and r.transitional_articles:
                    art_list = ", ".join(f"제{n}조" for n in r.transitional_articles)
                    art_p = self.doc.add_paragraph()
                    run_a = art_p.add_run(f"  → 영향 조문: {art_list} (해당 조는 직전 버전 적용 검토 필요)")
                    _set_font(run_a, size=9, color=_RED)

                if r.prev_version:
                    pv = r.prev_version
                    prev_p = self.doc.add_paragraph()
                    run_p = prev_p.add_run(
                        f"[직전 버전] {pv.announce_num} | 시행일: {pv.enforce_date} | "
                        f"URL: {_source_url(pv)}"
                    )
                    _set_font(run_p, size=9, color=_ORANGE)

            # 공기연장 관련 조문
            if r.relevant_articles:
                _heading(self.doc, f"공기연장 관련 조문 ({len(r.relevant_articles)}개)", 3)
                for art in r.relevant_articles:
                    art_title = self.doc.add_paragraph(style="List Paragraph")
                    run_t = art_title.add_run(
                        f"제{art['조번호']}조 {art['조제목']}"
                    )
                    _set_font(run_t, bold=True)

                    if art["조문내용"]:
                        art_body = self.doc.add_paragraph(style="List Paragraph")
                        run_c = art_body.add_run(art["조문내용"])
                        _set_font(run_c, size=9)

                    for para_item in art.get("항", []):
                        para_num = str(para_item.get("항번호") or "")
                        para_content = str(para_item.get("항내용") or "")
                        if para_content:
                            p_para = self.doc.add_paragraph(style="List Paragraph")
                            p_para.paragraph_format.left_indent = Cm(0.5)
                            # 항번호(숫자)를 원문자 ①②③...으로 변환, 범위 초과 시 숫자 그대로
                            try:
                                circle = chr(0x245F + int(para_num))  # ① = U+2460
                            except (ValueError, OverflowError):
                                circle = para_num
                            run_p = p_para.add_run(f"  {circle} {para_content}")
                            _set_font(run_p, size=9)

                        # 호(號) 단위 출력
                        for sub_item in para_item.get("호", []):
                            sub_num = str(sub_item.get("호번호") or "")
                            sub_content = str(sub_item.get("호내용") or "")
                            if sub_content:
                                s_para = self.doc.add_paragraph(style="List Paragraph")
                                s_para.paragraph_format.left_indent = Cm(1.0)
                                run_s = s_para.add_run(f"    {sub_num}. {sub_content}")
                                _set_font(run_s, size=9)

                            # 목(目) 단위 출력
                            for sub_sub in sub_item.get("목", []):
                                ss_num = str(sub_sub.get("목번호") or "")
                                ss_content = str(sub_sub.get("목내용") or "")
                                if ss_content:
                                    ss_para = self.doc.add_paragraph(style="List Paragraph")
                                    ss_para.paragraph_format.left_indent = Cm(1.5)
                                    run_ss = ss_para.add_run(f"      {ss_num}) {ss_content}")
                                    _set_font(run_ss, size=9)
            else:
                no_art = self.doc.add_paragraph()
                run_n = no_art.add_run("(공기연장 관련 조문 없음 또는 키워드 미탐지)")
                _set_font(run_n, size=9, color=_GRAY)

            self.doc.add_paragraph()

    def _add_review_list(self, results: list[MatchResult]) -> None:
        needs = [
            r for r in results
            if r.needs_user_review or r.transitional_flag or r.consistency_warning
        ]
        if not needs:
            return

        _heading(self.doc, "3. 사용자 확인 필요 항목", 1)
        for r in needs:
            p = self.doc.add_paragraph(style="List Bullet")
            reasons = []
            if r.transitional_flag:
                type_label = f"유형 {r.transitional_type}" if r.transitional_type else "유형 미확인"
                reasons.append(f"부칙 경과규정 탐지 ({type_label})")
            if r.consistency_warning:
                reasons.append("상위·하위법 시행일 불일치")
            if r.warning and not r.transitional_flag and not r.consistency_warning:
                reasons.append(r.warning)
            run = p.add_run(
                f"{r.display_name}" + (" — " + " / ".join(reasons) if reasons else "")
            )
            _set_font(run, color=_RED if r.transitional_flag else _ORANGE)

    # ── 공개 메서드 ───────────────────────────────────────────────────────────

    def generate(
        self,
        bid_date: date,
        results: list[MatchResult],
        output_path: Path,
    ) -> Path:
        """리포트 생성 후 저장 경로 반환"""
        self._bid_date = bid_date
        output_path.parent.mkdir(parents=True, exist_ok=True)

        self._add_cover(bid_date)
        self._add_summary_table(results)
        self._add_detail_sections(results)
        self._add_review_list(results)

        self.doc.save(str(output_path))
        logger.info("Word 리포트 저장: %s", output_path)
        return output_path
