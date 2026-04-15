"""
공기연장 간접비 산정 보고서 — Word 템플릿 생성 스크립트
유형 B: 상세형 (대형 국가계약 — 인덕원, 평택~오산, 당진기지 유형)

실행:
    python make_template_B.py

출력:
    reference/보고서_템플릿_B_상세형(국가계약).docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 상수 ─────────────────────────────────────────────────────────────────────

FONT_MAIN   = "한컴바탕"
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_GRAY  = RGBColor(89, 89, 89)
HEADER_BG   = "D9D9D9"
PH          = "【 입력 】"


# ── 공통 헬퍼 ────────────────────────────────────────────────────────────────

def _set_east_asia_font(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc, text="", font_name=FONT_MAIN, size_pt=11,
                  bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=0, space_after_pt=5):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.font.name  = font_name
        run.font.size  = Pt(size_pt)
        run.font.bold  = bold
        run.font.color.rgb = COLOR_BLACK
        _set_east_asia_font(run, font_name)
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after  = Pt(space_after_pt)
    return para


def add_heading(doc, text, level=1):
    cfg = {
        1: dict(size_pt=16, space_before_pt=14, space_after_pt=6),
        2: dict(size_pt=14, space_before_pt=10, space_after_pt=4),
        3: dict(size_pt=12, space_before_pt=6,  space_after_pt=3),
        4: dict(size_pt=11, space_before_pt=4,  space_after_pt=2),
    }[min(level, 4)]
    return add_paragraph(doc, text, bold=True, **cfg)


def add_body(doc, text=""):
    return add_paragraph(doc, text, size_pt=11, space_after_pt=4)


def add_note(doc, text=""):
    para = add_paragraph(doc, text, size_pt=10, space_after_pt=3)
    for run in para.runs:
        run.font.color.rgb = COLOR_GRAY
        run.font.italic    = True
    return para


def add_quote(doc, text):
    """들여쓰기 인용 단락 (법령 조문 등)."""
    para = add_paragraph(doc, text, size_pt=10, space_after_pt=3)
    para.paragraph_format.left_indent  = Cm(0.8)
    para.paragraph_format.right_indent = Cm(0.5)
    for run in para.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    return para


# ── 표 헬퍼 ──────────────────────────────────────────────────────────────────

def set_cell(cell, text, font_name=FONT_MAIN, size_pt=10,
             bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             bg_color: str | None = None):
    cell.text = ""
    para = cell.paragraphs[0]
    run  = para.add_run(text)
    run.font.name  = font_name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    para.alignment = align
    _set_east_asia_font(run, font_name)
    if bg_color:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg_color)
        tcPr.append(shd)
    return cell


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def set_table_border(table):
    tbl    = table._tbl
    tblPr  = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBdr.append(el)
    tblPr.append(tblBdr)


def add_header_row(table, headers, widths_cm=None):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i < len(row.cells):
            set_cell(row.cells[i], h, bold=True, bg_color=HEADER_BG)
    if widths_cm:
        set_col_widths(table, widths_cm)
    set_table_border(table)


def add_data_row(table, values, align=WD_ALIGN_PARAGRAPH.CENTER):
    row = table.add_row()
    for i, v in enumerate(values):
        if i < len(row.cells):
            set_cell(row.cells[i], v, align=align)
    return row


def make_result_section(doc, label="원도급사"):
    """원도급 or 하도급 산정 결과 섹션 (동일 구조)."""
    add_heading(doc, f"공기연장 간접비 {label} 집계표", level=3)
    add_note(doc, "※ [단위: 원, 영세율]")

    tbl = doc.add_table(rows=1, cols=7)
    add_header_row(tbl,
                   ["구  분", "간접노무비", "직접계상\n경비", "승률계상\n경비",
                    "경비 소계", "일반관리비", "이  윤"],
                   widths_cm=[2.5, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5])
    add_data_row(tbl, [label, PH, PH, PH, PH, PH, PH])
    doc.add_paragraph()

    add_heading(doc, "간접노무비", level=3)
    add_heading(doc, "대상인원", level=4)
    tbl2 = doc.add_table(rows=1, cols=6)
    add_header_row(tbl2,
                   ["No.", "소  속", "이  름", "직  무",
                    "간접노무비 산정 대상 기간", "비  고"],
                   widths_cm=[1.0, 3.0, 2.5, 2.5, 5.5, 2.0])
    for i in range(1, 5):
        add_data_row(tbl2, [str(i), PH, PH, PH,
                             PH + " ~ " + PH + "  (" + PH + "일)", ""])
    doc.add_paragraph()

    add_heading(doc, "급여내역", level=4)
    add_note(doc, "※ [단위: 원]  ※ A. 실비 산정 기간: " + PH + " ~ " + PH)
    tbl3 = doc.add_table(rows=2, cols=9)
    h1 = tbl3.rows[0].cells
    for idx, txt in enumerate(["No.", "소속", "이름", "직무",
                                 "A. 실비 산정", "A. 실비 산정",
                                 "A. 실비 산정", "A. 실비 산정", "합계"]):
        set_cell(h1[idx], txt, bold=True, bg_color=HEADER_BG)
    h2 = tbl3.rows[1].cells
    for idx, txt in enumerate(["No.", "소속", "이름", "직무",
                                 "① 급여", "② 일수",
                                 "③ 퇴직급여충당금\n(①/12)", "④ 소계\n(①+③)", "합  계"]):
        set_cell(h2[idx], txt, bold=True, bg_color=HEADER_BG)
    set_col_widths(tbl3, [1.0, 2.5, 2.0, 2.0, 2.5, 1.2, 2.5, 2.5, 2.5])
    set_table_border(tbl3)
    for i in range(1, 4):
        add_data_row(tbl3, [str(i), PH, PH, PH, PH, PH, PH, PH, PH])
    doc.add_paragraph()

    add_heading(doc, "경비", level=3)
    add_heading(doc, "직접계상비목", level=4)
    add_note(doc, "※ [단위: 원]")
    tbl4 = doc.add_table(rows=1, cols=5)
    add_header_row(tbl4,
                   ["비  목", "산정 기간", "금  액 (원)", "증빙자료", "비  고"],
                   widths_cm=[3.5, 4.0, 3.5, 3.5, 2.0])
    for item in ["전력비·수도광열비", "여비·교통비·통신비", "지급임차료",
                 "복리후생비", "소모품비", "세금과공과",
                 "국민건강·연금보험", "도서인쇄비", "지급수수료", "직접계상 소계"]:
        bold = "소계" in item
        row = tbl4.add_row()
        set_cell(row.cells[0], item, bold=bold,
                 bg_color=(HEADER_BG if bold else None),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for j in range(1, 5):
            set_cell(row.cells[j], PH, bold=bold,
                     bg_color=(HEADER_BG if bold else None))
    doc.add_paragraph()

    add_heading(doc, "승률계상비목", level=4)
    add_note(doc, "※ [단위: 원]")
    tbl5 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl5,
                   ["비  목", "산  출  방  식", "적용 요율", "금  액 (원)"],
                   widths_cm=[3.5, 7.0, 2.5, 3.5])
    for item, formula in [("산재보험료", "간접노무비  ×  요율"),
                          ("고용보험료", "간접노무비  ×  요율"),
                          ("승률계상 소계", "")]:
        bold = "소계" in item
        row = tbl5.add_row()
        set_cell(row.cells[0], item, bold=bold,
                 bg_color=(HEADER_BG if bold else None),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[1], formula if not bold else "",
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], PH + " %" if not bold else "",
                 bold=bold, bg_color=(HEADER_BG if bold else None))
        set_cell(row.cells[3], PH, bold=bold,
                 bg_color=(HEADER_BG if bold else None))
    doc.add_paragraph()

    add_heading(doc, "일반관리비 및 이윤", level=3)
    add_heading(doc, "일반관리비", level=4)
    tbl6 = doc.add_table(rows=2, cols=4)
    add_header_row(tbl6,
                   ["산  출  방  식", "적용 요율", "기준 금액 (원)", "일반관리비 (원)"],
                   widths_cm=[6.5, 2.5, 3.5, 4.0])
    set_cell(tbl6.rows[1].cells[0],
             "(간접노무비 + 경비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl6.rows[1].cells[1], PH + " %")
    set_cell(tbl6.rows[1].cells[2], PH)
    set_cell(tbl6.rows[1].cells[3], PH)
    doc.add_paragraph()

    add_heading(doc, "이윤", level=4)
    tbl7 = doc.add_table(rows=2, cols=4)
    add_header_row(tbl7,
                   ["산  출  방  식", "적용 요율", "기준 금액 (원)", "이  윤 (원)"],
                   widths_cm=[6.5, 2.5, 3.5, 4.0])
    set_cell(tbl7.rows[1].cells[0],
             "(간접노무비 + 경비 + 일반관리비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl7.rows[1].cells[1], PH + " %")
    set_cell(tbl7.rows[1].cells[2], PH)
    set_cell(tbl7.rows[1].cells[3], PH)
    doc.add_paragraph()


# ── 표지 / 제출문 / 요약문 ────────────────────────────────────────────────────
# (유형 A와 동일)

def make_cover(doc):
    for _ in range(6): doc.add_paragraph()
    add_paragraph(doc, PH + "  주식회사",
                  size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)
    for _ in range(2): doc.add_paragraph()
    add_paragraph(doc, PH + "  공사",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    add_paragraph(doc, "공기연장 간접비 산정 보고서",
                  size_pt=24, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)
    for _ in range(8): doc.add_paragraph()
    add_paragraph(doc, PH + " 년  " + PH + " 월",
                  size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    add_paragraph(doc, "사단법인 한국건설관리연구원",
                  size_pt=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)


def make_submission(doc):
    doc.add_page_break()
    add_paragraph(doc, "제   출   문",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=16)
    add_body(doc, PH + " (발주처명)  귀중")
    doc.add_paragraph()
    add_body(doc,
             "사단법인 한국건설관리연구원은 귀 기관으로부터 의뢰받은 "
             "「" + PH + " 공사」의 공기연장 간접비 산정 용역을 완료하고, "
             "그 결과를 다음과 같이 제출합니다.")
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)
    rows_data = [
        ("용 역 명", PH + "  공사 공기연장 간접비 산정"),
        ("용역기간", PH + " ~ " + PH),
        ("수행기관", "사단법인 한국건설관리연구원"),
        ("대  표  자", PH),
    ]
    for i, (k, v) in enumerate(rows_data):
        set_cell(tbl.rows[i].cells[0], k, bold=True, bg_color=HEADER_BG,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(tbl.rows[i].cells[1], v, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_col_widths(tbl, [4.0, 11.5])
    set_table_border(tbl)
    doc.add_paragraph()
    add_paragraph(doc, PH + " 년  " + PH + " 월",
                  size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)
    add_paragraph(doc, "사단법인 한국건설관리연구원",
                  size_pt=12, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)
    add_paragraph(doc, "원  장  " + PH + "  (인)",
                  size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)


def make_summary(doc):
    doc.add_page_break()
    add_paragraph(doc, "요      약      문",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=12)
    add_paragraph(doc, "< 산 정 결 과 >",
                  size_pt=13, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    add_note(doc, "※ 공사명: " + PH)
    add_note(doc, "※ 산정 기간: " + PH + " ~ " + PH + "  (" + PH + " 일)")
    doc.add_paragraph()
    add_paragraph(doc, "○ 원가계산서  (단위: 원, 영세율)",
                  size_pt=11, space_after_pt=3)
    tbl = doc.add_table(rows=1, cols=6)
    add_header_row(tbl,
                   ["구  분", "간접노무비", "경  비", "일반관리비", "이  윤", "합  계"],
                   widths_cm=[3.0, 3.0, 3.0, 3.0, 3.0, 3.0])
    for label in ["원도급사", "하도급사", "합  계"]:
        add_data_row(tbl, [label, PH, PH, PH, PH, PH])
    doc.add_paragraph()
    tbl2 = doc.add_table(rows=1, cols=3)
    add_header_row(tbl2,
                   ["구  분", "공기연장 간접비 합계 (원)", "비  고"],
                   widths_cm=[4.0, 7.5, 4.5])
    add_data_row(tbl2, ["최 종 청구액", PH, "VAT 별도"])


def make_toc(doc):
    doc.add_page_break()
    add_paragraph(doc, "목  차",
                  size_pt=16, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=10, space_after_pt=12)

    toc = [
        (1, "제1장.  개요"),
        (2, "1.1  사업의 개요 및 특성"),
        (3, "1.1.1  용역 대상 공사의 개요 및 특성"),
        (3, "1.1.2  공사 위치도"),
        (2, "1.2  계약의 현황"),
        (3, "1.2.1  계약의 주요 용어의 정의"),
        (3, "1.2.2  계약문서"),
        (3, "1.2.3  계약당사자 현황"),
        (2, "1.3  계약체결 및 변경계약의 경과"),
        (3, "1.3.1  원도급사 계약 현황"),
        (3, "1.3.2  하도급사 계약 현황"),
        (3, "1.3.3  공사기간 변경계약 현황"),
        (2, "1.4  과업 수행 절차"),
        (3, "1.4.1  수행 절차"),
        (3, "1.4.2  산정의 근거"),
        (1, "제2장.  계약의 성격"),
        (2, "2.1  국가계약법 적용 공사"),
        (2, "2.2  계속비 계약"),
        (2, "2.3  내역입찰 공사"),
        (1, "제3장.  공기연장 귀책 분석 및 계약금액조정 검토"),
        (2, "3.1  공기연장 귀책 분석"),
        (3, "3.1.1  지연사유 발생 경위"),
        (3, "3.1.2  공기지연 귀책사유 검토"),
        (2, "3.2  계약금액조정 관련 검토"),
        (3, "3.2.1  분석 개요"),
        (3, "3.2.2  계약금액조정 신청 관련 문서"),
        (3, "3.2.3  검토 의견"),
        (2, "3.3  공기연장 간접비 청구의 근거"),
        (3, "3.3.1  계약에 따른 청구의 근거"),
        (3, "3.3.2  관련 법령에 따른 청구의 근거"),
        (3, "3.3.3  사정변경 원칙에 따른 청구의 근거"),
        (3, "3.3.4  청구의 근거 검토"),
        (1, "제4장.  공기연장 간접비 산정"),
        (2, "4.1  공기연장 간접비 산정 대상 기간 검토"),
        (3, "4.1.1  공기연장 일수의 산정방식"),
        (3, "4.1.2  공기연장 일수 산정"),
        (2, "4.2  공기연장 간접비 산정 방식"),
        (3, "4.2.1  간접노무비 산정 방식"),
        (3, "4.2.2  경비 산정 방식"),
        (3, "4.2.3  일반관리비 및 이윤 산정 방식"),
        (3, "4.2.4  보증수수료 산정 방식"),
        (2, "4.3  원도급사 산정 결과"),
        (2, "4.4  하도급사 산정 결과"),
        (1, "제5장.  결론"),
        (1, "제6장.  첨부자료"),
        (2, "6.1  공기연장 간접비 산정근거"),
        (2, "6.2  계약문서"),
        (2, "6.3  수발신문서"),
        (2, "6.4  판례 및 관련 조항"),
        (2, "6.5  공기연장 간접비 증빙자료"),
    ]
    indent_map = {1: Cm(0), 2: Cm(0.8), 3: Cm(1.6), 4: Cm(2.4)}
    size_map   = {1: 12, 2: 11, 3: 10, 4: 10}
    bold_map   = {1: True, 2: False, 3: False, 4: False}
    for lvl, text in toc:
        para = add_paragraph(doc, text, size_pt=size_map[lvl],
                             bold=bold_map[lvl], space_after_pt=2)
        para.paragraph_format.left_indent = indent_map[lvl]


# ── 제1장: 개요 (확장형) ─────────────────────────────────────────────────────

def make_ch1(doc):
    doc.add_page_break()
    add_heading(doc, "제1장.  개요", level=1)

    # 1.1 사업의 개요 및 특성
    add_heading(doc, "1.1  사업의 개요 및 특성", level=2)
    add_heading(doc, "1.1.1  용역 대상 공사의 개요 및 특성", level=3)

    tbl = doc.add_table(rows=1, cols=3)
    add_header_row(tbl, ["구  분", "내  용", "비  고"],
                   widths_cm=[3.5, 10.0, 3.0])
    for k, v in [("공 사 명", PH + "  공사"), ("공사 위치", PH),
                 ("공사 규모", PH), ("발  주  청", PH),
                 ("계약상대자", PH + "  주식회사"), ("감  리  단", PH),
                 ("공 사 기 간", PH + " ~ " + PH), ("계약 금액", PH + " 천원")]:
        row = tbl.add_row()
        set_cell(row.cells[0], k, bold=True, bg_color=HEADER_BG)
        set_cell(row.cells[1], v, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], "")
    doc.add_paragraph()

    add_heading(doc, "1.1.2  공사 위치도", level=3)
    add_note(doc, "  ▶ [ 공사 위치도 / 조감도 삽입 ]")
    doc.add_paragraph()

    # 1.2 계약의 현황
    add_heading(doc, "1.2  계약의 현황", level=2)
    add_heading(doc, "1.2.1  계약의 주요 용어의 정의 (공사계약 일반조건 제2조)", level=3)
    add_quote(doc,
              "「공사계약 일반조건」제2조는 '계약문서', '발주자', '계약상대자', "
              "'감독관', '공사기간' 등 주요 용어를 정의하고 있습니다.")
    doc.add_paragraph()

    add_heading(doc, "1.2.2  계약문서 (공사계약 일반조건 제3조)", level=3)
    add_quote(doc,
              "공사계약 일반조건 제3조에 따르면 계약문서는 계약서, 설계서, 유의서, "
              "공사계약 일반조건, 공사계약 특수조건, 산출내역서 등으로 구성됩니다.")
    doc.add_paragraph()

    add_heading(doc, "1.2.3  계약당사자 현황", level=3)
    tbl2 = doc.add_table(rows=1, cols=3)
    add_header_row(tbl2, ["구  분", "기 관 명", "비  고"],
                   widths_cm=[3.5, 8.0, 5.0])
    for k in ["발주청", "계약상대자 (원도급)", "건설사업관리단", "하도급사 (해당 시)"]:
        add_data_row(tbl2, [k, PH, ""])
    doc.add_paragraph()

    # 1.3 계약체결 및 변경계약의 경과
    add_heading(doc, "1.3  계약체결 및 변경계약의 경과", level=2)
    add_heading(doc, "1.3.1  원도급사 계약 현황", level=3)
    add_note(doc, "※ [단위: 천원]")
    tbl3 = doc.add_table(rows=1, cols=7)
    add_header_row(tbl3,
                   ["구  분", "계약일", "착공일", "준공일", "증감일수", "계약금액\n(천원)", "비  고"],
                   widths_cm=[2.5, 2.8, 2.8, 2.8, 2.0, 3.0, 1.6])
    for label in ["최  초", "1회 변경", "2회 변경", "3회 변경", "합  계"]:
        add_data_row(tbl3, [label, PH, PH, PH, PH, PH, ""])
    doc.add_paragraph()

    add_heading(doc, "1.3.2  하도급사 계약 현황", level=3)
    add_note(doc, "※ 해당 하도급사 수만큼 반복 작성")
    tbl4 = doc.add_table(rows=1, cols=7)
    add_header_row(tbl4,
                   ["하도급사명", "계약일", "착공일", "준공일", "증감일수", "계약금액\n(천원)", "비  고"],
                   widths_cm=[3.0, 2.5, 2.5, 2.5, 2.0, 3.0, 2.0])
    for i in range(1, 4):
        add_data_row(tbl4, [PH + f"  ({i})", PH, PH, PH, PH, PH, ""])
    doc.add_paragraph()

    add_heading(doc, "1.3.3  공사기간 변경계약 현황", level=3)
    add_note(doc, "※ 원·하도급 공사기간 변경 경위 요약")
    tbl5 = doc.add_table(rows=1, cols=5)
    add_header_row(tbl5,
                   ["변경 차수", "변경계약일", "변경 전 준공일", "변경 후 준공일", "연장 사유 요약"],
                   widths_cm=[2.5, 2.8, 3.0, 3.0, 5.2])
    for i in range(1, 5):
        add_data_row(tbl5, [f"{i}회 변경", PH, PH, PH, PH],
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    # 1.4 과업 수행 절차
    add_heading(doc, "1.4  과업 수행 절차", level=2)
    add_heading(doc, "1.4.1  수행 절차", level=3)
    add_note(doc, "  ▶ [ 과업 수행 절차도 삽입 ]")
    steps = [
        "① 기초자료 조사 및 검토  →  계약문서, 설계서, 수발신 문서 등 검토",
        "② 공기지연 귀책 분석  →  지연사유 발생 경위 및 귀책 검토",
        "③ 절차적 요건 검토  →  계약금액조정 신청 적정성 확인",
        "④ 간접비 산정  →  간접노무비, 경비, 일반관리비, 이윤",
        "⑤ 보고서 작성 및 제출",
    ]
    for s in steps:
        add_body(doc, "  " + s)
    doc.add_paragraph()

    add_heading(doc, "1.4.2  산정의 근거", level=3)
    add_body(doc,
             "본 보고서는 계약서, 수발신 문서, 관계 기초자료 등 증거자료를 근거로 하여 "
             "관련 법령 및 금액 산정의 구체적인 기준인 계약조건, 국가계약법령, "
             "「(계약예규) 정부 입찰·계약 집행기준」, 관련 판례 등을 적용하여 산정하였습니다.")


# ── 제2장: 계약의 성격 ────────────────────────────────────────────────────────

def make_ch2(doc):
    doc.add_page_break()
    add_heading(doc, "제2장.  계약의 성격", level=1)

    add_heading(doc, "2.1  국가계약법 적용 공사", level=2)
    add_body(doc,
             "본건 공사는 국가가 계약상대자로 하여 체결하는 계약으로 "
             "「국가를 당사자로 하는 계약에 관한 법률」(이하 '국가계약법')의 적용을 받습니다.")
    add_quote(doc,
              "국가계약법 제5조: \"계약은 서로 대등한 입장에서 당사자의 합의에 따라 체결하여야 하며, "
              "당사자는 계약의 내용을 신의성실의 원칙에 따라 이를 이행하여야 한다.\"")
    add_body(doc,
             "또한 동법 제5조 제3항에서는 민법 제2조 제2항의 권리남용 금지 원칙에 따라 "
             "계약상대방의 권리를 부당하게 제한하는 특약을 금지하고 있습니다.")
    doc.add_paragraph()

    add_heading(doc, "2.2  계속비 계약", level=2)
    add_body(doc,
             "본건 공사는 「국가계약법」제21조 및 같은 법 시행령에 의거, "
             "총공사금액과 연부액을 명백히 하여 계속비 계약을 체결하였습니다.")
    add_quote(doc,
              "「국가재정법」제23조에 따른 계속비 사업에 대하여는 총액과 연부액을 명백히 하여 "
              "계속비계약을 체결하도록 정하고 있습니다.")
    add_note(doc, "※ 연도별 연부액 현황:")
    tbl = doc.add_table(rows=1, cols=4)
    add_header_row(tbl, ["연  도", "연부액 (천원)", "집행액 (천원)", "비  고"],
                   widths_cm=[3.0, 4.5, 4.5, 4.5])
    for year in [PH + "년도", PH + "년도", PH + "년도"]:
        add_data_row(tbl, [year, PH, PH, ""])
    doc.add_paragraph()

    add_heading(doc, "2.3  내역입찰 공사", level=2)
    add_body(doc,
             "본건 공사는 「(계약예규) 정부 입찰·계약 집행기준」제18조에 의거, "
             "입찰서에 산출내역서를 첨부하여 입찰하는 내역입찰 대상 공사입니다.")
    add_body(doc,
             "내역입찰공사계약의 구조상 설계도면, 물량내역서, 현장설명서 등 "
             "입찰의 기초가 되는 설계 관련 문서의 작성 책임은 발주자에게 귀속되고, "
             "이에 기초한 산출내역서의 작성 책임은 계약상대자에게 있습니다.")


# ── 제3장: 귀책 분석 및 계약금액조정 검토 ────────────────────────────────────

def make_ch3(doc):
    doc.add_page_break()
    add_heading(doc, "제3장.  공기연장 귀책 분석 및 계약금액조정 검토", level=1)

    # 3.1 귀책 분석
    add_heading(doc, "3.1  공기연장 귀책 분석", level=2)
    add_heading(doc, "3.1.1  지연사유 발생 경위", level=3)
    add_body(doc,
             "본건 공사는 아래와 같은 사유로 당초 계약 준공일의 준수가 곤란하게 되어 "
             "공사기간이 연장되었습니다.")
    doc.add_paragraph()
    add_note(doc, "※ 공기지연 경위를 시간순으로 서술 (공문 근거 포함)")
    tbl = doc.add_table(rows=1, cols=5)
    add_header_row(tbl,
                   ["No.", "일  자", "공문 / 사건 내용", "발신처 → 수신처", "비  고"],
                   widths_cm=[1.0, 2.8, 8.0, 4.0, 1.7])
    for i in range(1, 7):
        add_data_row(tbl,
                     [str(i), PH, PH, PH + " → " + PH, ""],
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    add_heading(doc, "3.1.2  공기지연 귀책사유 검토", level=3)
    add_body(doc,
             "앞서 살펴본 바와 같이, 본건 공사의 공기지연 사유는 아래와 같이 정리되며, "
             "이는 공사계약 일반조건 제26조 제3항에서 정한 '당공사의 책임으로 착공이 지연되거나 "
             "시공이 중단된 경우'에 해당하여 계약상대자의 책임없는 사유에 해당합니다.")
    doc.add_paragraph()
    tbl2 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl2,
                   ["귀책 주체", "지연 사유", "지연 기간", "근거 공문"],
                   widths_cm=[3.0, 7.0, 3.0, 3.5])
    for label in ["발주자 귀책", "인허가 지연", "민원·설계변경", "불가항력"]:
        add_data_row(tbl2, [label, PH, PH + " 일", PH],
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    add_body(doc,
             "따라서 공사계약 일반조건 제27조 제1항에 의거 계약상대자는 계약기간의 연장을 "
             "청구할 수 있으며, 제27조 제4항 및 제23조에 의하여 실비를 초과하지 아니하는 "
             "범위 안에서 계약금액의 조정을 청구할 수 있습니다.")
    doc.add_paragraph()

    # 3.2 계약금액조정 관련 검토
    add_heading(doc, "3.2  계약금액조정 관련 검토", level=2)
    add_heading(doc, "3.2.1  분석 개요", level=3)
    add_body(doc,
             "공사계약 일반조건 제20조 제10항에 의하면 계약금액조정 청구는 "
             "준공대가 수령 전까지 조정신청을 하여야 합니다. "
             "이에 따라 본건 공사의 간접비 청구가 절차적 요건을 준수하였는지 검토합니다.")
    doc.add_paragraph()

    add_heading(doc, "3.2.2  계약금액조정 신청 관련 문서", level=3)
    tbl3 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl3,
                   ["No.", "신청 일자", "문서 제목", "비  고"],
                   widths_cm=[1.0, 3.0, 10.0, 2.5])
    for i in range(1, 4):
        add_data_row(tbl3, [str(i), PH, PH, ""],
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    add_heading(doc, "3.2.3  검토 의견", level=3)
    add_body(doc,
             "본건 공사의 수발신 문서 등을 검토한 결과, 계약상대자는 각 공기지연 사유 발생 시 "
             "준공일 이전에 지체없이 계약기간 연장 및 계약금액 조정을 신청하였음을 확인하였습니다.")
    add_body(doc,
             "따라서 본건 공기연장 간접비 청구는 계약상 절차적 요건을 준수한 것으로 판단됩니다.")
    doc.add_paragraph()

    # 3.3 청구의 근거
    add_heading(doc, "3.3  공기연장 간접비 청구의 근거", level=2)

    add_heading(doc, "3.3.1  계약에 따른 청구의 근거", level=3)
    add_body(doc,
             "공사계약 일반조건 제23조 제1항 및 제27조 제4항에 따라 공사기간이 연장된 경우 "
             "실비를 초과하지 않는 범위 안에서 계약금액을 조정할 수 있습니다.")
    add_quote(doc,
              "공사계약 일반조건 제27조 제4항: \"제2항의 규정에 의하여 계약기간을 연장한 경우에는 "
              "제23조의 규정에 의하여 그 변경된 내용에 따라 실비를 초과하지 아니하는 범위 안에서 "
              "계약금액을 조정한다.\"")
    doc.add_paragraph()

    add_heading(doc, "3.3.2  관련 법령에 따른 청구의 근거", level=3)
    add_body(doc,
             "국가계약법 제19조 및 같은 법 시행령 제66조 제1항에 의거, "
             "공사기간 변경으로 계약금액 조정이 필요한 경우 실비 범위 내에서 조정합니다.")
    add_quote(doc,
              "국가계약법 시행령 제66조 제1항: \"공사기간·운반거리의 변경 등 기타 계약내용의 변경으로 "
              "인하여 계약금액을 조정하여야 할 필요가 있는 경우에는 그 변경된 내용에 따라 "
              "실비를 초과하지 아니하는 범위 안에서 이를 조정한다.\"")
    doc.add_paragraph()

    add_heading(doc, "3.3.3  사정변경 원칙에 따른 청구의 근거", level=3)
    add_body(doc,
             "계약 성립 당시 기초가 된 사정이 변경되어 기존 계약 내용을 그대로 유지하는 것이 "
             "신의성실 원칙에 반하는 경우, 민법상 사정변경의 원칙에 따라 계약 내용의 변경 또는 "
             "추가 비용 청구가 가능합니다.")
    doc.add_paragraph()

    add_heading(doc, "3.3.4  청구의 근거 검토", level=3)
    add_body(doc,
             "본건 공기연장 간접비는 계약조건(공사계약 일반조건 제23조, 제27조), "
             "관련 법령(국가계약법 시행령 제66조), 사정변경 원칙에 근거할 때 "
             "계약금액조정 대상으로 인정되며, 그 청구의 근거는 정당합니다.")


# ── 제4장: 간접비 산정 ────────────────────────────────────────────────────────

def make_ch4(doc):
    doc.add_page_break()
    add_heading(doc, "제4장.  공기연장 간접비 산정", level=1)

    # 4.1 대상 기간 검토
    add_heading(doc, "4.1  공기연장 간접비 산정 대상 기간 검토", level=2)
    add_heading(doc, "4.1.1  공기연장 일수의 산정방식", level=3)
    add_body(doc,
             "공기연장 일수는 변경계약서에 명시된 준공기한을 기준으로 산정하며, "
             "역일(曆日) 기준을 적용합니다.")
    doc.add_paragraph()

    add_heading(doc, "4.1.2  공기연장 일수 산정", level=3)
    tbl = doc.add_table(rows=1, cols=5)
    add_header_row(tbl,
                   ["구  분", "착공일", "준공일", "공사일수", "연장일수"],
                   widths_cm=[3.0, 3.5, 3.5, 2.5, 4.0])
    for label in ["당초 계약", "1회 변경", "2회 변경", "연장 합계"]:
        add_data_row(tbl, [label, PH, PH, PH + " 일", PH + " 일"])
    doc.add_paragraph()

    # 4.2 산정 방식
    add_heading(doc, "4.2  공기연장 간접비 산정 방식", level=2)

    add_heading(doc, "4.2.1  간접노무비 산정 방식", level=3)
    add_body(doc,
             "간접노무비는 공기연장 기간 중 현장 상주 관리인원의 노무비로 산정하며, "
             "노무비 단가는 아래 항목의 합계로 구성됩니다.")
    for item in [
        ("기본급",       "실제 지급 기본급 (급여명세서 기준)"),
        ("제수당",       "식대, 교통비, 직책수당 등 (급여명세서 기준)"),
        ("상여금",       "연간 상여금 ÷ 12 (지급 규정 또는 취업규칙 기준)"),
        ("퇴직급여충당금", "월 급여 합계 ÷ 12 (근로자퇴직급여보장법)"),
    ]:
        add_body(doc, f"  ▸ {item[0]}: {item[1]}")
    doc.add_paragraph()

    add_heading(doc, "4.2.2  경비 산정 방식", level=3)
    add_heading(doc, "직접계상비목", level=4)
    add_body(doc,
             "「(계약예규) 정부 입찰·계약 집행기준」제72조, 제73조에 따라 "
             "경비지출 관련 계약서, 요금고지서, 영수증 등 객관적인 자료에 의하여 확인된 금액을 산정합니다.")
    for item in ["전력비·수도광열비", "여비·교통비·통신비", "지급임차료",
                 "복리후생비", "소모품비", "세금과공과", "도서인쇄비",
                 "지급수수료", "국민건강·연금보험"]:
        add_body(doc, f"  • {item}")
    doc.add_paragraph()

    add_heading(doc, "승률계상비목", level=4)
    add_body(doc,
             "「(계약예규) 정부 입찰·계약 집행기준」제73조 제3항에 따라 "
             "해당 비목의 기준 비목 합계에 산출내역서상 요율을 곱하여 산정합니다.")
    tbl2 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl2,
                   ["비  목", "산  출  방  식", "적용 요율", "비  고"],
                   widths_cm=[3.5, 7.0, 2.5, 3.5])
    for item, formula in [("산재보험료", "간접노무비  ×  요율"),
                          ("고용보험료", "간접노무비  ×  요율")]:
        row = tbl2.add_row()
        set_cell(row.cells[0], item, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[1], formula, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], PH + " %")
        set_cell(row.cells[3], "산출내역서 기준")
    doc.add_paragraph()

    add_heading(doc, "4.2.3  일반관리비 및 이윤 산정 방식", level=3)
    tbl3 = doc.add_table(rows=1, cols=3)
    add_header_row(tbl3, ["비  목", "산  출  방  식", "적용 요율"],
                   widths_cm=[3.0, 9.0, 4.5])
    for item, formula in [
        ("일반관리비", "(간접노무비 + 경비)  ×  요율"),
        ("이  윤",    "(간접노무비 + 경비 + 일반관리비)  ×  요율"),
    ]:
        row = tbl3.add_row()
        set_cell(row.cells[0], item, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[1], formula, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], PH + " %  (산출내역서 기준)")
    doc.add_paragraph()

    add_heading(doc, "4.2.4  보증수수료 산정 방식", level=3)
    add_body(doc,
             "보증수수료는 「(계약예규) 정부 입찰·계약 집행기준」제73조 제4항에 의거, "
             "보증수수료 영수증 등 객관적인 자료에 의하여 확인된 금액을 기준으로 산정합니다.")
    doc.add_paragraph()

    # 4.3 원도급사 산정 결과
    add_heading(doc, "4.3  원도급사 산정 결과", level=2)
    make_result_section(doc, "원도급사")

    # 4.4 하도급사 산정 결과
    add_heading(doc, "4.4  하도급사 산정 결과", level=2)
    add_note(doc, "※ 하도급사별로 아래 구조를 반복 작성. 해당 없으면 절 삭제.")
    add_note(doc, "※ 하도급사 산출내역서 미확보 시: "
                  "\"원도급사 산출내역서 요율과 동일하게 적용\" 명기")

    # 전체 집계 먼저
    add_heading(doc, "하도급사 전체 집계표", level=3)
    add_note(doc, "※ [단위: 원, 영세율]")
    tbl_sub = doc.add_table(rows=1, cols=5)
    add_header_row(tbl_sub,
                   ["하도급사명", "간접노무비", "경  비", "일반관리비", "이  윤"],
                   widths_cm=[4.0, 3.5, 3.5, 3.5, 3.0])
    for i in range(1, 4):
        add_data_row(tbl_sub, [PH + f"  ({i})", PH, PH, PH, PH])
    add_data_row(tbl_sub, ["합  계", PH, PH, PH, PH])
    doc.add_paragraph()

    # 개별 하도급사
    add_heading(doc, PH + "  (하도급사 1 — 반복)", level=3)
    make_result_section(doc, PH + " (하도급사명)")


# ── 제5장: 결론 ──────────────────────────────────────────────────────────────

def make_ch5(doc):
    doc.add_page_break()
    add_heading(doc, "제5장.  결론", level=1)
    add_body(doc,
             "본건 공사의 공기연장 사유는 " + PH + " 등으로, "
             "이는 공사계약 일반조건 제26조 제3항에 따라 계약상대자의 책임없는 사유에 해당합니다.")
    add_body(doc,
             "공사계약 일반조건 제27조 제4항 및 제23조, 국가계약법 시행령 제66조 제1항에 의거, "
             "연장된 공사기간에 대한 간접비를 아래와 같이 청구합니다.")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=4)
    add_header_row(tbl,
                   ["구  분", "간접비 합계 (원)", "산정 기간", "비  고"],
                   widths_cm=[3.0, 5.5, 5.0, 3.0])
    for label in ["원도급사", "하도급사", "최 종 청구액"]:
        bold = label == "최 종 청구액"
        row = tbl.add_row()
        for j, v in enumerate([label, PH, PH, "VAT 별도" if bold else ""]):
            set_cell(row.cells[j], v, bold=bold,
                     bg_color=(HEADER_BG if bold else None))
    doc.add_paragraph()
    add_body(doc,
             "상기 금액은 실비를 초과하지 아니하는 범위 안에서 산정된 금액으로, "
             "관련 증빙자료는 첨부자료에 수록하였습니다.")


# ── 제6장: 첨부자료 ──────────────────────────────────────────────────────────

def make_ch6(doc):
    doc.add_page_break()
    add_heading(doc, "제6장.  첨부자료", level=1)
    add_note(doc, "※ 아래 각 절 순서대로 자료를 첨부합니다.")
    doc.add_paragraph()
    sections = [
        ("6.1  공기연장 간접비 산정근거",
         "급여명세서, 계좌이체 내역, 노무비 상세 계산서, 보험요율 확인서 등"),
        ("6.2  계약문서",
         "공사도급계약서, 산출내역서, 변경계약서 (전 차수)"),
        ("6.3  수발신문서",
         "귀책사유 관련 공문 원본 (발신/수신 순)"),
        ("6.4  판례 및 관련 조항",
         "참조 판례, 국가계약법 관련 조항 발췌"),
        ("6.5  공기연장 간접비 증빙자료",
         "전력·수도 고지서, 임대차계약서, 영수증, 보증수수료 영수증 등"),
    ]
    for title, desc in sections:
        add_heading(doc, title, level=2)
        add_note(doc, "  ▶ " + desc)
        add_body(doc, "  [ 자료 첨부 위치 ]")
        doc.add_paragraph()


# ── 메인 ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_MAIN)

    make_cover(doc)
    make_submission(doc)
    make_summary(doc)
    make_toc(doc)
    make_ch1(doc)
    make_ch2(doc)
    make_ch3(doc)
    make_ch4(doc)
    make_ch5(doc)
    make_ch6(doc)

    out_path = Path("reference/보고서_템플릿_B_상세형(국가계약).docx")
    doc.save(str(out_path))
    print(f"저장 완료: {out_path}")
    print(f"  ※ 【 입력 】 표시 부분을 실제 값으로 교체하여 사용하세요.")


if __name__ == "__main__":
    main()
