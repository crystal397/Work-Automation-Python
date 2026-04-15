"""
귀책분석 데이터 JSON → docx 생성 (데이터 드리븐)

입력: output/귀책분석_data.json
출력: output/02_귀책분석_[프로젝트명]_[날짜].docx

JSON 스키마 (모든 필드는 선택적이며 기본값이 있음):
  project_name            str
  chapter_heading         str   (기본: "제2장  귀책분석")
  section_3_1_heading     str   (기본: "3.1. 공기연장 귀책 분석")
  section_3_1_1_heading   str   (기본: "3.1.1. 지연사유 발생 경위")
  section_3_1_2_heading   str   (기본: "3.1.2. 공기지연 귀책 사유 검토")
  table_intro_sentence    str   (기본 문구 포함)
  background_paragraphs   list[str]      3.1.1 도입부 서술 단락
  items                   list[dict]     공문 목록 (show_in_table 플래그 포함)
  detail_narratives       list[dict]     공문 목록 표 이후 상세 서술 단락 블록
  pre_diagram_paragraphs  list[str]      3.1.2 도식표 전 서술 단락
  accountability_diagram  list[dict]     귀책사유 도식표 행
  conclusion_paragraphs   list[str]      3.1.2 결론 서술 단락
  summary                 str            종합 단락
"""

from __future__ import annotations

import json
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 폰트·색상 상수 ──────────────────────────────────────────────────────────
FONT_KO = "맑은 고딕"

# 색상 (hex, 6자리)
COLOR_HEADER_GRAY   = "D9D9D9"   # 표 헤더 배경
COLOR_ALT_ROW       = "F2F2F2"   # 짝수 행 배경
COLOR_SECTION_BG    = "E8F0F8"   # 소절 헤딩 배경 (연한 청색)

# A4 세로 유효 너비 (좌우 여백 각 2.5 cm 기준)
PAGE_WIDTH_CM = 16.5


# ── 내부 유틸 ────────────────────────────────────────────────────────────────

def _set_font(run, size_pt: int, bold: bool = False, color_hex: str | None = None):
    """run 에 한글 폰트·크기·굵기·색상 적용"""
    run.font.name = FONT_KO
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_KO)
    rPr.insert(0, rFonts)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold: bool = False,
               size_pt: int = 9, align: str = "left",
               bg_hex: str | None = None):
    """셀 텍스트 설정 (줄바꿈 \n 지원)"""
    cell.text = ""
    if bg_hex:
        _set_cell_bg(cell, bg_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }

    lines = str(text).split("\n") if text else [""]
    for idx, line in enumerate(lines):
        if idx == 0:
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()
        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(line)
        _set_font(run, size_pt, bold=bold)


def _set_col_widths(table, widths_cm: list[float]):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _add_paragraph(doc: Document, text: str,
                   size_pt: int = 10, bold: bool = False,
                   align: str = "left",
                   space_before_pt: int = 0,
                   space_after_pt: int = 4) -> None:
    """본문 서술 단락 추가"""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after  = Pt(space_after_pt)

    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    _set_font(run, size_pt, bold=bold)


def _add_heading(doc: Document, text: str,
                 size_pt: int = 11, bold: bool = True,
                 space_before_pt: int = 8,
                 space_after_pt: int = 4) -> None:
    """섹션 헤딩 단락 추가"""
    _add_paragraph(doc, text, size_pt=size_pt, bold=bold,
                   space_before_pt=space_before_pt,
                   space_after_pt=space_after_pt)


# ── 공문 목록 표 (5열) ────────────────────────────────────────────────────────
# 열: 일자(2.2) | 문서번호(3.8) | 제목(auto) | 발신(2.5) | 수신(2.5)
_CORR_FIXED_CM = [2.2, 3.8, 0.0, 2.5, 2.5]   # 0.0 = 나머지 채움

def _build_correspondence_table(doc: Document, table_items: list[dict]) -> None:
    """5열 공문 목록 표 생성"""
    headers = ["일자", "문서번호", "제목", "발신", "수신"]
    n_cols  = len(headers)
    n_rows  = 1 + len(table_items)

    fixed_total = sum(w for w in _CORR_FIXED_CM if w > 0)
    auto_width  = PAGE_WIDTH_CM - fixed_total
    col_widths  = [auto_width if w == 0 else w for w in _CORR_FIXED_CM]

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h,
                   bold=True, size_pt=9, align="center",
                   bg_hex=COLOR_HEADER_GRAY)

    # 데이터
    for row_idx, item in enumerate(table_items):
        row  = table.rows[row_idx + 1]
        bg   = COLOR_ALT_ROW if row_idx % 2 == 1 else None
        vals = [
            item.get("date", ""),
            item.get("doc_number", ""),
            item.get("subject", ""),
            item.get("sender", ""),
            item.get("receiver", ""),
        ]
        aligns = ["center", "left", "left", "center", "center"]
        for col_idx, (val, aln) in enumerate(zip(vals, aligns)):
            _cell_text(row.cells[col_idx], val,
                       size_pt=9, align=aln, bg_hex=bg)

    _set_col_widths(table, col_widths)


# ── 귀책사유 도식표 (3열) ─────────────────────────────────────────────────────
# 열: 공기지연 사유(auto) | 관련 근거(4.0) | 비용부담자(2.5)
_DIAG_FIXED_CM = [0.0, 4.0, 2.5]

def _build_accountability_diagram(doc: Document,
                                  rows_data: list[dict]) -> None:
    """귀책사유 도식표 생성"""
    headers = ["공기지연 사유", "관련 근거", "비용부담자"]
    n_cols  = len(headers)
    n_rows  = 1 + len(rows_data)

    fixed_total = sum(w for w in _DIAG_FIXED_CM if w > 0)
    auto_width  = PAGE_WIDTH_CM - fixed_total
    col_widths  = [auto_width if w == 0 else w for w in _DIAG_FIXED_CM]

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h,
                   bold=True, size_pt=9, align="center",
                   bg_hex=COLOR_HEADER_GRAY)

    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        bg  = COLOR_ALT_ROW if row_idx % 2 == 1 else None
        # 공기지연 사유: delay_cause 또는 cause 키 모두 허용 + 지연일수 병기
        cause_text = row_data.get("delay_cause") or row_data.get("cause", "")
        delay_days = row_data.get("delay_days", 0)
        if delay_days and int(delay_days) > 0:
            cause_text = f"{cause_text} ({delay_days}일)"
        _cell_text(row.cells[0], cause_text,
                   size_pt=9, align="left", bg_hex=bg)
        _cell_text(row.cells[1], row_data.get("basis", ""),
                   size_pt=9, align="left", bg_hex=bg)
        _cell_text(row.cells[2], row_data.get("responsible_party", ""),
                   size_pt=9, align="center", bold=True, bg_hex=bg)

    _set_col_widths(table, col_widths)


# ── 스키마 검증 ──────────────────────────────────────────────────────────────

# 잘못된 필드명 → 올바른 필드명 안내
_FORBIDDEN_FIELDS = {
    "intro_paragraph":   "background_paragraphs (리스트)",
    "contract_info":     "(제거 — 내용은 background_paragraphs에 포함)",
    "chapter_title":     "(제거 — 사용하지 않음)",
    "section_title":     "(제거 — 사용하지 않음)",
}
_FORBIDDEN_NARRATIVE_KEYS = {
    "body":       "paragraphs (리스트)",
    "section_no": "(제거 — 사용하지 않음)",
}
_FORBIDDEN_ITEM_KEYS = {
    "summary": "causal_description",
}

_REQUIRED_FIELDS = [
    "project_name",
    "total_delay_days",
    "background_paragraphs",
    "items",
    "detail_narratives",
    "pre_diagram_paragraphs",
    "accountability_diagram",
    "conclusion_paragraphs",
    "summary",
]


def _validate_data(data: dict, data_path: Path) -> None:
    """
    귀책분석_data.json 스키마 검증.
    문제 발견 시 경고를 출력하고, 치명적 오류는 ValueError 발생.
    """
    warnings: list[str] = []
    errors:   list[str] = []

    # 1) 금지 최상위 필드
    for bad, good in _FORBIDDEN_FIELDS.items():
        if bad in data:
            errors.append(
                f"  [오류] 잘못된 필드 '{bad}' 사용 — 올바른 필드: {good}\n"
                f"     → 이 필드의 내용은 보고서에 출력되지 않습니다."
            )

    # 2) 필수 필드 누락 / 빈값
    for field in _REQUIRED_FIELDS:
        val = data.get(field)
        if val is None:
            errors.append(f"  [오류] 필수 필드 누락: '{field}'")
        elif isinstance(val, list) and len(val) == 0:
            errors.append(f"  [오류] 필수 필드가 빈 배열 []: '{field}'  → 내용을 채워야 합니다.")
        elif isinstance(val, str) and val.strip() == "":
            errors.append(f"  [오류] 필수 필드가 빈 문자열: '{field}'  → 내용을 채워야 합니다.")

    # 3) total_delay_days 타입 및 값
    tdd = data.get("total_delay_days")
    if tdd is not None and (not isinstance(tdd, int) or tdd <= 0):
        errors.append(
            f"  [오류] 'total_delay_days' 값이 올바르지 않습니다: {tdd!r}\n"
            f"     → 변경계약서 연장일수 합산 정수(양수)를 기재하세요."
        )

    # 4) conclusion_paragraphs 단락 수
    cp = data.get("conclusion_paragraphs", [])
    if isinstance(cp, list) and 0 < len(cp) < 2:
        warnings.append(
            f"  [경고] 'conclusion_paragraphs' 단락이 {len(cp)}개입니다. 2단락 이상을 권장합니다."
        )

    # 5) detail_narratives 하위 키 검증
    narratives = data.get("detail_narratives", [])
    if isinstance(narratives, list):
        for i, block in enumerate(narratives):
            if not isinstance(block, dict):
                continue
            for bad_key, good_key in _FORBIDDEN_NARRATIVE_KEYS.items():
                if bad_key in block:
                    errors.append(
                        f"  [오류] detail_narratives[{i}]에 잘못된 키 '{bad_key}' 사용\n"
                        f"     → 올바른 키: {good_key}\n"
                        f"     → 이 블록의 내용은 보고서에 출력되지 않습니다."
                    )
            if "paragraphs" not in block and "body" not in block:
                paras = block.get("paragraphs")
                if paras is None:
                    warnings.append(
                        f"  [경고] detail_narratives[{i}]에 'paragraphs' 키가 없습니다 (제목: {block.get('title') or block.get('label', '?')})"
                    )

    # 6) items 하위 키 검증
    items_list = data.get("items", [])
    if isinstance(items_list, list):
        for i, item in enumerate(items_list):
            if not isinstance(item, dict):
                continue
            for bad_key, good_key in _FORBIDDEN_ITEM_KEYS.items():
                if bad_key in item and "causal_description" not in item:
                    warnings.append(
                        f"  [경고] items[{i}]에 '{bad_key}' 사용 ('{good_key}' 권장) — 표 렌더링에는 영향 없음"
                    )

    # 7) accountability_diagram 일수 합계 교차검증
    diagram = data.get("accountability_diagram", [])
    if isinstance(diagram, list) and len(diagram) > 0:
        diagram_total = sum(
            item.get("delay_days", 0)
            for item in diagram
            if isinstance(item, dict) and isinstance(item.get("delay_days"), (int, float))
        )
        tdd = data.get("total_delay_days")
        if isinstance(tdd, int) and tdd > 0 and int(diagram_total) != tdd:
            errors.append(
                f"  [오류] accountability_diagram 항목 delay_days 합계({int(diagram_total)}일)가\n"
                f"     total_delay_days({tdd}일)와 다릅니다.\n"
                f"     → 변경계약서 차수별 연장일수를 다시 확인하세요."
            )

    # 8) items 날짜 시간순 검증
    def _parse_date(s: str) -> date | None:
        m = re.match(r"(\d{4})\.(\d{1,2})\.(\d{1,2})", s.strip()) if s else None
        if m:
            try:
                return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            except ValueError:
                return None
        return None

    if isinstance(items_list, list):
        prev_d: date | None = None
        prev_i = -1
        prev_date_str = ""
        for i, item in enumerate(items_list):
            if not isinstance(item, dict):
                continue
            d = _parse_date(item.get("date", ""))
            if d is None:
                continue
            if prev_d is not None and d < prev_d:
                warnings.append(
                    f"  [경고] items[{i}] 날짜({item['date']})가 items[{prev_i}]({prev_date_str})보다 앞입니다.\n"
                    f"     → items는 시행일(발신일) 기준 시간순으로 정렬되어야 합니다."
                )
            prev_d = d
            prev_i = i
            prev_date_str = item.get("date", "")

    # 9) items에서 시행일/접수일 혼용 가능성 감지
    #    같은 doc_number(공문번호)를 가진 항목이 1~2일 간격으로 2개 존재하면 경고
    if isinstance(items_list, list):
        doc_map: dict[str, list[tuple[int, str]]] = {}
        for i, item in enumerate(items_list):
            if not isinstance(item, dict):
                continue
            doc_num = item.get("doc_number", "").strip()
            date_str = item.get("date", "")
            if doc_num:
                doc_map.setdefault(doc_num, []).append((i, date_str))
        for doc_num, entries in doc_map.items():
            if len(entries) == 2:
                d0 = _parse_date(entries[0][1])
                d1 = _parse_date(entries[1][1])
                if d0 and d1:
                    diff = abs((d1 - d0).days)
                    if 1 <= diff <= 3:
                        warnings.append(
                            f"  [경고] 공문번호 '{doc_num}'가 {diff}일 간격으로 2번 등장합니다.\n"
                            f"     → items[{entries[0][0]}]({entries[0][1]})과 items[{entries[1][0]}]({entries[1][1]})\n"
                            f"     → 시행일(발신일)과 접수일이 분리됐을 가능성 — date 필드는 시행일 기준으로 통일,\n"
                            f"       접수일은 note 필드에 '접수: YYYY.MM.DD' 형식으로 병기하세요."
                        )

    # 10) causal_description에서 시행일/접수일 혼용 감지
    #     "시행: ...(YYYY.MM.DD.)" 와 "접수: ...(YYYY.MM.DD.)" 가 동시에 언급된 항목에서
    #     items[].date가 시행일이 아닌 접수일을 쓰고 있으면 경고
    _SIHAENG_RE = re.compile(r"시행\s*[：:].+?\((\d{4}\.\d{1,2}\.\d{1,2})\.?\)")
    _JEOPSU_RE  = re.compile(r"접수\s*[：:].+?\((\d{4}\.\d{1,2}\.\d{1,2})\.?\)")
    if isinstance(items_list, list):
        for i, item in enumerate(items_list):
            if not isinstance(item, dict):
                continue
            desc = item.get("causal_description", "") or ""
            m_s = _SIHAENG_RE.search(desc)
            m_j = _JEOPSU_RE.search(desc)
            if m_s and m_j:
                sihaeng_str = m_s.group(1)
                jeopsu_str  = m_j.group(1)
                item_date   = item.get("date", "")
                # date가 시행일이 아니라 접수일과 일치하는 경우
                if (item_date.replace(" ", "") == jeopsu_str.replace(" ", "")
                        and item_date.replace(" ", "") != sihaeng_str.replace(" ", "")):
                    warnings.append(
                        f"  [경고] items[{i}]({item_date}) date 필드가 접수일({jeopsu_str})을 사용하고 있습니다.\n"
                        f"     → causal_description에 시행일({sihaeng_str})이 별도 기재되어 있습니다.\n"
                        f"     → date 필드를 시행일({sihaeng_str})로 수정하고,\n"
                        f"       note 필드에 '접수: {jeopsu_str}' 형식으로 병기하세요."
                    )

    # 11) detail_narratives 연도 혼재 경고
    #     공기연장 서술 블록에서 연도를 추출하여, 같은 블록 안에 서로 다른 연도의 날짜가
    #     3개 이상 등장하면 "원인 발생 기간과 간접비 청구 기간이 섞였을 가능성" 경고
    _DATE_YEAR_RE = re.compile(r"\b(20\d{2})\.\s*\d{1,2}\.\s*\d{1,2}")
    if isinstance(narratives, list):
        for i, block in enumerate(narratives):
            if not isinstance(block, dict):
                continue
            paras = block.get("paragraphs", [])
            if not isinstance(paras, list):
                continue
            full_text = " ".join(str(p) for p in paras)
            years_found = set(_DATE_YEAR_RE.findall(full_text))
            if len(years_found) >= 2:
                label = block.get("title") or block.get("label", f"[{i}]")
                warnings.append(
                    f"  [경고] detail_narratives '{label}' 블록에 {sorted(years_found)}년 날짜가 혼재합니다.\n"
                    f"     → 이것이 의도된 경우(원인 발생 기간 ≠ 간접비 청구 기간)라면 정상입니다.\n"
                    f"     → 서술이 두 개념을 명확히 구분하고 있는지 직접 확인하세요."
                )

    # 출력
    has_error = bool(errors)
    if warnings or errors:
        print()
        print("=" * 60)
        print(f"  귀책분석_data.json 스키마 검증 — {data_path.parent.name}")
        print("=" * 60)
        for msg in errors:
            print(msg)
        for msg in warnings:
            print(msg)
        print("=" * 60)

    if has_error:
        raise ValueError(
            f"\n귀책분석_data.json에 오류가 있습니다. 위 내용을 확인하고 수정한 뒤 다시 실행하세요.\n"
            f"올바른 필드명은 사용안내.txt 또는 귀책분석_schema.json을 참고하세요."
        )


# ── 메인 generate 함수 ────────────────────────────────────────────────────────

def generate(output_dir: Path, project_name: str = "") -> Path:
    """
    귀책분석_data.json 을 읽어 구조화된 docx 를 생성.

    문서 구조:
      [장 제목]
      3.1.  공기연장 귀책 분석
      3.1.1. 지연사유 발생 경위
        - background_paragraphs (서술 단락)
        - 공문 목록 표 (5열, show_in_table=True 항목)
        - detail_narratives (상세 서술 단락 블록 — DL-0351, 변경계약, DL-0805 등)
      3.1.2. 공기지연 귀책 사유 검토
        - pre_diagram_paragraphs (도식표 전 서술)
        - 귀책사유 도식표
        - conclusion_paragraphs (결론 서술)
      [종합] summary
    """
    data_path = output_dir / "귀책분석_data.json"
    if not data_path.exists():
        raise FileNotFoundError(
            f"데이터 파일이 없습니다: {data_path}\n"
            "준비 단계를 먼저 완료하세요."
        )

    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    # ── 스키마 검증 ────────────────────────────────────────────────────────────
    _validate_data(data, data_path)

    pname   = project_name or data.get("project_name", "")
    items   = data.get("items", [])
    summary = data.get("summary", "").strip()

    if not items:
        raise ValueError("귀책분석_data.json 에 items 가 없습니다.")

    # show_in_table=True 인 항목만 공문 목록 표에 사용
    table_items = [it for it in items if it.get("show_in_table", True)]

    # ── 섹션 제목 (JSON 오버라이드 가능, 없으면 기본값) ──────────────────────
    chapter_heading    = data.get("chapter_heading",       "제3장  귀책분석")
    sec_3_1            = data.get("section_3_1_heading",   "3.1. 공기연장 귀책 분석")
    sec_3_1_1          = data.get("section_3_1_1_heading", "3.1.1. 지연사유 발생 경위")
    sec_3_1_2          = data.get("section_3_1_2_heading", "3.1.2. 공기지연 귀책 사유 검토")
    table_intro        = data.get(
        "table_intro_sentence",
        "이와 관련해 발주자 및 계약상대자, 건설사업관리단 사이에 주고받은 문서 및 "
        "공사지연 사유를 확인할 수 있는 주요 자료를 시간 순서로 정리하면 아래와 같습니다.",
    )

    # ── 문서 생성 ────────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── 장 제목 ──────────────────────────────────────────────────────────────
    _add_heading(doc, chapter_heading,
                 size_pt=13, bold=True,
                 space_before_pt=0, space_after_pt=6)

    if pname:
        _add_paragraph(doc, f"◈ 대상 공사: {pname}",
                       size_pt=10, space_after_pt=2)

    total_delay_days = data.get("total_delay_days")
    if total_delay_days is not None and total_delay_days != 0:
        _add_paragraph(doc, f"◈ 공기연장 일수: {total_delay_days}일",
                       size_pt=10, space_after_pt=8)

    # ── 3.1 ──────────────────────────────────────────────────────────────────
    _add_heading(doc, sec_3_1,
                 size_pt=11, bold=True,
                 space_before_pt=6, space_after_pt=4)

    # ── 3.1.1 ────────────────────────────────────────────────────────────────
    _add_heading(doc, sec_3_1_1,
                 size_pt=10, bold=True,
                 space_before_pt=4, space_after_pt=4)

    # 배경 서술 단락
    for para_text in data.get("background_paragraphs", []):
        _add_paragraph(doc, para_text, size_pt=10, align="justify",
                       space_after_pt=4)

    # 공문 목록 표 도입 문장
    if table_items:
        _add_paragraph(doc, table_intro, size_pt=10, align="justify",
                       space_before_pt=4, space_after_pt=4)
        _build_correspondence_table(doc, table_items)

    # 상세 서술 단락 블록 (DL-0351 서술, 변경계약 체결, DL-0805 서술 등)
    for block in data.get("detail_narratives", []):
        title = (block.get("title") or block.get("label") or "").strip()
        if title:
            _add_heading(doc, f"▶ {title}",
                         size_pt=10, bold=True,
                         space_before_pt=6, space_after_pt=3)
        for para_text in block.get("paragraphs", []):
            _add_paragraph(doc, para_text, size_pt=10, align="justify",
                           space_before_pt=4, space_after_pt=4)

    # ── 3.1.2 ────────────────────────────────────────────────────────────────
    _add_heading(doc, sec_3_1_2,
                 size_pt=10, bold=True,
                 space_before_pt=8, space_after_pt=4)

    # 도식표 전 서술
    for para_text in data.get("pre_diagram_paragraphs", []):
        _add_paragraph(doc, para_text, size_pt=10, align="justify",
                       space_after_pt=4)

    # 귀책사유 도식표
    diagram = data.get("accountability_diagram", [])
    if diagram:
        _add_paragraph(doc, "", size_pt=4, space_after_pt=2)  # 여백
        _build_accountability_diagram(doc, diagram)

    # 결론 서술
    for para_text in data.get("conclusion_paragraphs", []):
        _add_paragraph(doc, para_text, size_pt=10, align="justify",
                       space_before_pt=4, space_after_pt=4)

    # ── 종합 단락 ────────────────────────────────────────────────────────────
    if summary:
        # [종합] 접두어가 이미 포함된 경우 중복 출력 방지
        summary_clean = summary
        for prefix in ("[종합]", "[ 종합 ]"):
            summary_clean = summary_clean.lstrip()
            if summary_clean.startswith(prefix):
                summary_clean = summary_clean[len(prefix):].lstrip()
        _add_paragraph(doc, "", size_pt=4, space_after_pt=2)
        _add_paragraph(doc, "[ 종합 ]  " + summary_clean,
                       size_pt=10, align="justify",
                       space_before_pt=4, space_after_pt=4)

    # ── 출처 파일 목록 (scan_result.json 조회) ──────────────────────────────
    _append_source_list(doc, output_dir, items)

    # ── 저장 ────────────────────────────────────────────────────────────────
    today    = datetime.today().strftime("%Y%m%d")
    safe     = re.sub(r'[\\/:*?"<>|]', "_", pname) if pname else "프로젝트"
    out_name = f"02_귀책분석_{safe}_{today}.docx"
    out_path = output_dir / out_name

    doc.save(str(out_path))
    print(f"\n생성 완료: {out_path}")
    return out_path


def _append_source_list(doc: Document, output_dir: Path, items: list[dict]) -> None:
    """
    docx 말미에 [참조 원본 파일 목록] 섹션 추가.
    data.json items의 scan_no 필드로 scan_result.json의 file_path를 조회.
    scan_no가 없는 항목은 파일명(source_file)만 표시.
    """
    scan_result_path = output_dir / "scan_result.json"
    if not scan_result_path.exists():
        return

    with open(scan_result_path, encoding="utf-8") as f:
        scan_data = json.load(f)
    scan_items = scan_data.get("items", [])

    import re as _re

    def _norm_date(s: str) -> str:
        """날짜 문자열에서 숫자만 추출 후 8자리(YYYYMMDD) 정규화"""
        nums = _re.findall(r"\d+", str(s))
        if len(nums) >= 3:
            y, m, d = nums[0], nums[1].zfill(2), nums[2].zfill(2)
            return y + m + d
        return "".join(nums)

    # scan_no → file_path 맵 + 파일명 → file_path 맵 + 날짜 → [scan_items] 맵
    scan_map: dict[int, str] = {}
    fname_map: dict[str, str] = {}
    date_map: dict[str, list] = {}
    for si in scan_items:
        no = si.get("no")
        fp = si.get("file_path", "")
        if no is not None and fp:
            try:
                scan_map[int(no)] = fp
            except (ValueError, TypeError):
                pass
        if fp:
            fname = fp.replace("\\", "/").split("/")[-1]
            fname_map[fname] = fp
        norm = _norm_date(si.get("date", ""))
        if norm:
            date_map.setdefault(norm, []).append(si)

    # 섹션 제목
    _add_paragraph(doc, "", size_pt=6, space_after_pt=4)
    _add_heading(doc, "[ 참조 원본 파일 목록 ]",
                 size_pt=10, bold=True,
                 space_before_pt=10, space_after_pt=4)
    _add_paragraph(doc,
                   "아래는 본 귀책분석 작성에 참조된 원본 파일의 경로입니다. "
                   "검수 시 해당 파일을 직접 열어 내용을 확인하십시오.",
                   size_pt=9, space_after_pt=4)

    # correspondence_texts.md 파싱: 섹션별 본문 텍스트 + 파일명 맵
    corr_sections: list[tuple[str, str]] = []  # [(파일명, 본문텍스트), ...]
    corr_path = output_dir / "correspondence_texts.md"
    if corr_path.exists():
        corr_raw = corr_path.read_text(encoding="utf-8")
        # 각 섹션: ## [N] 제목 ... **파일**: path ... ### 전문 ... 본문
        sec_pattern = _re.compile(
            r"## \[\d+\].+?\n.*?\*\*파일\*\*:\s*(.+?)\n.*?### 전문\n(.+?)(?=\n## \[|\Z)",
            _re.DOTALL
        )
        for m in sec_pattern.finditer(corr_raw):
            rel_path = m.group(1).strip()
            body     = m.group(2)
            fname    = rel_path.replace("\\", "/").split("/")[-1]
            # 파일명으로 scan_result에서 full_path 조회
            fp = fname_map.get(fname, "")
            corr_sections.append((fp or rel_path, body))

    def _find_in_corr(item: dict) -> str:
        """item의 날짜·공문번호·제목 키워드가 언급된 correspondence 섹션 경로 반환."""
        search_terms: list[str] = []
        raw_date = item.get("date", "").strip(".").replace(".", "").replace("-", "")
        if len(raw_date) >= 6:
            search_terms.append(raw_date[:8])
        doc_no = item.get("doc_number", "")
        if doc_no:
            # 공문번호에서 숫자+한글 토큰 추출
            search_terms += _re.findall(r'[\w가-힣]{3,}', doc_no)

        best_path, best_score = "", 0
        for fp, body in corr_sections:
            score = sum(1 for t in search_terms if t in body)
            if score > best_score:
                best_score, best_path = score, fp
        return best_path if best_score >= 1 else ""

    seen_paths: set[str] = set()
    for item in items:
        no_val   = item.get("no", "?")
        scan_no  = item.get("scan_no")
        src_file = item.get("source_file", "")
        subject  = item.get("subject", "")[:50]
        date     = item.get("date", "")
        is_direct = True

        # 전체 경로 조회 (우선순위: scan_no → source_file 파일명 역조회)
        full_path = ""
        if scan_no is not None:
            try:
                full_path = scan_map.get(int(scan_no), "")
            except (ValueError, TypeError):
                pass

        if not full_path and src_file:
            full_path = fname_map.get(src_file, "")

        # 날짜 기반 매칭
        if not full_path:
            item_date = _norm_date(item.get("date", ""))
            candidates = date_map.get(item_date, [])
            if len(candidates) == 1:
                full_path = candidates[0].get("file_path", "")
            elif len(candidates) > 1:
                item_subj_tokens = set(_re.findall(r'[\w가-힣]{3,}', item.get("subject", "")))
                best_c, best_score = None, 0
                for c in candidates:
                    c_tokens = set(_re.findall(r'[\w가-힣]{3,}', c.get("subject", "")))
                    score = len(item_subj_tokens & c_tokens)
                    if score > best_score:
                        best_score, best_c = score, c
                if best_c and best_score > 0:
                    full_path = best_c.get("file_path", "")

        # 직접 경로 없으면 언급 문서에서 탐색
        if not full_path:
            full_path = _find_in_corr(item)
            if full_path:
                is_direct = False

        # 중복 경로 skip
        if full_path and full_path in seen_paths:
            continue
        if full_path:
            seen_paths.add(full_path)

        _add_paragraph(doc,
                       f"No.{no_val}  [{date}]  {subject}",
                       size_pt=9, bold=True, space_before_pt=4, space_after_pt=1)
        if full_path:
            label = "" if is_direct else "※ 직접 경로 아님 (아래 문서에서 언급됨)\n"
            _add_paragraph(doc, label + full_path, size_pt=8, space_after_pt=3)
        else:
            _add_paragraph(doc, "(경로 확인 불가)", size_pt=8, space_after_pt=3)
