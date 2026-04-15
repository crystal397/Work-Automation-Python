"""
포맷별 텍스트 추출기 — docx / pdf / hwp / 이미지
각 포맷마다 여러 방법을 순서대로 시도하고, 모두 실패하면 파일명만 기록.
"""

from __future__ import annotations

import io
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

# config는 항상 사용 가능하므로 모듈 상단에서 1회만 import
sys.path.insert(0, str(Path(__file__).parent.parent))
import config  # noqa: E402

# ── DRM 워터마크 필터 ────────────────────────────────────────────────────────────

_DRM_LINE_RE = re.compile(
    r"(POSCO\s+ENC\s+CONFIDENTIAL|이\s*문서는\s*포스코이앤씨의\s*허락없이|"
    r"이\s*SAE\s*포스코이앤씨의|이\s*BAS\s*포스코이앤씨의|"
    r"이\s*ENE\s*BAROMN\s*허락없이|이\s*SAS\s*포스코이앤씨의|"
    r"이\s*SME\s*포스코이앤씨의|CONFIDENTIAL)",
    re.IGNORECASE,
)
# DRM 워터마크에 함께 찍히는 단독 타임스탬프 줄 (YYYY-MM-DD HH:MM:SS)
_DRM_TIMESTAMP_RE = re.compile(r"^\s*\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}\s*$")

def _strip_drm(text: str) -> str:
    """DRM 워터마크 줄(및 함께 찍힌 타임스탬프 줄)을 제거하고 반환."""
    lines = [
        ln for ln in text.splitlines()
        if not _DRM_LINE_RE.search(ln) and not _DRM_TIMESTAMP_RE.match(ln)
    ]
    return "\n".join(lines)


# ── 데이터 구조 ─────────────────────────────────────────────────────────────────

@dataclass
class ExtractResult:
    file_path: Path
    text: str = ""
    method: str = ""
    quality: str = "OK"   # OK / WARN / FAIL
    error: str = ""

    @property
    def short_name(self) -> str:
        return self.file_path.name

    def preview(self, chars: int = 500) -> str:
        return self.text[:chars]


# ── DOCX ────────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> ExtractResult:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        quality = "OK" if len(text) > 50 else "WARN"
        return ExtractResult(path, text, "python-docx", quality)
    except Exception as e:
        return ExtractResult(path, "", "python-docx", "FAIL", str(e))


# ── PDF ─────────────────────────────────────────────────────────────────────────

def _extract_pdf_plumber(path: Path) -> ExtractResult:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        text = _strip_drm("\n".join(parts))
        quality = "OK" if len(text.strip()) > 50 else "WARN"
        return ExtractResult(path, text, "pdfplumber", quality)
    except Exception as e:
        return ExtractResult(path, "", "pdfplumber", "FAIL", str(e))


def _extract_pdf_mupdf(path: Path) -> ExtractResult:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        parts = []
        for page in doc:
            t = page.get_text()
            if t:
                parts.append(t)
        doc.close()
        text = _strip_drm("\n".join(parts))
        quality = "OK" if len(text.strip()) > 50 else "WARN"
        return ExtractResult(path, text, "pymupdf", quality)
    except Exception as e:
        return ExtractResult(path, "", "pymupdf", "FAIL", str(e))


def _extract_pdf_ocr(path: Path) -> ExtractResult:
    """스캔 PDF → 이미지 변환 → OCR"""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH

        doc = fitz.open(str(path))
        parts = []
        for page_num, page in enumerate(doc):
            if page_num >= config.PDF_OCR_MAX_PAGES:
                break
            mat = fitz.Matrix(2.0, 2.0)  # 2x 해상도
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))
            t = pytesseract.image_to_string(img, lang=config.TESSERACT_LANG)
            if t.strip():
                parts.append(t)
        doc.close()
        text = "\n".join(parts)
        quality = "OK" if len(text) > 50 else "WARN"
        return ExtractResult(path, text, "pdf-ocr", quality)
    except Exception as e:
        return ExtractResult(path, "", "pdf-ocr", "FAIL", str(e))


def extract_pdf(path: Path) -> ExtractResult:
    # 1. pdfplumber
    r = _extract_pdf_plumber(path)
    if r.quality == "OK":
        return r
    # 2. PyMuPDF
    r2 = _extract_pdf_mupdf(path)
    if r2.quality == "OK":
        return r2
    # 3. OCR
    return _extract_pdf_ocr(path)


# ── HWP / HWPX ─────────────────────────────────────────────────────────────────

def _extract_hwp_hwp5txt(path: Path) -> ExtractResult:
    try:
        result = subprocess.run(
            ["hwp5txt", str(path)],
            capture_output=True, text=True, timeout=30, encoding="utf-8", errors="replace"
        )
        text = result.stdout.strip()
        quality = "OK" if len(text) > 50 else "WARN"
        return ExtractResult(path, text, "hwp5txt", quality,
                             result.stderr[:200] if result.returncode != 0 else "")
    except FileNotFoundError:
        return ExtractResult(path, "", "hwp5txt", "FAIL", "hwp5txt not installed")
    except Exception as e:
        return ExtractResult(path, "", "hwp5txt", "FAIL", str(e))


def _extract_hwp_libreoffice(path: Path) -> ExtractResult:
    """LibreOffice로 docx 변환 후 python-docx 읽기"""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, str(path)],
                capture_output=True, timeout=60
            )
            if result.returncode != 0:
                return ExtractResult(path, "", "libreoffice", "FAIL",
                                     result.stderr.decode(errors="replace")[:200])
            converted = Path(tmpdir) / (path.stem + ".docx")
            if not converted.exists():
                return ExtractResult(path, "", "libreoffice", "FAIL", "변환 파일 없음")
            return _extract_docx(converted)
    except FileNotFoundError:
        return ExtractResult(path, "", "libreoffice", "FAIL", "LibreOffice not installed")
    except Exception as e:
        return ExtractResult(path, "", "libreoffice", "FAIL", str(e))


def extract_hwp(path: Path) -> ExtractResult:
    r = _extract_hwp_hwp5txt(path)
    if r.quality == "OK":
        return r
    return _extract_hwp_libreoffice(path)


# ── 이미지 OCR ──────────────────────────────────────────────────────────────────

def extract_image(path: Path) -> ExtractResult:
    try:
        import pytesseract
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH
        img = Image.open(str(path))
        text = pytesseract.image_to_string(img, lang=config.TESSERACT_LANG)
        quality = "OK" if len(text.strip()) > 30 else "WARN"
        return ExtractResult(path, text, "image-ocr", quality)
    except Exception as e:
        return ExtractResult(path, "", "image-ocr", "FAIL", str(e))


# ── 통합 진입점 ─────────────────────────────────────────────────────────────────

SUPPORTED_EXT = {
    ".docx": _extract_docx,
    ".pdf":  extract_pdf,
    ".hwp":  extract_hwp,
    ".hwpx": extract_hwp,
    ".jpg":  extract_image,
    ".jpeg": extract_image,
    ".png":  extract_image,
    ".tif":  extract_image,
    ".tiff": extract_image,
}


def extract_file(path: Path) -> ExtractResult:
    """파일 하나를 텍스트로 추출. 지원하지 않는 형식은 FAIL 반환."""
    ext = path.suffix.lower()
    handler = SUPPORTED_EXT.get(ext)
    if handler is None:
        return ExtractResult(path, "", "unsupported", "FAIL", f"지원 안 되는 형식: {ext}")
    return handler(path)


def extract_first_pages(path: Path, max_chars: int = 2000) -> ExtractResult:
    """
    빠른 공문 판별용 — 전체 텍스트 앞부분만 추출.
    PDF는 1~2페이지만, docx는 앞 단락만 읽음.
    """
    ext = path.suffix.lower()

    # DRM 워터마크만 있는 이미지 PDF를 본문으로 오인하지 않도록:
    # 추출 텍스트가 100자 미만(워터마크 제거 후)이면 OCR로 폴백.
    _MIN_TEXT_LEN = 100

    if ext == ".pdf":
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages[:2]:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
                        if sum(len(p) for p in parts) >= max_chars:
                            break
            text = "\n".join(parts)[:max_chars]
            if len(_strip_drm(text).strip()) >= _MIN_TEXT_LEN:
                return ExtractResult(path, text, "pdfplumber-fast", "OK")
        except Exception:
            pass
        # 폴백: PyMuPDF + OCR — fitz 한 번만 open
        fitz_doc = None
        try:
            import fitz
            fitz_doc = fitz.open(str(path))
        except Exception:
            pass

        if fitz_doc is not None:
            try:
                # PyMuPDF 텍스트 추출
                pages_text = [t for page in list(fitz_doc)[:2] if (t := page.get_text())]
                text = "\n".join(pages_text)[:max_chars]
                if len(_strip_drm(text).strip()) >= _MIN_TEXT_LEN:
                    return ExtractResult(path, text, "pymupdf-fast", "OK")

                # OCR 폴백 (같은 doc 재사용, 추가 open 없음)
                try:
                    import pytesseract
                    from PIL import Image
                    import io as _io
                    pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH
                    parts = []
                    for page in list(fitz_doc)[:2]:
                        mat = fitz.Matrix(2.0, 2.0)
                        pix = page.get_pixmap(matrix=mat)
                        img = Image.open(_io.BytesIO(pix.tobytes("png")))
                        t = pytesseract.image_to_string(img, lang=config.TESSERACT_LANG)
                        if t.strip():
                            parts.append(t)
                    text = "\n".join(parts)[:max_chars]
                    if text.strip():
                        return ExtractResult(path, text, "pdf-ocr-fast", "OK")
                except Exception:
                    pass
            finally:
                fitz_doc.close()

    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            char_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
                    char_count += len(para.text)
                    if char_count >= max_chars:
                        break
            text = "\n".join(parts)[:max_chars]
            return ExtractResult(path, text, "python-docx-fast", "OK" if text else "WARN")
        except Exception:
            pass

    elif ext in (".hwp", ".hwpx"):
        r = _extract_hwp_hwp5txt(path)
        if r.quality == "OK":
            r.text = r.text[:max_chars]
            return r

    # 다른 형식은 전체 추출 후 잘라서 반환
    r = extract_file(path)
    r.text = r.text[:max_chars]
    return r
