"""
귀책분석 자동화 시스템 — 진입점

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  워크플로우 (신규 프로젝트 추가 시)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  STEP 1 — 스캔 + 프롬프트 생성
    python main.py scanprepare "C:\\...\\수신자료" [--project <이름>]
      → scan_summary.md   : 스캔 결과 요약 (검토·편집 후 필요 시 재prepare)
      → scan_borderline.md: 경계선 공문 목록 (포함 여부 직접 판단)
      → scan_result.json  : 공문 목록 (직접 편집 가능)
      → prompt_for_claude.md: Claude Code에 전달할 귀책분석 작성 지시서

  STEP 2 — Claude Code가 귀책분석 JSON 작성
    Claude Code가 prompt_for_claude.md 를 읽고
    output/<프로젝트명>/귀책분석_data.json 을 직접 저장.

  STEP 3 — 사전 검증 + docx 생성
    python main.py finish [프로젝트명]
      → validate: 스키마 오류·소결 누락·delay_days 불일치 점검
        오류 발견 시 Claude에게 전달할 수정 요청 블록 자동 출력
      → generate: 02_귀책분석_[프로젝트명]_[날짜].docx (검증 통과 시에만)

  STEP 4 — 품질 검증
    python main.py compare [프로젝트명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  전체 명령 목록
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  python main.py learn                                          # reference 보고서 패턴 학습
  python main.py scanprepare <공문폴더경로> [--project <이름>]  # 스캔 + 프롬프트 생성 (콤보)
  python main.py validate    [프로젝트명]                       # JSON 사전 검증
  python main.py generate    [프로젝트명]                       # docx 생성
  python main.py finish      [프로젝트명]                       # validate + generate (콤보)
  python main.py compare     [프로젝트명]                       # 단일 품질 비교
  python main.py compare-all                                    # 전체 품질 비교
  python main.py rescan-all                                     # 전체 프로젝트 일괄 재스캔

개별 명령 (scanprepare·finish 분해 시):
  python main.py scan    <공문폴더경로> [--project <이름>]  # 스캔만
  python main.py prepare [프로젝트명]                       # 프롬프트 생성만

  --project 옵션: scanprepare/scan 단계에서 폴더명 대신 원하는 프로젝트명 지정
    예) python main.py scanprepare "C:\\...\\수신자료" --project 인덕원5공구
"""

import io
import sys
from pathlib import Path

# Windows 콘솔 UTF-8
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

sys.path.insert(0, str(Path(__file__).parent))

import json

import config
from src.reference_learner import learn_all
from src.correspondence_scanner import scan
from src.prompt_builder import build as build_prompt
from src.report_generator import generate, _validate_data


def _print_header():
    print("=" * 60)
    print("  귀책분석 자동화 시스템  v1.0")
    print("=" * 60)


def _resolve_project_dir(project_name: str | None) -> tuple[str, Path]:
    """
    프로젝트명 → (project_name, project_dir) 반환.
    project_name 이 None 이면 .current_project 에서 읽는다.
    둘 다 없으면 오류 출력 후 종료.
    """
    if project_name:
        pdir = config.get_project_dir(project_name)
    else:
        project_name = config.load_current_project()
        if not project_name:
            print(
                "[오류] 프로젝트명을 알 수 없습니다.\n"
                "  먼저 scan 을 실행하거나 명령에 프로젝트명을 직접 지정하세요.\n"
                "  예) python main.py prepare 인덕원5공구"
            )
            sys.exit(1)
        pdir = config.get_project_dir(project_name)

    if not pdir.exists():
        print(
            f"[오류] 프로젝트 폴더가 없습니다: {pdir}\n"
            f"  먼저 python main.py scan 을 실행하세요."
        )
        sys.exit(1)

    return project_name, pdir


# ── Step 1: learn ────────────────────────────────────────────────────────────

def _token_overlap_score(a: str, b: str) -> float:
    """
    두 문자열의 단어 겹침 비율 (0.0 ~ 1.0).
    공백·기호 분리 + 한국어 숫자 경계 분리까지 적용하여 공통 토큰 수 / min(len_a, len_b)
    """
    import re as _re
    def _tokens(s: str) -> set[str]:
        # 1) 특수문자를 공백으로 치환
        norm = _re.sub(r"[\s\-~_\[\]()\.·×,·!?·]", " ", s)
        # 2) 공백 분리 + 2자 이상 토큰
        base = {p for p in norm.split() if len(p) >= 2}
        # 3) 한국어+숫자 경계에서 추가 분리 (예: "부산도시철도양산선1공구현장" → "양산선1", "1공구")
        extra: set[str] = set()
        for seg in _re.findall(r"[가-힣]+[0-9]+[가-힣]*|[0-9]+[가-힣]+", norm.replace(" ", "")):
            if len(seg) >= 2:
                extra.add(seg)
        # 4) 숫자 포함 한국어 키워드 단독 (예: "11공구", "5공구", "1공구")
        for seg in _re.findall(r"[0-9]+공구", norm):
            extra.add(seg)
        return base | extra

    ta, tb = _tokens(a), _tokens(b)
    if not ta or not tb:
        return 0.0
    common = ta & tb
    # 복합어 내 부분문자열 일치: 3자 이상 토큰이 상대 토큰에 포함되면 추가 인정
    # (예: "양산선1공구" ∈ "부산도시철도양산선1공구현장", "군부대" ∈ "원주군부대")
    for ta_tok in ta:
        if len(ta_tok) < 3:
            continue
        for tb_tok in tb:
            if len(tb_tok) < 3 or ta_tok == tb_tok:
                continue
            if ta_tok in tb_tok or tb_tok in ta_tok:
                shorter = ta_tok if len(ta_tok) <= len(tb_tok) else tb_tok
                common.add(shorter)
    return len(common) / min(len(ta), len(tb))


def _print_reference_output_mapping():
    """
    reference 파일명 ↔ output 폴더명 퍼지 매칭 결과를 출력.
    학습 완료 후 진단용으로 호출한다.
    """
    ref_dir = config.REFERENCE_DIR
    out_dir = config.OUTPUT_DIR

    ref_files = (
        sorted(ref_dir.glob("*.docx")) + sorted(ref_dir.glob("*.DOCX"))
        + sorted(ref_dir.glob("*.pdf")) + sorted(ref_dir.glob("*.PDF"))
    )
    # 템플릿 파일 제외 (보고서_템플릿* / 비용산정기준*)
    ref_files = [
        f for f in ref_files
        if not any(kw in f.name for kw in ["템플릿", "비용산정기준", "template", "Template"])
    ]

    if not out_dir.exists():
        return

    out_folders = [d for d in sorted(out_dir.iterdir()) if d.is_dir()]
    if not out_folders:
        return

    print("\n[참고] reference ↔ output 폴더 매칭 진단")
    print("─" * 60)

    THRESHOLD = 0.35  # 이 점수 이상이면 매칭으로 간주

    matched_refs: set[str] = set()
    for folder in out_folders:
        best_score = 0.0
        best_ref: Path | None = None
        for rf in ref_files:
            score = _token_overlap_score(folder.name, rf.stem)
            if score > best_score:
                best_score = score
                best_ref = rf
        if best_ref and best_score >= THRESHOLD:
            tag = "✅"
            matched_refs.add(best_ref.name)
        else:
            tag = "❓"
            best_ref = None

        ref_label = best_ref.name if best_ref else "(매칭 없음)"
        print(f"  {tag} output/{folder.name}/")
        print(f"       ← reference/{ref_label}")

    # 매칭 안 된 reference 파일
    unmatched = [rf for rf in ref_files if rf.name not in matched_refs]
    if unmatched:
        print()
        print("  ⚠️  output 폴더 없이 reference만 존재하는 파일:")
        for rf in unmatched:
            print(f"     · {rf.name}")

    print("─" * 60)


def cmd_learn():
    _print_header()
    print("\n[1단계] REFERENCE 보고서 학습")
    print("─" * 60)

    if not config.REFERENCE_DIR.exists():
        print(f"[오류] reference 폴더가 없습니다: {config.REFERENCE_DIR}")
        sys.exit(1)

    # learn 결과는 프로젝트 무관 공통 → OUTPUT_DIR 루트에 저장
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = learn_all(config.REFERENCE_DIR, config.OUTPUT_DIR)

    print("\n" + "─" * 60)
    print("REFERENCE 학습 완료.")
    print(f"\n→ {out_path}")

    _print_reference_output_mapping()

    print("\n다음 단계: python main.py scan <공문폴더경로>")


# ── Step 2: scan ─────────────────────────────────────────────────────────────

def _parse_scan_args(args: list[str]) -> tuple[list[str], str | None]:
    """
    scan 인수를 파싱한다.
    --project <이름> 옵션을 분리하고 나머지는 경로 목록으로 반환.
    반환: (경로 목록, 프로젝트명 또는 None)
    """
    paths: list[str] = []
    project: str | None = None
    i = 0
    while i < len(args):
        if args[i] == "--project" and i + 1 < len(args):
            project = args[i + 1]
            i += 2
        else:
            paths.append(args[i])
            i += 1
    return paths, project


def _extract_project_name(path: Path) -> str:
    """
    경로의 상위 폴더를 탐색하여 프로젝트명을 추출한다.

    지원 패턴 (순서대로 시도):
      1. YYMMDD_업체명_프로젝트명 공기연장 간접비  (언더스코어 두 개)
         예) 251120_경남기업_평택기지-오산 천연가스 공급설비 건설공사 2공구 공기연장 간접비
             → 평택기지-오산 천연가스 공급설비 건설공사 2공구
      2. YYMMDD 프로젝트명 공기연장 간접비  (공백 구분, 업체명 없음)
         예) 240226 울산-포항 전철2공구 공기연장 간접비
             → 울산-포항 전철2공구
      3. YYMMDD_업체명 프로젝트명 공기연장 간접비  (언더스코어 한 개, 업체+프로젝트 공백 구분)
         예) 230920_디엘이앤씨 울릉공항건설공사 공기연장 간접비
             → 디엘이앤씨 울릉공항건설공사

    패턴 모두 실패하면 상위 폴더 중 '간접비'/'공기연장'이 포함된 가장 가까운 폴더명을 반환.
    그것도 없으면 입력 경로의 상위 폴더명을 반환한다.
    """
    import re
    patterns = [
        re.compile(r"^\d{6}_[^_]+_(.+?)\s*(?:공기연장\s*)?간접비.*$"),   # 언더스코어 두 개
        re.compile(r"^\d{6}\s+(.+?)\s*(?:공기연장\s*)?간접비.*$"),        # 공백 구분
        re.compile(r"^\d{6}_(.+?)\s*(?:공기연장\s*)?간접비.*$"),          # 언더스코어 한 개
    ]
    candidates = [path, *path.parents]
    for parent in candidates:
        for pat in patterns:
            m = pat.match(parent.name)
            if m:
                return m.group(1).strip()
    # 패턴 미매칭 — '간접비' 또는 '공기연장'이 포함된 가장 가까운 상위 폴더명 사용
    for parent in candidates:
        if any(kw in parent.name for kw in ("간접비", "공기연장", "클레임")):
            return parent.name
    # 최후 폴백: 입력 경로 자체의 폴더명
    return path.name


def cmd_scan(raw_args: list[str]):
    _print_header()
    print("\n[2단계] 업체 공문 스캔 및 필터링")
    print("─" * 60)

    import os

    vendor_path_strs, project_name_arg = _parse_scan_args(raw_args)

    # 경로 결정: 인수 → 환경변수 → 오류
    if vendor_path_strs:
        vendor_dirs = [Path(p) for p in vendor_path_strs]
    else:
        env_path = os.environ.get("CORRESPONDENCE_DIR", "")
        if env_path:
            vendor_dirs = [Path(p.strip()) for p in env_path.split(";") if p.strip()]
        else:
            print("[오류] 공문 폴더 경로를 지정하세요.")
            print("  python main.py scan <경로1> [경로2] ... [--project <이름>]")
            print("  또는 .env 파일에 CORRESPONDENCE_DIR=<경로> 설정")
            sys.exit(1)

    for d in vendor_dirs:
        if not d.exists():
            print(f"[오류] 폴더가 존재하지 않습니다: {d}")
            sys.exit(1)

    # 프로젝트명 결정: --project > 상위 폴더 패턴 추출 > 스캔 폴더명
    if project_name_arg:
        project_name = project_name_arg
    else:
        project_name = _extract_project_name(vendor_dirs[0])

    project_dir = config.get_project_dir(project_name)
    project_dir.mkdir(parents=True, exist_ok=True)

    # 현재 프로젝트 저장
    config.save_current_project(project_name)

    print(f"\n프로젝트 폴더: output/{project_dir.name}/")
    result_path = scan(vendor_dirs, project_dir)

    print("\n" + "─" * 60)
    print("스캔 완료.")
    print(f"\n→ output/{project_dir.name}/scan_summary.md  를 열어 공문 목록을 확인하세요.")
    print(f"→ output/{project_dir.name}/scan_result.json 을 편집하여 항목을 추가/제외한 후")
    print(f"  python main.py prepare 를 실행하세요.")


# ── Step 3: prepare ──────────────────────────────────────────────────────────

def cmd_prepare(project_name: str | None):
    _print_header()
    print("\n[3단계] claude.ai 프롬프트 조립")
    print("─" * 60)

    project_name, project_dir = _resolve_project_dir(project_name)
    print(f"\n프로젝트 폴더: output/{project_dir.name}/")

    try:
        prompt_path = build_prompt(project_dir, project_name)
    except FileNotFoundError as e:
        print(f"[오류] {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"[오류] {e}")
        sys.exit(1)

    print("\n" + "─" * 60)
    print("프롬프트 생성 완료.")
    print(f"\n→ {prompt_path}")
    print("\n다음 단계 (둘 중 하나 선택):")
    print()
    print("  ★ Claude Code 사용 (터미널에서 바로):")
    print(f'    "output/{project_dir.name}/prompt_for_claude.md 읽고')
    print(f'     귀책분석_data.json 생성해줘"')
    print(f"    → Claude Code가 파일을 읽고 분석하여 JSON을 직접 저장")
    print()
    print("  ★ claude.ai 사용 (웹 브라우저):")
    print("    1. prompt_for_claude.md 전체 내용을 복사")
    print("    2. claude.ai 새 대화창에 붙여넣기")
    print("    3. claude.ai가 작성한 JSON 전체를 복사")
    print(f"    4. output/{project_dir.name}/귀책분석_data.json 파일로 저장")
    print()
    print("  → 완료 후: python main.py generate")


# ── Step 4: generate ─────────────────────────────────────────────────────────

def cmd_generate(project_name: str | None):
    _print_header()
    print("\n[4단계] docx 생성")
    print("─" * 60)

    project_name, project_dir = _resolve_project_dir(project_name)
    # 폴더명에서 날짜 접두사·불필요 접미사를 제거한 순수 프로젝트명 추출
    extracted = _extract_project_name(project_dir)
    display_name = extracted if extracted != project_dir.name else project_name
    print(f"\n프로젝트 폴더: output/{project_dir.name}/")

    try:
        out_path = generate(project_dir, display_name)
    except FileNotFoundError as e:
        print(f"[오류] {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"[오류] {e}")
        sys.exit(1)

    print("\n" + "─" * 60)
    print("docx 생성 완료.")
    print(f"\n→ {out_path}")


# ── Step 3.5: validate ───────────────────────────────────────────────────────

def cmd_validate(project_name: str | None, _stop_on_error: bool = True) -> bool:
    """검증 실행. 반환: True=hard_error 없음, False=hard_error 있음."""
    _print_header()
    print("\n[3.5단계] 귀책분석_data.json 사전 검증")
    print("─" * 60)

    project_name, project_dir = _resolve_project_dir(project_name)
    print(f"\n프로젝트 폴더: output/{project_dir.name}/")

    # ── JSON 파일 탐색 ────────────────────────────────────────────────
    import glob as _glob
    import io as _io

    candidates = _glob.glob(str(project_dir / "*data.json"))
    if not candidates:
        print(f"\n[오류] 귀책분석_data.json 파일이 없습니다: {project_dir}")
        print("  먼저 Claude에게 prompt_for_claude.md를 전달하여 JSON을 생성하세요.")
        sys.exit(1)
    data_path = Path(candidates[0])

    with open(data_path, encoding="utf-8") as f:
        try:
            data = json.load(f)
        except json.JSONDecodeError as e:
            print(f"\n[오류] JSON 파싱 실패: {e}")
            print("─" * 60)
            print("  귀책분석_data.json 파일이 JSON 형식 오류입니다.")
            print(f"  오류 위치: {e}")
            print("  JSON 전체를 다시 검토하고 올바른 형식으로 재출력해 주세요.")
            print("─" * 60)
            sys.exit(1)

    # ── 분류: hard_errors(generate 중단) / soft_warnings(경고만) ──────
    hard_errors:  list[str] = []   # Claude 수정 필요 → generate 불가
    soft_warnings: list[str] = []  # 확인 권고 → generate는 가능

    # ── scan_result.json 로드 (교차검증용) ───────────────────────────
    scan_items_by_no: dict[int, dict] = {}
    scan_doc_numbers: set[str] = set()
    scan_result_path = project_dir / "scan_result.json"
    if scan_result_path.exists():
        try:
            with open(scan_result_path, encoding="utf-8") as _sf:
                _sdata = json.load(_sf)
            for idx_s, _si in enumerate(_sdata.get("items", []), 1):
                _no = _si.get("no", idx_s)
                scan_items_by_no[int(_no)] = _si
                _dn = (_si.get("doc_number") or "").strip()
                if _dn:
                    scan_doc_numbers.add(_dn)
        except Exception:
            pass

    # ── [경고] items scan_no 교차검증 ────────────────────────────────
    data_items = data.get("items", [])
    if isinstance(data_items, list) and scan_items_by_no:
        for i, di in enumerate(data_items):
            sno = di.get("scan_no") or di.get("no")
            if sno is not None:
                try:
                    sno = int(sno)
                except (ValueError, TypeError):
                    continue
                if sno not in scan_items_by_no:
                    soft_warnings.append(
                        f"  [경고] items[{i}] scan_no={sno} 가 scan_result.json에 없습니다.\n"
                        f"     수동 추가 항목이면 무시하세요. 아니라면 번호를 확인하세요."
                    )

    # ── [오류] accountability_diagram 합계 행 삽입 감지 ─────────────
    diagram = data.get("accountability_diagram", [])
    _SUM_ROW_KW = {"합계", "소계", "계", "total", "sum"}
    if isinstance(diagram, list):
        for i, row in enumerate(diagram):
            if not isinstance(row, dict):
                continue
            cause_raw = (row.get("cause") or row.get("delay_cause") or "").strip()
            if cause_raw.lower() in _SUM_ROW_KW or cause_raw in _SUM_ROW_KW:
                total = data.get("total_delay_days", "?")
                hard_errors.append(
                    f"  [오류] accountability_diagram[{i}] cause='{cause_raw}' — 합계 행이 삽입되어 있습니다.\n"
                    f"     이 행을 삭제하세요. report_generator가 delay_days를 자동 합산하므로\n"
                    f"     합계 행이 있으면 total_delay_days({total})의 2배로 오류 발생합니다.\n"
                    f"     accountability_diagram에는 실제 지연 사유 항목만 포함하세요."
                )

    # ── [오류] accountability_diagram delay_days 누락 ────────────────
    if isinstance(diagram, list):
        for i, row in enumerate(diagram):
            if not isinstance(row, dict):
                continue
            dd = row.get("delay_days")
            if dd is None:
                cause = (row.get("cause") or row.get("delay_cause") or f"항목 {i+1}")[:30]
                hard_errors.append(
                    f"  [오류] accountability_diagram[{i}] ('{cause}')에\n"
                    f"     delay_days 필드가 없습니다. 일수 미확정 시 0으로 기재하세요.\n"
                    f"     ※ 모든 항목 delay_days 합계 = total_delay_days({data.get('total_delay_days','?')})\n"
                    f"     ※ '합계' 행 별도 추가 금지 (이중 계산됨)"
                )

    # ── [경고] accountability_diagram source_docs 누락 ───────────────
    if isinstance(diagram, list):
        for i, row in enumerate(diagram):
            if not isinstance(row, dict):
                continue
            cause_label = (row.get("cause") or row.get("delay_cause") or f"항목 {i+1}")[:30]
            sd = row.get("source_docs")
            if sd is None or (isinstance(sd, list) and len(sd) == 0):
                soft_warnings.append(
                    f"  [경고] accountability_diagram[{i}] ('{cause_label}')\n"
                    f"     source_docs 필드가 없거나 비어 있습니다.\n"
                    f"     근거 공문번호·변경계약 차수를 리스트로 추가하세요.\n"
                    f"     예: \"source_docs\": [\"제22-0123호\", \"3차 변경계약\"]"
                )
            elif isinstance(sd, list) and scan_doc_numbers:
                # source_docs 내 번호가 scan에 있는지 교차검증
                for ref in sd:
                    ref_str = str(ref).strip()
                    if ref_str and ref_str not in scan_doc_numbers:
                        # 부분 매칭 허용 (공문번호가 앞뒤 텍스트와 붙어있을 수 있음)
                        matched = any(ref_str in dn or dn in ref_str for dn in scan_doc_numbers)
                        if not matched:
                            soft_warnings.append(
                                f"  [경고] accountability_diagram[{i}] source_docs의\n"
                                f"     '{ref_str}' 가 scan_result.json 공문번호 목록에 없습니다.\n"
                                f"     공문번호 표기가 다를 수 있으니 직접 확인하세요."
                            )

    # ── [오류/경고] detail_narratives 소결 누락 ─────────────────────
    # 결론 문구는 계약 유형 무관하게 동일 → hard_error
    # 조항 인용은 민간계약 등 비표준 조항 가능 → soft_warning
    _SOGYEOL_CONCLUSION = "계약상대자의 책임 없는 사유"
    _SOGYEOL_CLAUSE_KW  = [
        "제8절", "제22조", "제25조", "제26조", "제27조", "제74조",
        "일반조건", "계약조건", "해당 조항", "계약 조항",
    ]
    narratives = data.get("detail_narratives", [])
    if isinstance(narratives, list):
        for i, block in enumerate(narratives):
            if not isinstance(block, dict):
                continue
            paras = block.get("paragraphs", [])
            if not isinstance(paras, list) or len(paras) == 0:
                continue
            block_title = block.get("title") or block.get("label", f"[{i}]")
            paras_text = " ".join(str(p) for p in paras)
            if _SOGYEOL_CONCLUSION not in paras_text:
                hard_errors.append(
                    f"  [오류] detail_narratives '{block_title}' 블록에 소결 결론이 없습니다.\n"
                    f"     마지막 단락에 \"계약상대자의 책임 없는 사유에 해당합니다\"를 포함시키세요."
                )
            elif not any(kw in paras_text for kw in _SOGYEOL_CLAUSE_KW):
                soft_warnings.append(
                    f"  [경고] detail_narratives '{block_title}' 블록에 근거 조항 인용이 없습니다.\n"
                    f"     소결에 공사계약일반조건 제XX조 등 근거 조항을 인용하세요."
                )

    # ── [경고] responsible_party 유효값 확인 ─────────────────────────
    # compare CHECK 6과 동일한 기준 사용
    # 국가/지방계약: 발주처·발주자·시공사·수급인  /  민간계약: 도급인·수급인  /  기타: 공동귀책
    _VALID_PARTIES = {"발주처", "발주자", "도급인", "시공사", "수급인", "공동귀책"}
    import re as _re_rp
    def _strip_rp(v: str) -> str:
        v = v.replace("\n", "").strip()
        m = _re_rp.match(r"^(발주처|발주자|도급인|시공사|수급인|공동귀책)", v)
        return m.group(1) if m else v
    if isinstance(diagram, list):
        for i, row in enumerate(diagram):
            if not isinstance(row, dict):
                continue
            rp = (row.get("responsible_party") or "").strip()
            if rp and _strip_rp(rp) not in _VALID_PARTIES:
                cause_label = (row.get("cause") or row.get("delay_cause") or f"항목 {i+1}")[:30]
                soft_warnings.append(
                    f"  [경고] accountability_diagram[{i}] ('{cause_label}')\n"
                    f"     responsible_party='{rp}' 가 표준값이 아닙니다.\n"
                    f"     표준값: 발주처 / 발주자 / 도급인 / 시공사 / 수급인 / 공동귀책\n"
                    f"     괄호 부연 허용: 예) '발주처(한국가스공사)', '도급인 (수급인의 책임 없는 사유)'"
                )

    # ── [경고] items 날짜 미확정 / doc_number OCR 아티팩트 / scan_no null ─
    import re as _re
    _DATE_UNCERTAIN = ("xx", "?", "미확정", "불명", "확인필요")
    if isinstance(data_items, list):
        for i, di in enumerate(data_items):
            if not isinstance(di, dict):
                continue
            item_label = f"items[{i}](no={di.get('no','?')})"

            # 날짜 미확정
            dt = (di.get("date") or "").strip()
            if any(tok in dt for tok in _DATE_UNCERTAIN):
                soft_warnings.append(
                    f"  [경고] {item_label} date='{dt}' — 미확정 날짜입니다.\n"
                    f"     원문 공문에서 실제 발신일을 확인 후 수정하세요."
                )

            # doc_number OCR 아티팩트 의심
            # 한글이 포함된 번호에 소문자 영문 2자 이상이 연속 등장하면 OCR 오인식 가능성
            dn = (di.get("doc_number") or "").strip()
            if (dn
                    and _re.search(r"[가-힣]", dn)
                    and _re.search(r"[a-z]{2,}", dn)):
                soft_warnings.append(
                    f"  [경고] {item_label} doc_number='{dn}' — OCR 오인식 의심.\n"
                    f"     한글 번호에 소문자 영문이 섞여 있습니다 (예: -ys, -bts).\n"
                    f"     원문 공문에서 실제 문서번호를 확인 후 수정하세요."
                )

            # scan_no null → 수동 추가 항목
            sno = di.get("scan_no")
            if sno is None:
                soft_warnings.append(
                    f"  [경고] {item_label} scan_no=null — 스캔 미수록 수동 추가 항목.\n"
                    f"     원본 파일이 scan_result.json에 없으므로 내용을 직접 검토하세요."
                )

    # ── 스키마 검증 (_validate_data) — 출력 캡처 ─────────────────────
    captured = _io.StringIO()
    old_stdout = sys.stdout
    sys.stdout = captured
    schema_error = None
    try:
        _validate_data(data, data_path)
    except ValueError as e:
        schema_error = str(e)
    finally:
        sys.stdout = old_stdout
    schema_output = captured.getvalue()

    # ── 결과 출력 ─────────────────────────────────────────────────────
    has_schema_error = schema_error is not None

    if schema_output.strip():
        print(schema_output, end="")

    # soft_warnings 출력 (generate는 가능)
    if soft_warnings:
        print()
        print("─" * 60)
        print("  ⚠️  경고 (generate는 가능하지만 확인 권고)")
        print("─" * 60)
        for w in soft_warnings:
            print(w)

    # 오류 없으면 통과
    if not hard_errors and not has_schema_error:
        print("\n" + "─" * 60)
        if soft_warnings:
            print("⚠️  경고 있음 — 위 항목 확인 후 generate 를 실행하세요.")
        else:
            print("✅ 검증 통과 — 오류 없음.")
        print("─" * 60)
        return True

    # hard_errors 있으면 Claude 수정 요청 출력 후 중단
    print()
    print("=" * 60)
    print("  ▼ Claude에게 전달할 수정 요청 (아래를 복사하여 붙여넣기)")
    print("=" * 60)
    print()
    print("귀책분석_data.json에 다음 오류가 있습니다. 수정 후 전체 JSON을 다시 출력해 주세요.\n")
    for err in hard_errors:
        print(err)
    if has_schema_error:
        print(f"\n[스키마 오류] {schema_error}")
    print()
    print("=" * 60)
    if _stop_on_error:
        sys.exit(1)
    return False


# ── compare: output vs reference 품질 검증 ──────────────────────────────────

def _find_reference_file(project_dir: Path) -> "Path | None":
    """output 폴더명과 퍼지 매칭으로 reference 파일을 찾는다."""
    ref_dir = config.REFERENCE_DIR
    if not ref_dir.exists():
        return None
    ref_files = (
        sorted(ref_dir.glob("*.docx")) + sorted(ref_dir.glob("*.DOCX"))
        + sorted(ref_dir.glob("*.pdf")) + sorted(ref_dir.glob("*.PDF"))
    )
    ref_files = [
        f for f in ref_files
        if not any(kw in f.name for kw in ["템플릿", "비용산정기준", "template", "Template"])
    ]
    best_score, best_ref = 0.0, None
    for rf in ref_files:
        s = _token_overlap_score(project_dir.name, rf.stem)
        if s > best_score:
            best_score, best_ref = s, rf
    return best_ref if best_score >= 0.35 else None


def _extract_text_from_docx(path: Path) -> str:
    """docx에서 전체 텍스트를 추출한다. 실패 시 빈 문자열."""
    try:
        from docx import Document as _Document
        doc = _Document(str(path))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except Exception:
        return ""


def _extract_text_from_pdf(path: Path) -> str:
    """pdf에서 텍스트를 추출한다. 실패 시 빈 문자열."""
    try:
        import pdfminer.high_level as _pml
        import io as _io2
        buf = _io2.StringIO()
        with open(path, "rb") as fh:
            _pml.extract_text_to_fp(fh, buf, output_type="text")
        return buf.getvalue()
    except Exception:
        return ""


def _extract_ref_text(ref_path: Path) -> str:
    """reference 파일에서 텍스트를 추출한다."""
    if ref_path.suffix.lower() == ".pdf":
        return _extract_text_from_pdf(ref_path)
    return _extract_text_from_docx(ref_path)


def _extract_ref_tables(ref_path: Path) -> list[list[list[str]]]:
    """reference docx에서 표를 추출한다. [{행: [셀, ...]}]"""
    if ref_path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document as _Document
        doc = _Document(str(ref_path))
        result = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            result.append(rows)
        return result
    except Exception:
        return []


def _find_accountability_table_in_ref(tables: list[list[list[str]]]) -> "list[list[str]] | None":
    """
    reference 표 목록에서 귀책사유 도식 요약 표를 찾는다.
    reference에는 공문현황 상세표(items)와 귀책사유 도식표 두 개가 있을 수 있다.
    귀책사유 도식표는 상세표보다 뒤에 등장하고 행 수가 더 적으므로 마지막 매칭 표를 반환한다.

    제외 기준 (법령 조문 인용 표):
    - 데이터 행 첫 번째 셀이 "1)", "가.", "나." 등 법령 열거 번호 패턴으로 시작하면 제외.
    - 이런 표는 지방계약법/공사계약일반조건 조문을 나열하는 표이며 귀책 도식표가 아님.
    """
    import re as _re
    kws = {"공기지연 사유", "공기지연사유", "비용부담자", "비용부담", "관련 근거", "관련근거",
           "귀책", "지연일수", "귀책사유"}
    # 법령 조문 열거 패턴: "1)", "2)", "가.", "나.", "①", "②" 등으로 시작하는 셀
    _LAW_PAT = _re.compile(r"^(?:[0-9]+\)|[가-힣]\.|[①-⑳]|제\s*[0-9]+)")

    matched: list[list[list[str]]] = []
    for table in tables:
        if not table or len(table[0]) < 2:  # 1컬럼 표는 법령 인용 박스 → 제외
            continue
        # 헤더 행에 귀책 관련 키워드가 있는지 확인
        header_text = " ".join(cell.replace(" ", "") for cell in table[0])
        if not any(kw.replace(" ", "") in header_text for kw in kws):
            continue
        # 데이터 행(두 번째 행 이후)이 법령 조문 열거면 제외
        data_rows = table[1:] if len(table) > 1 else []
        if data_rows:
            first_cell = data_rows[0][0].strip() if data_rows[0] else ""
            if _LAW_PAT.match(first_cell):
                continue  # 법령 조문 인용 표 → 건너뜀
        matched.append(table)

    if not matched:
        return None
    # 여러 표가 매칭된 경우: 마지막 표가 귀책사유 도식 요약 표
    return matched[-1]


def _find_delay_days_in_text(text: str) -> "list[int]":
    """텍스트에서 'N일' 패턴 숫자를 추출한다 (3자리 이상만)."""
    import re as _re
    # "159일", "387일" 등 100 이상 숫자
    hits = [int(m) for m in _re.findall(r"\b([1-9]\d{2,})\s*일\b", text)]
    return hits


def _check_item(label: str, passed: bool, detail: str = "") -> tuple[str, bool, str]:
    return (label, passed, detail)


def _compare_one(project_dir: Path, ref_path: Path) -> dict:
    """
    단일 프로젝트의 data.json과 reference를 비교하여 체크 결과를 반환한다.
    반환: {checks: [(label, ok, detail)], score: (ok, total), ref_name: str}
    """
    import re as _re

    checks: list[tuple[str, bool, str]] = []

    # ── data.json 로드 ──
    import glob as _glob
    candidates = _glob.glob(str(project_dir / "*data.json"))
    if not candidates:
        return {"error": "data.json 없음", "ref_name": ref_path.name}
    with open(candidates[0], encoding="utf-8") as f:
        data = json.load(f)

    # ── reference 텍스트·표 추출 ──
    ref_text = _extract_ref_text(ref_path)
    ref_tables = _extract_ref_tables(ref_path)
    ref_acc_table = _find_accountability_table_in_ref(ref_tables)

    total_dd = data.get("total_delay_days")

    # CHECK 1: total_delay_days가 reference 텍스트에 등장하는지
    # PDF 텍스트 역순 추출 대응: "일184" 형태도 함께 탐색
    if total_dd:
        fwd = bool(_re.search(str(total_dd) + r"\s*일", ref_text))   # "184일"
        rev = bool(_re.search(r"일\s*" + str(total_dd), ref_text))   # "일184" (PDF 역순)
        found_in_ref = fwd or rev

        # reference 텍스트가 너무 짧으면 추출 실패로 간주 → warning으로 격하
        ref_too_short = len(ref_text.strip()) < 500
        if ref_too_short:
            checks.append(_check_item(
                f"total_delay_days({total_dd}일) — reference에 등장",
                None,  # type: ignore
                "reference 텍스트 추출 실패 (스캔 PDF 또는 추출 오류) — 수동 확인 필요"
            ))
        else:
            checks.append(_check_item(
                f"total_delay_days({total_dd}일) — reference에 등장",
                found_in_ref,
                "" if found_in_ref else f"reference에서 '{total_dd}일' 미발견 (분산 기재 또는 reference 불일치 가능)"
            ))
    else:
        checks.append(_check_item("total_delay_days 설정 여부", False, "data.json에 total_delay_days 없음"))

    # CHECK 2: accountability_diagram delay_days 합 = total_delay_days
    diag = data.get("accountability_diagram", [])
    diag_sum = sum(r.get("delay_days", 0) for r in diag if isinstance(r, dict))
    checks.append(_check_item(
        f"귀책 도식표 합계({diag_sum}일) = total_delay_days({total_dd}일)",
        diag_sum == total_dd,
        "" if diag_sum == total_dd else f"불일치: 합계 {diag_sum} ≠ {total_dd}"
    ))

    # CHECK 3: output docx의 귀책 도식표 행 수 = accountability_diagram 행 수
    # (reference는 법령 조문 인용 표만 있고 프로젝트별 요약 표가 없으므로, output docx를 직접 검증)
    diag_rows = len([r for r in diag if isinstance(r, dict)])
    out_docx_files = sorted(project_dir.glob("02_귀책분석_*.docx"))
    out_acc_table: "list[list[str]] | None" = None
    if out_docx_files:
        out_tables = _extract_ref_tables(out_docx_files[-1])  # 최신 파일
        # output의 귀책 요약 표: 헤더에 '공기지연 사유' 또는 '귀책' 포함
        _ACC_KWS = {"공기지연 사유", "공기지연사유", "귀책사유", "귀책", "지연사유"}
        for t in reversed(out_tables):
            if not t or len(t[0]) < 2:
                continue
            h = " ".join(c.replace(" ", "") for c in t[0])
            if any(kw.replace(" ", "") in h for kw in _ACC_KWS):
                # 법령 조문 열거 패턴 제외: 데이터 첫 행이 "1)", "가.", "①" 등으로 시작
                import re as _re2
                _LAW2 = _re2.compile(r"^(?:[0-9]+\)|[가-힣]\.|[①-⑳]|제\s*[0-9]+)")
                dr = t[1:] if len(t) > 1 else []
                if dr and _LAW2.match(dr[0][0].strip() if dr[0] else ""):
                    continue
                out_acc_table = t
                break

    if out_acc_table:
        sum_kws2 = {"합계", "소계", "계", "total", "sum"}
        out_data_rows = [
            r for r in out_acc_table[1:]
            if not any(cell.strip() in sum_kws2 or cell.strip().lower() in sum_kws2
                       for cell in r)
        ]
        out_row_count = len(out_data_rows)
        match = (diag_rows == out_row_count)
        checks.append(_check_item(
            f"output docx 귀책 표 행 수 일치 (data {diag_rows}행 vs docx {out_row_count}행)",
            match,
            "" if match else f"data {diag_rows}행 ≠ docx {out_row_count}행 (생성 누락 의심)"
        ))
    else:
        checks.append(_check_item(
            "output docx 귀책 표 행 수 일치",
            None,  # type: ignore
            "output docx에서 귀책 표를 찾지 못함 (generate 미실행 또는 형식 변경)"
        ))

    # CHECK 4: output docx 귀책 표 컬럼 수 (3=지방계약/민간, 5=국가계약)
    if out_acc_table and out_acc_table[0]:
        col_count = len(out_acc_table[0])
        ok4 = col_count in (3, 5)
        checks.append(_check_item(
            f"output docx 귀책 표 컬럼 수 ({col_count}컬럼)",
            ok4,
            f"{'지방계약/민간' if col_count == 3 else ('국가계약' if col_count == 5 else '비표준')} 형식"
            if ok4 else f"예상 외 컬럼 수: {col_count}"
        ))
    else:
        checks.append(_check_item("output docx 귀책 표 컬럼 수", None, "귀책 표 미발견"))  # type: ignore

    # CHECK 5: 필수 섹션 존재 여부 (data.json 기준)
    req_fields = {
        "background_paragraphs": "배경 단락",
        "pre_diagram_paragraphs": "도식표 전 단락",
        "conclusion_paragraphs": "결론 단락",
        "summary": "종합 요약",
    }
    for field, label in req_fields.items():
        val = data.get(field)
        ok = bool(val) and (len(val) >= 1 if isinstance(val, list) else len(str(val)) > 10)
        checks.append(_check_item(f"필수 섹션 — {label}", ok, "" if ok else f"'{field}' 비어 있음"))

    # CHECK 6: responsible_party 비표준값 감지 (줄바꿈·괄호 내용 제거 후 판단)
    # 표준값: 국가/지방계약(발주처·발주자·시공사·수급인) + 민간계약(도급인·수급인) + 공동귀책
    _STD_RP = {"발주처", "발주자", "도급인", "시공사", "수급인", "공동귀책", ""}
    def _normalize_rp(v: str) -> str:
        import re as _re2
        v = v.replace("\n", "").strip()
        # 괄호/부연 포함 값에서 기본 키워드만 추출
        # 예: "발주처(한국가스공사)" → "발주처", "도급인 (수급인의 책임 없는 사유)" → "도급인"
        m = _re2.match(r"^(발주처|발주자|도급인|시공사|수급인|공동귀책)", v)
        return m.group(1) if m else v
    non_std = [
        r.get("responsible_party", "") for r in diag
        if isinstance(r, dict) and _normalize_rp(r.get("responsible_party", "")) not in _STD_RP
    ]
    checks.append(_check_item(
        "responsible_party 표준값 사용 여부",
        len(non_std) == 0,
        "" if not non_std else f"비표준값: {non_std[:3]}"
    ))

    # CHECK 7: items에 delay_days 또는 causal_description이 채워져 있는지
    items = data.get("items", [])
    empty_items = [
        i.get("no", idx + 1) for idx, i in enumerate(items)
        if isinstance(i, dict) and not (i.get("causal_description") or i.get("subject"))
    ]
    checks.append(_check_item(
        "items 내용 충실도 (causal_description/subject 채워짐)",
        len(empty_items) == 0,
        "" if not empty_items else f"비어있는 항목 no: {empty_items[:5]}"
    ))

    ok_count = sum(1 for _, ok, _ in checks if ok is True)
    total_count = sum(1 for _, ok, _ in checks if ok is not None)

    return {
        "checks": checks,
        "score": (ok_count, total_count),
        "ref_name": ref_path.name,
    }


def _print_compare_result(project_name: str, result: dict):
    """비교 결과를 콘솔에 출력한다."""
    print(f"\n[{project_name}]")
    if "error" in result:
        print(f"  ❌ {result['error']}")
        return

    print(f"  참조: reference/{result['ref_name']}")
    for label, ok, detail in result["checks"]:
        if ok is True:
            icon = "✅"
        elif ok is False:
            icon = "❌"
        else:
            icon = "⚠️ "
        line = f"  {icon} {label}"
        if detail:
            line += f"\n       → {detail}"
        print(line)

    ok_n, total_n = result["score"]
    pct = int(ok_n / total_n * 100) if total_n else 0
    print(f"  점수: {ok_n}/{total_n} ({pct}%)")


def cmd_compare(project_name: str | None):
    """단일 프로젝트 output vs reference 품질 비교."""
    _print_header()
    print("\n[품질 검증] output vs reference 비교")
    print("─" * 60)

    project_name, project_dir = _resolve_project_dir(project_name)
    ref_path = _find_reference_file(project_dir)

    if not ref_path:
        print(f"\n[오류] '{project_name}' 에 매칭되는 reference 파일을 찾지 못했습니다.")
        print("  reference/ 폴더에 해당 프로젝트의 완성본 보고서가 있는지 확인하세요.")
        sys.exit(1)

    result = _compare_one(project_dir, ref_path)
    _print_compare_result(project_name, result)
    print()


def cmd_compare_all():
    """전체 프로젝트 output vs reference 품질 비교 (일괄)."""
    _print_header()
    print("\n[품질 검증 — 전체] output 14개 vs reference 대조")
    print("─" * 60)

    out_dir = config.OUTPUT_DIR
    if not out_dir.exists():
        print("[오류] output/ 폴더가 없습니다.")
        sys.exit(1)

    project_dirs = sorted([d for d in out_dir.iterdir() if d.is_dir()])
    if not project_dirs:
        print("[오류] output/ 아래 프로젝트 폴더가 없습니다.")
        sys.exit(1)

    summary_rows: list[tuple[str, int, int, list[str]]] = []  # name, ok, total, failures

    for pd in project_dirs:
        ref_path = _find_reference_file(pd)
        if not ref_path:
            print(f"\n  ⚠️  [{pd.name}] reference 매칭 없음 — 건너뜀")
            continue

        result = _compare_one(pd, ref_path)
        _print_compare_result(pd.name, result)

        ok_n, total_n = result.get("score", (0, 0))
        failures = [label for label, ok, _ in result.get("checks", []) if ok is False]
        summary_rows.append((pd.name, ok_n, total_n, failures))

    # ── 전체 요약 ──
    print("\n" + "=" * 60)
    print("  전체 요약")
    print("=" * 60)
    print(f"  {'프로젝트':<40} {'점수':>8}  {'판정'}")
    print("  " + "-" * 55)
    for name, ok_n, total_n, failures in summary_rows:
        pct = int(ok_n / total_n * 100) if total_n else 0
        verdict = "✅" if pct == 100 else ("⚠️ " if pct >= 70 else "❌")
        short = name[:38]
        print(f"  {short:<40} {ok_n}/{total_n} ({pct:>3}%)  {verdict}")

    # ── 반복 실패 패턴 ──
    from collections import Counter as _Counter
    all_failures = [f for _, _, _, fs in summary_rows for f in fs]
    if all_failures:
        print()
        print("  반복 실패 항목 (2개 이상 프로젝트):")
        for label, cnt in _Counter(all_failures).most_common():
            if cnt >= 2:
                print(f"    [{cnt}회] {label}")
    print("─" * 60)

    # ── 결과 파일 저장 ──
    result_path = out_dir / "compare_result.txt"
    # (stdout 캡처 없이 간단히 요약만 저장)
    lines = ["output vs reference 품질 검증 결과\n"]
    for name, ok_n, total_n, failures in summary_rows:
        pct = int(ok_n / total_n * 100) if total_n else 0
        lines.append(f"{name}: {ok_n}/{total_n} ({pct}%)")
        for f in failures:
            lines.append(f"  ❌ {f}")
    result_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"\n  결과 저장: {result_path}")


# ── 유틸리티: rescan-all ─────────────────────────────────────────────────────

def cmd_rescan_all():
    """
    전체 프로젝트 일괄 재스캔 + 재prepare.

    - scan_result.json 백업 (scan_result_backup.json, 없는 경우만)
    - 새 코드로 재스캔 (캐시 재사용, OCR 없음)
    - scan_result.json 이 없는 폴더는 건너뜀
    - 이전 scan_result.json vs 새 결과 비교 → 신규 발견 항목 리포트
    """
    import shutil

    _print_header()
    print("\n[일괄 재스캔 + 재prepare]")
    print("─" * 60)

    out_dir = config.OUTPUT_DIR
    if not out_dir.exists():
        print("[오류] output/ 폴더가 없습니다.")
        sys.exit(1)

    def _backup(proj_dir: Path) -> bool:
        src = proj_dir / "scan_result.json"
        dst = proj_dir / "scan_result_backup.json"
        if src.exists() and not dst.exists():
            shutil.copy2(src, dst)
            return True
        return False

    def _new_items(proj_dir: Path) -> list[dict]:
        old_p = proj_dir / "scan_result_backup.json"
        new_p = proj_dir / "scan_result.json"
        if not old_p.exists() or not new_p.exists():
            return []
        old = json.loads(old_p.read_text(encoding="utf-8"))
        new = json.loads(new_p.read_text(encoding="utf-8"))
        old_files = {Path(i.get("file_path", "")).name for i in old.get("items", [])}
        return [i for i in new.get("items", []) if Path(i.get("file_path", "")).name not in old_files]

    results: list[dict] = []
    proj_dirs = sorted([d for d in out_dir.iterdir() if d.is_dir()])

    for proj_dir in proj_dirs:
        scan_json = proj_dir / "scan_result.json"
        if not scan_json.exists():
            continue
        name = proj_dir.name
        print(f"\n[{name}]")

        backed = _backup(proj_dir)
        if backed:
            print("  백업 완료 → scan_result_backup.json")

        d = json.loads(scan_json.read_text(encoding="utf-8"))
        vendor_dirs_raw = d.get("vendor_dirs") or ([d["vendor_dir"]] if d.get("vendor_dir") else [])
        if not vendor_dirs_raw or not vendor_dirs_raw[0]:
            print("  [SKIP] vendor_dirs 없음")
            continue

        vendor_dirs = [Path(p) for p in vendor_dirs_raw]

        # ── scan (직접 함수 호출 — exe 환경에서도 동작) ─────────────────────
        try:
            scan(vendor_dirs, proj_dir)
        except Exception as e:
            print(f"  [ERROR] scan 실패: {e}")
            results.append({"name": name, "ok": False, "step": "scan"})
            continue

        # ── prepare (직접 함수 호출) ─────────────────────────────────────────
        try:
            build_prompt(proj_dir, name)
        except Exception as e:
            print(f"  [ERROR] prepare 실패: {e}")
            results.append({"name": name, "ok": False, "step": "prepare"})
            continue

        news = _new_items(proj_dir)
        if news:
            print(f"  신규 발견 {len(news)}건:")
            for item in news:
                date = item.get("date", "")
                subj = item.get("subject", "")[:40]
                print(f"    - {date} | {subj}")
        else:
            print("  신규 항목 없음")

        results.append({"name": name, "ok": True, "new": len(news)})

    print("\n" + "=" * 60)
    print("  완료 요약")
    print("=" * 60)
    for r in results:
        if r["ok"]:
            flag = f"  [신규 {r['new']}건]" if r.get("new") else ""
            print(f"  ✅ {r['name']}{flag}")
        else:
            print(f"  ❌ {r['name']} ({r['step']} 실패)")
    print("─" * 60)


# ── 콤보 명령 ────────────────────────────────────────────────────────────────

def cmd_scan_prepare(raw_args: list[str]):
    """scan + prepare 한 번에 실행."""
    cmd_scan(raw_args)
    # scan이 current_project를 저장하므로 project_name=None으로 자동 참조
    print()
    cmd_prepare(None)


def cmd_finish(project_name: str | None):
    """validate 후 오류 없으면 generate 자동 실행."""
    ok = cmd_validate(project_name, _stop_on_error=False)
    if ok:
        print()
        cmd_generate(project_name)


# ── 인터랙티브 메뉴 (더블클릭 실행 시) ─────────────────────────────────────

def _ask(prompt: str, default: str = "") -> str:
    """사용자 입력을 받는다. 빈 입력이면 default 반환."""
    suffix = f" [{default}]" if default else ""
    try:
        val = input(f"{prompt}{suffix}: ").strip()
    except (EOFError, KeyboardInterrupt):
        return default
    return val if val else default


def _pause():
    """작업 완료 후 창이 바로 닫히지 않도록 대기한다."""
    try:
        input("\n  Enter 를 누르면 메뉴로 돌아갑니다...")
    except (EOFError, KeyboardInterrupt):
        pass


def cmd_interactive():
    """더블클릭 실행 시 진입하는 인터랙티브 메뉴."""
    while True:
        _print_header()

        # 마지막 작업 프로젝트 표시
        last = config.load_current_project()
        if last:
            print(f"  마지막 프로젝트: {last}")
        print()
        print("  ┌─────────────────────────────────────────────────────────┐")
        print("  │  작업 순서                                              │")
        print("  │                                                         │")
        print("  │  [1] 스캔 + 프롬프트 생성   ← 수신자료 폴더 경로 입력  │")
        print("  │       ↓                                                 │")
        print("  │  [--] Claude Code 가 JSON 작성  (이 프로그램 밖에서)   │")
        print("  │       ↓                                                 │")
        print("  │  [2] docx 생성              ← JSON 작성 완료 후        │")
        print("  │       ↓                                                 │")
        print("  │  [3] 품질 검증 (단일)   또는   [4] 전체 품질 검증      │")
        print("  │                                                         │")
        print("  │  ─────────────────────────────────────────────────────  │")
        print("  │  0. 종료                                                │")
        print("  └─────────────────────────────────────────────────────────┘")
        print()

        choice = _ask("  번호 선택").strip()

        if choice == "0":
            print("\n  종료합니다.")
            break

        elif choice == "1":
            # ── scanprepare ───────────────────────────────────────
            print()
            print("  수신자료 폴더 경로를 입력하세요.")
            print("  (여러 경로는 세미콜론(;)으로 구분)")
            path_str = _ask("  경로")
            if not path_str:
                print("  [취소] 경로가 입력되지 않았습니다.")
                _pause()
                continue

            project = _ask("  프로젝트명 (엔터 = 자동 감지)")
            raw: list[str] = [p.strip().strip('"').strip("'")
                               for p in path_str.split(";") if p.strip()]
            scan_args = raw[:]
            if project:
                scan_args += ["--project", project]
            print()
            try:
                cmd_scan_prepare(scan_args)
            except SystemExit:
                pass

            # ── 완료 후 안내 (prompt 파일이 실제로 생성된 경우만) ───────
            proj_now = config.load_current_project()
            if proj_now:
                proj_dir = config.get_project_dir(proj_now)
                prompt_ok = (proj_dir / "prompt_for_claude.md").exists()
                scan_ok   = (proj_dir / "scan_result.json").exists()
                print()
                if prompt_ok:
                    print("  ══════════════════════════════════════════════════")
                    print("  ★ 1단계 완료 — 지금 해야 할 일")
                    print("  ══════════════════════════════════════════════════")
                    print(f"  저장 위치: output\\{proj_dir.name}\\")
                    print()
                    print("  [1] scan_summary.md 를 열어 공문 목록을 확인하세요.")
                    print("      → 관련 없는 공문이 있으면 scan_result.json 에서 해당 줄 삭제")
                    print("      → OCR ⚠️ 표시 항목은 원본 파일로 날짜·공문번호 확인")
                    print()
                    print("  [2] Claude Code 에 다음과 같이 입력하세요:")
                    print(f'      "output\\{proj_dir.name}\\prompt_for_claude.md 읽고')
                    print(f'       귀책분석_data.json 생성해줘"')
                    print()
                    print("  [3] JSON 저장 완료 후 이 프로그램에서 [2] docx 생성 실행")
                    print("  ══════════════════════════════════════════════════")
                elif scan_ok:
                    # 스캔은 됐지만 prepare 실패 (reference_patterns.md 없음 등)
                    print("  ══════════════════════════════════════════════════")
                    print("  ⚠️  스캔은 완료됐으나 프롬프트 생성에 실패했습니다.")
                    print(f"  저장 위치: output\\{proj_dir.name}\\")
                    print()
                    print("  배포 패키지에 output\\reference_patterns.md 파일이")
                    print("  있는지 확인하고 관리자에게 문의하세요.")
                    print("  ══════════════════════════════════════════════════")
            _pause()

        elif choice == "2":
            # ── finish ────────────────────────────────────────────
            print()
            project = _ask("  프로젝트명 (엔터 = 마지막 프로젝트)", last or "")
            print()
            try:
                cmd_finish(project or None)
            except SystemExit:
                pass
            _pause()

        elif choice == "3":
            # ── compare ───────────────────────────────────────────
            print()
            project = _ask("  프로젝트명 (엔터 = 마지막 프로젝트)", last or "")
            print()
            try:
                cmd_compare(project or None)
            except SystemExit:
                pass
            _pause()

        elif choice == "4":
            # ── compare-all ───────────────────────────────────────
            print()
            try:
                cmd_compare_all()
            except SystemExit:
                pass
            _pause()

        else:
            print(f"\n  '{choice}' 는 유효하지 않은 선택입니다.")
            _pause()


# ── CLI 파싱 ─────────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]

    if not args:
        # 인수 없이 실행 — GUI (더블클릭 포함)
        try:
            from gui import run_gui
            run_gui()
        except Exception as _gui_err:
            if getattr(sys, "frozen", False):
                # exe 환경: 콘솔 없으므로 input() 사용 불가 — 오류 다이얼로그만 표시
                try:
                    import tkinter.messagebox as _mb
                    _mb.showerror("실행 오류",
                                  f"GUI를 시작할 수 없습니다.\n\n{_gui_err}")
                except Exception:
                    pass
            else:
                # 개발 환경 폴백: 콘솔 메뉴
                try:
                    cmd_interactive()
                except (KeyboardInterrupt, EOFError):
                    print("\n  종료합니다.")
        return

    cmd = args[0].lower()

    if cmd == "learn":
        cmd_learn()

    elif cmd == "scan":
        cmd_scan(args[1:])

    elif cmd == "prepare":
        project_name = args[1] if len(args) > 1 else None
        cmd_prepare(project_name)

    elif cmd == "validate":
        project_name = args[1] if len(args) > 1 else None
        cmd_validate(project_name)

    elif cmd == "generate":
        project_name = args[1] if len(args) > 1 else None
        cmd_generate(project_name)

    elif cmd in ("scanprepare", "scan-prepare", "sp"):
        cmd_scan_prepare(args[1:])

    elif cmd in ("finish", "checkgenerate", "check-generate", "cg"):
        project_name = args[1] if len(args) > 1 else None
        cmd_finish(project_name)

    elif cmd == "compare":
        project_name = args[1] if len(args) > 1 else None
        cmd_compare(project_name)

    elif cmd in ("compare-all", "compareall", "ca"):
        cmd_compare_all()

    elif cmd in ("rescan-all", "rescanall", "ra"):
        cmd_rescan_all()

    else:
        print(f"알 수 없는 명령: {cmd}")
        print(__doc__)
        sys.exit(1)


if __name__ == "__main__":
    main()
