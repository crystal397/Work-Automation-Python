
import sys
import os
import io
from pathlib import Path
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

def extract_doc_info(filepath):
    """Extract headings, paragraphs, and tables from a docx file."""
    doc = Document(filepath)

    headings = []
    paragraphs_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith('Heading'):
            headings.append((para.style.name, text))
        paragraphs_text.append((para.style.name, text))

    tables_data = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)
        tables_data.append(rows_data)

    return headings, paragraphs_text, tables_data


def find_accountability_table(tables_data):
    """Find the 귀책사유 정리표 (accountability table) - looking for 3-col table with 공기지연사유."""
    for i, table in enumerate(tables_data):
        for row in table:
            combined = ' '.join(row)
            if '공기지연사유' in combined or '귀책' in combined or '비용부담' in combined or '관련근거' in combined:
                return i, table
    # Also try finding any 3-column table
    for i, table in enumerate(tables_data):
        if table and len(table[0]) == 3:
            return i, table
    return None, None


def print_separator(title):
    print("\n" + "="*80)
    print(f"  {title}")
    print("="*80)


def analyze_project(project_name, gen_path, ref_path):
    print_separator(f"PROJECT: {project_name}")

    print(f"\n[Generated]: {os.path.basename(gen_path)}")
    print(f"[Reference]: {os.path.basename(ref_path)}")

    # Load documents
    try:
        gen_headings, gen_paras, gen_tables = extract_doc_info(gen_path)
        print(f"\n  Generated: {len(gen_paras)} paragraphs, {len(gen_tables)} tables, {len(gen_headings)} headings")
    except Exception as e:
        print(f"  ERROR loading generated: {e}")
        return

    try:
        ref_headings, ref_paras, ref_tables = extract_doc_info(ref_path)
        print(f"  Reference: {len(ref_paras)} paragraphs, {len(ref_tables)} tables, {len(ref_headings)} headings")
    except Exception as e:
        print(f"  ERROR loading reference: {e}")
        return

    # --- CHECK 1: 귀책사유 정리표 ---
    print("\n" + "-"*60)
    print("CHECK 1: 귀책사유 정리표 (accountability table)")
    print("-"*60)

    # Generated
    acc_idx, acc_table = find_accountability_table(gen_tables)
    if acc_table:
        print(f"\n  [Generated] Found accountability table at table index {acc_idx} ({len(acc_table)} rows):")
        for j, row in enumerate(acc_table):
            if j == 0:
                print(f"    HEADER: {row}")
            else:
                print(f"    Row {j}: {row}")
    else:
        print("\n  [Generated] No accountability table found!")
        # Show all tables summary
        print(f"  All tables ({len(gen_tables)} total):")
        for i, t in enumerate(gen_tables):
            cols = len(t[0]) if t else 0
            header = t[0] if t else []
            print(f"    Table {i}: {len(t)} rows x {cols} cols | header: {header}")

    # Reference
    ref_acc_idx, ref_acc_table = find_accountability_table(ref_tables)
    if ref_acc_table:
        print(f"\n  [Reference] Found accountability table at table index {ref_acc_idx} ({len(ref_acc_table)} rows):")
        for j, row in enumerate(ref_acc_table):
            if j == 0:
                print(f"    HEADER: {row}")
            else:
                print(f"    Row {j}: {row}")
    else:
        print("\n  [Reference] No accountability table found!")
        print(f"  All tables ({len(ref_tables)} total):")
        for i, t in enumerate(ref_tables):
            cols = len(t[0]) if t else 0
            header = t[0] if t else []
            print(f"    Table {i}: {len(t)} rows x {cols} cols | header: {header}")

    # --- CHECK 2: 섹션 헤딩 ---
    print("\n" + "-"*60)
    print("CHECK 2: 섹션 헤딩 (chapter headings)")
    print("-"*60)

    print("\n  [Generated] All headings:")
    for style, text in gen_headings:
        print(f"    [{style}] {text}")

    print("\n  [Reference] All headings:")
    for style, text in ref_headings:
        print(f"    [{style}] {text}")

    # Check for 제3장 vs 제2장
    gen_heading_texts = [t for _, t in gen_headings]
    has_3jang = any('제3장' in t or '3장' in t for t in gen_heading_texts)
    has_2jang = any('제2장' in t or '2장' in t for t in gen_heading_texts)
    print(f"\n  Contains 제3장: {has_3jang}")
    print(f"  Contains 제2장: {has_2jang}")

    # --- CHECK 3: detail_narratives 소제목 ---
    print("\n" + "-"*60)
    print("CHECK 3: detail_narratives 소제목 (▶ X차 공기연장)")
    print("-"*60)

    subheadings = [(s, t) for s, t in gen_paras if '▶' in t or '차 공기연장' in t]
    if subheadings:
        print(f"\n  [Generated] Found {len(subheadings)} subheading(s):")
        for s, t in subheadings:
            print(f"    [{s}] {t}")
    else:
        print("\n  [Generated] No ▶ subheadings found.")

    # Also check in reference
    ref_subheadings = [(s, t) for s, t in ref_paras if '▶' in t or '차 공기연장' in t]
    if ref_subheadings:
        print(f"\n  [Reference] Found {len(ref_subheadings)} subheading(s):")
        for s, t in ref_subheadings[:10]:
            print(f"    [{s}] {t}")
    else:
        print("\n  [Reference] No ▶ subheadings found (may use different style).")

    # --- CHECK 4: 전체 내용 비교 ---
    print("\n" + "-"*60)
    print("CHECK 4: 전체 내용 비교 (content comparison)")
    print("-"*60)

    # Key reference sections/keywords to look for
    key_sections = ['공기연장 개요', '귀책사유', '간접비', '공기지연', '계약', '발주처', '시공사', '공사개요',
                    '공기연장 경위', '귀책 분석', '공사기간', '지연', '총괄']

    gen_text_all = ' '.join([t for _, t in gen_paras])
    ref_text_all = ' '.join([t for _, t in ref_paras])

    print("\n  Key content checks (present in reference vs generated):")
    for kw in key_sections:
        in_ref = kw in ref_text_all
        in_gen = kw in gen_text_all
        status = "✓" if in_gen else "✗"
        print(f"    [{status}] '{kw}' - ref: {in_ref}, gen: {in_gen}")

    # Show first few and last few paragraphs of generated
    print(f"\n  [Generated] First 15 non-empty paragraphs:")
    count = 0
    for s, t in gen_paras:
        if t and count < 15:
            print(f"    [{s}] {t[:120]}")
            count += 1

    print(f"\n  [Generated] Last 10 non-empty paragraphs:")
    non_empty = [(s, t) for s, t in gen_paras if t]
    for s, t in non_empty[-10:]:
        print(f"    [{s}] {t[:120]}")

    # --- Show all tables summary ---
    print(f"\n  [Generated] All tables summary:")
    for i, t in enumerate(gen_tables):
        cols = len(t[0]) if t else 0
        print(f"    Table {i}: {len(t)} rows x {cols} cols")
        if t:
            print(f"      Header: {t[0]}")
            if len(t) > 1:
                print(f"      Row 1:  {t[1]}")
            if len(t) > 2:
                print(f"      Row 2:  {t[2]}")


# Define file paths (스크립트 위치 기준 상대경로)
_BASE = Path(__file__).parent
projects = [
    {
        "name": "송도 11-1공구",
        "gen": str(_BASE / "output" / "송도11-1공구 기반시설 건설공사 (1-1구역) 현장" / "02_귀책분석_송도11-1공구 기반시설 건설공사 (1-1구역) 현장_20260409.docx"),
        "ref": str(_BASE / "reference" / "송도 11-1공구 기반시설 건설공사(1-1구역)_공기연장 간접비 산정 보고서_260309.docx"),
    },
    {
        "name": "인덕원~동탄 5공구",
        "gen": str(_BASE / "output" / "인덕원~동탄 복선전철 제5공구 노반신설 기타공사" / "02_귀책분석_인덕원~동탄 복선전철 제5공구 노반신설 기타공사_20260409.docx"),
        "ref": str(_BASE / "reference" / "인덕원~동탄 5공구_공기연장 간접비 산정 보고서_260317.docx"),
    },
    {
        "name": "평택기지 2공구",
        "gen": str(_BASE / "output" / "평택기지-오산 천연가스 공급설비 건설공사 2공구" / "02_귀책분석_평택기지-오산 천연가스 공급설비 건설공사 2공구_20260409.docx"),
        "ref": str(_BASE / "reference" / "평택기지~오산 제2공구 천연가스 공급시설 건설공사 공기연장 간접비 보고서_260318.docx"),
    },
]

for p in projects:
    analyze_project(p["name"], p["gen"], p["ref"])

print("\n" + "="*80)
print("  ANALYSIS COMPLETE")
print("="*80)
