"""
공기연장 간접비 산정 보고서 — Word 템플릿 생성 스크립트

실행:
    python make_template.py

출력:
    reference/보고서_템플릿.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── 상수 ─────────────────────────────────────────────────────────────────────

FONT_MAIN   = "한컴바탕"   # 본문/헤딩 공통
FONT_BODY   = "함초롬바탕" # 바탕글
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_GRAY  = RGBColor(89, 89, 89)

PLACEHOLDER_LIGHT = "[      ]"   # 짧은 값 입력란
PH = "【 입력 】"                 # 눈에 띄는 입력란


# ── 헬퍼: 단락 서식 ──────────────────────────────────────────────────────────

def _fmt(para, font_name=FONT_MAIN, size_pt=11, bold=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before_pt=0, space_after_pt=6):
    """단락 폰트·크기·정렬·간격 일괄 설정."""
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after  = Pt(space_after_pt)
    for run in para.runs:
        run.font.name        = font_name
        run.font.size        = Pt(size_pt)
        run.font.bold        = bold
        run.font.color.rgb   = COLOR_BLACK
        # 한글 폰트 설정 (동아시아)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font_name)
    return para


def add_paragraph(doc, text="", style=None, font_name=FONT_MAIN, size_pt=11,
                  bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=0, space_after_pt=6):
    """단락 추가 + 서식 적용."""
    if style:
        para = doc.add_paragraph(text, style=style)
    else:
        para = doc.add_paragraph(text)
    if not para.runs and text:
        para.add_run(text)
    _fmt(para, font_name=font_name, size_pt=size_pt, bold=bold, align=align,
         space_before_pt=space_before_pt, space_after_pt=space_after_pt)
    return para


def add_heading(doc, text, level=1):
    """
    level 1 → 장(章) 제목  : 16pt bold, 앞 12pt 여백
    level 2 → 절(節) 제목  : 14pt bold, 앞 9pt
    level 3 → 항(項) 제목  : 12pt bold, 앞 6pt
    """
    cfg = {
        1: dict(size_pt=16, space_before_pt=12, space_after_pt=6),
        2: dict(size_pt=14, space_before_pt=9,  space_after_pt=4),
        3: dict(size_pt=12, space_before_pt=6,  space_after_pt=3),
    }[level]
    return add_paragraph(doc, text, bold=True, **cfg)


def add_body(doc, text=""):
    """일반 본문 단락."""
    return add_paragraph(doc, text, size_pt=11, space_after_pt=4)


def add_note(doc, text=""):
    """회색 주석/여백 안내 단락."""
    para = add_paragraph(doc, text, size_pt=10, space_after_pt=3)
    for run in para.runs:
        run.font.color.rgb = COLOR_GRAY
        run.font.italic    = True
    return para


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx_break_type := __import__(
        "docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:br"))
    docx_break_type.set(qn("w:type"), "page")
    return para


# ── 헬퍼: 표 ─────────────────────────────────────────────────────────────────

def set_cell(cell, text, font_name=FONT_MAIN, size_pt=10,
             bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             bg_color: str | None = None):
    """셀 텍스트·서식 설정."""
    cell.text = ""
    para = cell.paragraphs[0]
    run  = para.add_run(text)
    run.font.name  = font_name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    para.alignment = align
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    if bg_color:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg_color)
        tcPr.append(shd)
    return cell


def set_col_widths(table, widths_cm):
    """열 너비 설정 (cm 단위 리스트)."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def set_table_border(table):
    """표 전체 테두리 설정."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


HEADER_BG = "D9D9D9"  # 헤더 행 배경 (연회색)


def add_header_row(table, headers, widths_cm=None):
    """표 헤더 행 추가 (회색 배경 + bold)."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i < len(row.cells):
            set_cell(row.cells[i], h, bold=True, bg_color=HEADER_BG)
    if widths_cm:
        set_col_widths(table, widths_cm)
    set_table_border(table)


def add_data_row(table, values, align=WD_ALIGN_PARAGRAPH.CENTER):
    """데이터 행 추가."""
    row = table.add_row()
    for i, v in enumerate(values):
        if i < len(row.cells):
            set_cell(row.cells[i], v, align=align)
    return row


# ── 섹션별 생성 함수 ──────────────────────────────────────────────────────────

def make_cover(doc):
    """표지"""
    for _ in range(6):
        doc.add_paragraph()

    p = add_paragraph(doc, PH + "  주식회사",
                      size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)

    for _ in range(2):
        doc.add_paragraph()

    add_paragraph(doc, PH + "  공사",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)

    add_paragraph(doc, "공기연장 간접비 산정 보고서",
                  size_pt=24, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)

    for _ in range(8):
        doc.add_paragraph()

    add_paragraph(doc, PH + " 년  " + PH + " 월",
                  size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    add_paragraph(doc, "사단법인 한국건설관리연구원",
                  size_pt=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)


def make_submission(doc):
    """제출문"""
    doc.add_page_break()
    add_paragraph(doc, "제   출   문",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=16)

    add_body(doc, PH + " (발주처명)  귀중")
    doc.add_paragraph()
    add_body(doc,
             "사단법인 한국건설관리연구원은 귀 기관으로부터 의뢰받은 "
             "「" + PH + " 공사」의 공기연장 간접비 산정 용역을 완료하고, "
             "그 결과를 다음과 같이 제출합니다.")
    doc.add_paragraph()

    # 용역 정보 표
    tbl = doc.add_table(rows=4, cols=2)
    headers_data = [
        ("용 역 명", PH + "  공사 공기연장 간접비 산정"),
        ("용역기간", PH + " 년  " + PH + " 월  " + PH + " 일  ~  " + PH + " 년  " + PH + " 월  " + PH + " 일"),
        ("수행기관", "사단법인 한국건설관리연구원"),
        ("대  표  자", PH),
    ]
    for i, (k, v) in enumerate(headers_data):
        set_cell(tbl.rows[i].cells[0], k, bold=True, bg_color=HEADER_BG,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(tbl.rows[i].cells[1], v,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
    set_col_widths(tbl, [4.0, 11.5])
    set_table_border(tbl)

    doc.add_paragraph()
    add_paragraph(doc, PH + " 년  " + PH + " 월",
                  size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=4)
    add_paragraph(doc, "사단법인 한국건설관리연구원",
                  size_pt=12, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)
    add_paragraph(doc, "원  장  " + PH + "  (인)",
                  size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)


def make_summary(doc):
    """요약문 — 산정 결과"""
    doc.add_page_break()
    add_paragraph(doc, "요      약      문",
                  size_pt=20, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=12)

    add_paragraph(doc, "< 산 정 결 과 >",
                  size_pt=13, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)

    add_note(doc, "※ 공사명: " + PH)
    add_note(doc, "※ 산정 기간: " + PH + " ~ " + PH + "  (" + PH + " 일)")
    doc.add_paragraph()

    # 원가계산서 집계표
    add_paragraph(doc, "○ 원가계산서  (단위: 원, 영세율)",
                  size_pt=11, space_after_pt=3)
    tbl = doc.add_table(rows=1, cols=5)
    add_header_row(tbl,
                   ["비  목", "간접노무비", "경  비", "일반관리비", "이  윤"],
                   widths_cm=[3.5, 3.5, 3.5, 3.5, 3.5])
    for label in ["원도급", "하도급 (해당 시)", "합  계"]:
        add_data_row(tbl, [label, PH, PH, PH, PH])

    doc.add_paragraph()

    # 합계 강조
    tbl2 = doc.add_table(rows=1, cols=3)
    add_header_row(tbl2,
                   ["구  분", "공기연장 간접비 합계 (원)", "비  고"],
                   widths_cm=[4.0, 7.5, 4.5])
    add_data_row(tbl2, ["최 종 청구액", PH, "VAT 별도"])


def make_toc(doc):
    """목차"""
    doc.add_page_break()
    add_paragraph(doc, "목  차",
                  size_pt=16, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=10, space_after_pt=12)

    toc_items = [
        (1, "제1장.  개요"),
        (2, "1.1  공사의 개요"),
        (2, "1.2  계약현황"),
        (3, "1.2.1  공사 계약현황"),
        (1, "제2장.  청구의 사유"),
        (2, "2.1  청구의 사유 개요"),
        (3, "2.1.1  계약에 의한 청구권"),
        (3, "2.1.2  사정변경의 원칙"),
        (2, "2.2  공기지연의 귀책사유 분석"),
        (3, "2.2.1  공기지연 귀책사유 관련 근거"),
        (3, "2.2.2  공기지연 귀책사유 구분"),
        (1, "제3장.  청구의 방법"),
        (2, "3.1  공기연장 기간의 산정"),
        (3, "3.1.1  공기연장 일수의 산정방식"),
        (3, "3.1.2  공기연장 일수 산정"),
        (2, "3.2  공기연장에 따른 간접비 산정방식"),
        (3, "3.2.1  간접노무비"),
        (3, "3.2.2  경비"),
        (3, "3.2.3  일반관리비"),
        (3, "3.2.4  이윤"),
        (3, "3.2.5  보증수수료"),
        (1, "제4장.  공기연장에 따른 간접비 산정 결과"),
        (2, "4.1  간접비 집계표"),
        (2, "4.2  간접노무비"),
        (3, "4.2.1  대상인원"),
        (3, "4.2.2  급여내역"),
        (2, "4.3  경비"),
        (3, "4.3.1  직접계상비목"),
        (3, "4.3.2  승률계상비목"),
        (2, "4.4  일반관리비 및 이윤"),
        (3, "4.4.1  일반관리비"),
        (3, "4.4.2  이윤"),
        (1, "제5장.  첨부자료"),
        (2, "5.1  공기연장 간접비 산정근거"),
        (2, "5.2  계약문서"),
        (2, "5.3  수발신문서"),
        (2, "5.4  공기연장 간접비 증빙자료"),
    ]
    indent_map = {1: 0, 2: Cm(0.8), 3: Cm(1.6)}
    size_map   = {1: 12, 2: 11, 3: 10}
    bold_map   = {1: True, 2: False, 3: False}

    for lvl, text in toc_items:
        para = add_paragraph(doc, text,
                             size_pt=size_map[lvl],
                             bold=bold_map[lvl],
                             space_after_pt=2)
        para.paragraph_format.left_indent = indent_map[lvl]


# ── 제1장 ────────────────────────────────────────────────────────────────────

def make_ch1(doc):
    doc.add_page_break()
    add_heading(doc, "제1장.  개요", level=1)

    # 1.1 공사의 개요
    add_heading(doc, "1.1  공사의 개요", level=2)

    tbl = doc.add_table(rows=7, cols=3)
    header_data = [
        ("계약명",   PH + "  공사",    ""),
        ("위치",     PH,               ""),
        ("공사규모", PH,               ""),
        ("발주청",   PH,               ""),
        ("계약상대자", PH + "  주식회사", ""),
        ("공사기간", PH + " ~ " + PH,  ""),
        ("계약금액", PH + " 천원",     ""),
    ]
    set_cell(tbl.rows[0].cells[0], "구  분", bold=True, bg_color=HEADER_BG)
    set_cell(tbl.rows[0].cells[1], "내  용", bold=True, bg_color=HEADER_BG)
    set_cell(tbl.rows[0].cells[2], "비  고", bold=True, bg_color=HEADER_BG)
    for i, (k, v, note) in enumerate(header_data):
        row = tbl.add_row()
        set_cell(row.cells[0], k, bold=True, bg_color=HEADER_BG,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], v, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], note, align=WD_ALIGN_PARAGRAPH.CENTER)
    # 첫 번째 행(미사용) 삭제 후 재구성 — 위 방식으로 충분
    set_col_widths(tbl, [3.5, 10.0, 2.5])
    set_table_border(tbl)
    doc.add_paragraph()

    # 1.2 계약현황
    add_heading(doc, "1.2  계약현황", level=2)
    add_heading(doc, "1.2.1  공사 계약현황", level=3)
    add_note(doc, "※ [단위 : 천원]")

    tbl2 = doc.add_table(rows=1, cols=6)
    add_header_row(tbl2,
                   ["구  분", "계약일", "계약금액(천원)", "착공일", "준공일", "비  고"],
                   widths_cm=[2.5, 3.0, 3.5, 3.0, 3.0, 1.5])
    for label in ["최  초", "1회 변경", "2회 변경", "합  계 / 연장일수"]:
        add_data_row(tbl2, [label, PH, PH, PH, PH, PH])
    doc.add_paragraph()


# ── 제2장 ────────────────────────────────────────────────────────────────────

def make_ch2(doc):
    doc.add_page_break()
    add_heading(doc, "제2장.  청구의 사유", level=1)

    # 2.1
    add_heading(doc, "2.1  청구의 사유 개요", level=2)
    add_heading(doc, "2.1.1  계약에 의한 청구권", level=3)
    add_body(doc,
             "「" + PH + " (근거법령)」에 따르면, 공사기간이 연장된 경우 계약상대자는 "
             "연장 기간에 대한 간접비를 발주자에게 청구할 수 있습니다.")
    add_body(doc,
             "관련 조항: " + PH + " 제" + PH + "조 (조항 내용 기재)")
    doc.add_paragraph()

    add_heading(doc, "2.1.2  사정변경의 원칙", level=3)
    add_body(doc,
             "민법상 사정변경의 원칙에 따라, 계약 체결 당시에 예측할 수 없었던 "
             "현저한 사정 변경이 발생한 경우 계약 조건의 변경 또는 추가 비용 청구가 가능합니다.")
    doc.add_paragraph()

    # 2.2
    add_heading(doc, "2.2  공기지연의 귀책사유 분석", level=2)
    add_heading(doc, "2.2.1  공기지연 귀책사유 관련 근거", level=3)
    add_body(doc, "본 공사의 공기지연 귀책사유와 관련한 주요 공문 및 계약 변경 내역은 다음과 같습니다.")
    doc.add_paragraph()

    # 공문 목록 표
    tbl = doc.add_table(rows=1, cols=4)
    add_header_row(tbl,
                   ["No.", "일자", "제목 (발신처 → 수신처)", "비고"],
                   widths_cm=[1.0, 2.8, 11.0, 1.7])
    for i in range(1, 6):
        add_data_row(tbl,
                     [str(i),
                      PH,
                      PH + "  ( " + PH + " → " + PH + " )",
                      ""])
    doc.add_paragraph()

    add_heading(doc, "2.2.2  공기지연 귀책사유 구분", level=3)
    add_body(doc, "위 공문 및 계약 변경 내역을 종합하면, 공기지연의 귀책사유는 아래와 같이 구분됩니다.")
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl2,
                   ["귀책 주체", "지연 사유", "지연 기간", "비고"],
                   widths_cm=[3.0, 8.0, 3.5, 2.0])
    for label in ["발주자 귀책", "설계 변경", "민원·인허가", "불가항력"]:
        add_data_row(tbl2, [label, PH, PH + " 일", ""])
    doc.add_paragraph()

    add_body(doc,
             "따라서 계약상대자는 연장된 공사기간에 대한 추가공사비를 발주자에게 청구할 수 있습니다.")


# ── 제3장 ────────────────────────────────────────────────────────────────────

def make_ch3(doc):
    doc.add_page_break()
    add_heading(doc, "제3장.  청구의 방법", level=1)

    # 3.1 공기연장 기간
    add_heading(doc, "3.1  공기연장 기간의 산정", level=2)
    add_heading(doc, "3.1.1  공기연장 일수의 산정방식", level=3)
    add_body(doc,
             "공기연장 일수는 변경계약서에 명시된 준공기한을 기준으로 산정하며, "
             "역일(曆日) 기준을 적용합니다.")
    doc.add_paragraph()

    add_heading(doc, "3.1.2  공기연장 일수 산정", level=3)
    tbl = doc.add_table(rows=1, cols=4)
    add_header_row(tbl,
                   ["구  분", "기  간", "일  수", "비  고"],
                   widths_cm=[3.5, 7.0, 2.5, 3.5])
    for label in ["당초 준공", "1회 연장", "2회 연장", "공기연장 합계"]:
        add_data_row(tbl, [label, PH + " ~ " + PH, PH + " 일", ""])
    doc.add_paragraph()

    # 3.2 산정방식
    add_heading(doc, "3.2  공기연장에 따른 간접비 산정방식", level=2)

    # 3.2.1 간접노무비
    add_heading(doc, "3.2.1  간접노무비", level=3)
    add_body(doc,
             "간접노무비는 공기연장 기간 동안 현장에 상주하여 공사 관리업무를 수행한 "
             "직원의 노무비로 산정합니다. 노무비 단가는 기본급, 제수당, 상여금, "
             "퇴직급여충당금의 합계액으로 구성됩니다.")
    add_body(doc, "  ① 기본급: 실제 지급 기본급")
    add_body(doc, "  ② 제수당: 식대, 교통비, 직책수당 등")
    add_body(doc, "  ③ 상여금: 연간 상여금 ÷ 12")
    add_body(doc, "  ④ 퇴직급여충당금: 월 급여 합계 ÷ 12")
    doc.add_paragraph()

    # 3.2.2 경비
    add_heading(doc, "3.2.2  경비", level=3)

    add_body(doc, "① 승률계상비목")
    add_body(doc,
             "승률계상비목 경비는 「" + PH + " 집행기준」에 의거 도급내역서에 "
             "반영된 요율을 적용하여 산정합니다.")
    add_body(doc, "  • 산재보험료 = 간접노무비  ×  " + PH + " %  (산출내역서상 요율)")
    add_body(doc, "  • 고용보험료 = 간접노무비  ×  " + PH + " %  (산출내역서상 요율)")
    doc.add_paragraph()

    add_body(doc, "② 직접계상비목")
    add_body(doc,
             "직접계상비목 경비는 「" + PH + " 집행기준」에 따라 실제 지출된 "
             "계약서, 요금고지서, 영수증 등 객관적인 자료를 기준으로 산정합니다.")
    add_body(doc,
             "  해당 비목: 전력비·수도광열비, 여비·교통비·통신비, 지급임차료, "
             "복리후생비, 소모품비, 세금과공과, 도서인쇄비, 지급수수료 등")
    doc.add_paragraph()

    add_body(doc, "③ 보증수수료")
    add_body(doc,
             "보증수수료는 「" + PH + " 집행기준」제73조 제4항에 의거하여 "
             "보증수수료 영수증 등 객관적인 자료에 의하여 확인된 금액으로 반영합니다.")
    doc.add_paragraph()

    # 3.2.3 일반관리비
    add_heading(doc, "3.2.3  일반관리비", level=3)
    add_body(doc,
             "일반관리비는 간접노무비와 경비의 합계액에 산출내역서상 요율을 곱하여 계상합니다.")

    tbl2 = doc.add_table(rows=2, cols=2)
    add_header_row(tbl2, ["산  출  방  식", "적용 요율"],
                   widths_cm=[11.0, 5.5])
    set_cell(tbl2.rows[1].cells[0],
             "일반관리비  =  (간접노무비 + 경비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl2.rows[1].cells[1],
             PH + " %  (산출내역서 기준)",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 3.2.4 이윤
    add_heading(doc, "3.2.4  이윤", level=3)
    add_body(doc,
             "이윤은 간접노무비, 경비, 일반관리비의 합계액에 산출내역서상 요율을 곱하여 계상합니다.")

    tbl3 = doc.add_table(rows=2, cols=2)
    add_header_row(tbl3, ["산  출  방  식", "적용 요율"],
                   widths_cm=[11.0, 5.5])
    set_cell(tbl3.rows[1].cells[0],
             "이윤  =  (간접노무비 + 경비 + 일반관리비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl3.rows[1].cells[1],
             PH + " %  (산출내역서 기준)",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 3.2.5 보증수수료
    add_heading(doc, "3.2.5  보증수수료", level=3)
    add_body(doc,
             "보증수수료는 연장된 공사기간에 해당하는 계약보증서 및 하자보증서의 "
             "보증수수료로, 증빙자료 확인 금액을 기준으로 계상합니다.")


# ── 제4장 ────────────────────────────────────────────────────────────────────

def make_ch4(doc):
    doc.add_page_break()
    add_heading(doc, "제4장.  공기연장에 따른 간접비 산정 결과", level=1)

    # 4.1 집계표
    add_heading(doc, "4.1  간접비 집계표", level=2)
    add_note(doc, "※ [단위: 원, 영세율]")

    tbl = doc.add_table(rows=1, cols=7)
    add_header_row(tbl,
                   ["구  분", "간접노무비", "직접계상\n경비", "승률계상\n경비",
                    "경비 소계", "일반관리비", "이  윤"],
                   widths_cm=[2.5, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5])
    for label in ["원도급", "하도급", "합  계"]:
        row = tbl.add_row()
        set_cell(row.cells[0], label, bold=(label == "합  계"),
                 bg_color=(HEADER_BG if label == "합  계" else None))
        for j in range(1, 7):
            set_cell(row.cells[j], PH,
                     bg_color=(HEADER_BG if label == "합  계" else None))
    doc.add_paragraph()

    # 최종 합계 강조
    add_note(doc, "※ 보증수수료는 경비에 포함하여 계상합니다.")
    tbl_total = doc.add_table(rows=2, cols=3)
    add_header_row(tbl_total,
                   ["구  분", "공기연장 간접비 합계 (원)", "비  고"],
                   widths_cm=[4.0, 9.0, 3.5])
    add_data_row(tbl_total, ["최 종 청구액", PH, "VAT 별도"])
    doc.add_paragraph()

    # 4.2 간접노무비
    add_heading(doc, "4.2  간접노무비", level=2)
    add_heading(doc, "4.2.1  대상인원", level=3)

    tbl2 = doc.add_table(rows=1, cols=6)
    add_header_row(tbl2,
                   ["No.", "소  속", "이  름", "직  무",
                    "간접노무비 산정 대상 기간", "비  고"],
                   widths_cm=[1.0, 3.5, 2.5, 2.5, 5.5, 1.5])
    for i in range(1, 6):
        add_data_row(tbl2,
                     [str(i), PH, PH, PH,
                      PH + " ~ " + PH + "  (" + PH + "일)", ""])
    doc.add_paragraph()

    add_heading(doc, "4.2.2  급여내역", level=3)
    add_note(doc, "※ [단위: 원]  ※ A. 실비 산정 기간: " + PH + " ~ " + PH)

    tbl3 = doc.add_table(rows=2, cols=9)
    # 헤더 2행 구성
    h1 = tbl3.rows[0].cells
    set_cell(h1[0], "No.",     bold=True, bg_color=HEADER_BG)
    set_cell(h1[1], "소속",    bold=True, bg_color=HEADER_BG)
    set_cell(h1[2], "이름",    bold=True, bg_color=HEADER_BG)
    set_cell(h1[3], "직무",    bold=True, bg_color=HEADER_BG)
    set_cell(h1[4], "A. 실비 산정",  bold=True, bg_color=HEADER_BG)
    set_cell(h1[5], "A. 실비 산정",  bold=True, bg_color=HEADER_BG)
    set_cell(h1[6], "A. 실비 산정",  bold=True, bg_color=HEADER_BG)
    set_cell(h1[7], "A. 실비 산정",  bold=True, bg_color=HEADER_BG)
    set_cell(h1[8], "합계",    bold=True, bg_color=HEADER_BG)
    h2 = tbl3.rows[1].cells
    set_cell(h2[0], "No.",     bold=True, bg_color=HEADER_BG)
    set_cell(h2[1], "소속",    bold=True, bg_color=HEADER_BG)
    set_cell(h2[2], "이름",    bold=True, bg_color=HEADER_BG)
    set_cell(h2[3], "직무",    bold=True, bg_color=HEADER_BG)
    set_cell(h2[4], "① 급여",         bold=True, bg_color=HEADER_BG)
    set_cell(h2[5], "② 일수",         bold=True, bg_color=HEADER_BG)
    set_cell(h2[6], "③ 퇴직급여충당금\n(①/12)", bold=True, bg_color=HEADER_BG)
    set_cell(h2[7], "④ 소계\n(①+③)", bold=True, bg_color=HEADER_BG)
    set_cell(h2[8], "합  계",         bold=True, bg_color=HEADER_BG)
    set_col_widths(tbl3, [1.0, 2.5, 2.0, 2.0, 2.5, 1.2, 2.5, 2.5, 2.5])
    set_table_border(tbl3)
    for i in range(1, 5):
        add_data_row(tbl3, [str(i), PH, PH, PH, PH, PH, PH, PH, PH])
    doc.add_paragraph()

    # 4.3 경비
    add_heading(doc, "4.3  경비", level=2)
    add_heading(doc, "4.3.1  직접계상비목", level=3)
    add_note(doc, "※ [단위: 원]")

    tbl4 = doc.add_table(rows=1, cols=5)
    add_header_row(tbl4,
                   ["비  목", "산정 기간", "금  액 (원)", "증빙자료", "비  고"],
                   widths_cm=[3.5, 4.0, 3.5, 3.5, 2.0])
    direct_items = [
        "전력비·수도광열비", "여비·교통비·통신비", "지급임차료",
        "복리후생비", "소모품비", "세금과공과",
        "국민건강·연금보험", "도서인쇄비", "지급수수료",
        "직접계상비목 소계",
    ]
    for item in direct_items:
        bold = item == "직접계상비목 소계"
        row = tbl4.add_row()
        set_cell(row.cells[0], item, bold=bold,
                 bg_color=(HEADER_BG if bold else None),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for j in range(1, 5):
            set_cell(row.cells[j], PH if not bold else PH,
                     bold=bold,
                     bg_color=(HEADER_BG if bold else None))
    doc.add_paragraph()

    add_heading(doc, "4.3.2  승률계상비목", level=3)
    add_note(doc, "※ [단위: 원]")

    tbl5 = doc.add_table(rows=1, cols=4)
    add_header_row(tbl5,
                   ["비  목", "산  출  방  식", "적용 요율", "금  액 (원)"],
                   widths_cm=[3.5, 7.0, 2.5, 3.5])
    rate_items = [
        ("산재보험료",  "간접노무비  ×  요율"),
        ("고용보험료",  "간접노무비  ×  요율"),
        ("승률계상비목 소계", ""),
    ]
    for item, formula in rate_items:
        bold = "소계" in item
        row = tbl5.add_row()
        set_cell(row.cells[0], item, bold=bold,
                 bg_color=(HEADER_BG if bold else None),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[1], formula if not bold else "",
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], PH + " %" if not bold else "",
                 bold=bold, bg_color=(HEADER_BG if bold else None))
        set_cell(row.cells[3], PH,
                 bold=bold, bg_color=(HEADER_BG if bold else None))
    doc.add_paragraph()

    # 4.4 일반관리비 및 이윤
    add_heading(doc, "4.4  일반관리비 및 이윤", level=2)
    add_heading(doc, "4.4.1  일반관리비", level=3)

    tbl6 = doc.add_table(rows=2, cols=4)
    add_header_row(tbl6,
                   ["산  출  방  식", "적용 요율", "기준 금액 (원)", "일반관리비 (원)"],
                   widths_cm=[6.5, 2.5, 3.5, 4.0])
    set_cell(tbl6.rows[1].cells[0],
             "(간접노무비 + 경비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl6.rows[1].cells[1], PH + " %")
    set_cell(tbl6.rows[1].cells[2], PH)
    set_cell(tbl6.rows[1].cells[3], PH)
    doc.add_paragraph()

    add_heading(doc, "4.4.2  이윤", level=3)

    tbl7 = doc.add_table(rows=2, cols=4)
    add_header_row(tbl7,
                   ["산  출  방  식", "적용 요율", "기준 금액 (원)", "이  윤 (원)"],
                   widths_cm=[6.5, 2.5, 3.5, 4.0])
    set_cell(tbl7.rows[1].cells[0],
             "(간접노무비 + 경비 + 일반관리비)  ×  요율",
             align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(tbl7.rows[1].cells[1], PH + " %")
    set_cell(tbl7.rows[1].cells[2], PH)
    set_cell(tbl7.rows[1].cells[3], PH)
    doc.add_paragraph()


# ── 제5장 ────────────────────────────────────────────────────────────────────

def make_ch5(doc):
    doc.add_page_break()
    add_heading(doc, "제5장.  첨부자료", level=1)
    add_note(doc, "※ 아래 각 절에 해당 자료를 순서대로 첨부합니다.")
    doc.add_paragraph()

    sections = [
        ("5.1  공기연장 간접비 산정근거",
         "급여명세서, 계좌이체 내역, 노무비 상세 계산서 등"),
        ("5.2  계약문서",
         "공사도급계약서, 산출내역서, 변경계약서 등"),
        ("5.3  수발신문서",
         "귀책사유 관련 공문 원본 (발신/수신 순)"),
        ("5.4  공기연장 간접비 증빙자료",
         "전력·수도 고지서, 임대차계약서, 영수증, 보증수수료 영수증 등"),
    ]
    for title, desc in sections:
        add_heading(doc, title, level=2)
        add_note(doc, "  ▶ " + desc)
        add_body(doc, "  [ 자료 첨부 위치 ]")
        doc.add_paragraph()


# ── 메인 ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # 페이지 여백 설정 (A4, 상하 2.5cm, 좌우 3.0cm)
    section = doc.sections[0]
    section.page_height  = Cm(29.7)
    section.page_width   = Cm(21.0)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin  = Cm(3.0)
    section.right_margin = Cm(3.0)

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_MAIN)

    # 각 섹션 생성
    make_cover(doc)
    make_submission(doc)
    make_summary(doc)
    make_toc(doc)
    make_ch1(doc)
    make_ch2(doc)
    make_ch3(doc)
    make_ch4(doc)
    make_ch5(doc)

    # 저장
    out_path = Path("reference/보고서_템플릿.docx")
    doc.save(str(out_path))
    print(f"저장 완료: {out_path}")
    print(f"  ※ 【 입력 】 표시 부분을 실제 값으로 교체하여 사용하세요.")


if __name__ == "__main__":
    main()
