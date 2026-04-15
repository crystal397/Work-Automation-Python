"""
유형별 워드 템플릿 생성기
실행: python src/generator/build_templates.py

13개 레퍼런스를 전체 비교 → 유형별 공통 섹션은 고정 텍스트로,
현장마다 다른 섹션은 {{ }} Jinja2 자리표시자로 교체합니다.

생성 파일:
  src/generator/templates/template_A.docx
  src/generator/templates/template_B.docx
  src/generator/templates/template_C.docx
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

ROOT = Path(__file__).parent.parent.parent
REF  = ROOT / "reference"
TMPL = Path(__file__).parent / "templates"
TMPL.mkdir(exist_ok=True)

HEADING_STYLES = {
    "1. 개요 2", "1.1. 개요3", "1.1.1. 개요 4",
    "가) 개요 5", "① 개요 7",
}

STYLE_BASE = {
    "A": REF / "광주도시철도 2호선 11공구_공기연장 간접비 산정 보고서_250617.docx",
    "B": REF / "인덕원~동탄 4공구_공기연장 간접비 산정 보고서_260311.docx",
    "C": REF / "한진대전 공기연장에 따른 추가공사비 사감정보고서_240712.docx",
}

TYPE_FILES = {
    "A": ["광주도시철도 2호선 11공구","광주도시철도 2호선 13공구","광주도시철도 2호선 14공구",
          "송도 11-1공구","양산선2공구"],
    "B": ["당진기지1단계","대한민국 축구종합센터","세종안성10공구","신세계건설_원주군부대",
          "인덕원~동탄 4공구","인덕원~동탄 5공구","평택기지~오산"],
    "C": ["한진대전"],
}


# ══════════════════════════════════════════════════════════════════════════════
# 텍스트 치환 (runs 병합 후 교체)
# ══════════════════════════════════════════════════════════════════════════════

def _replace_para(para, src: str, dst: str) -> bool:
    full = "".join(r.text for r in para.runs)
    if src not in full:
        return False
    new = full.replace(src, dst)
    if para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_in_doc(doc: Document, pairs: list[tuple[str, str]]):
    def _do(para):
        for src, dst in pairs:
            _replace_para(para, src, dst)

    for p in doc.paragraphs:
        _do(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _do(p)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if not part:
                continue
            for p in part.paragraphs:
                _do(p)
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _do(p)


# ══════════════════════════════════════════════════════════════════════════════
# 섹션 단위 → 자리표시자 교체
# 단락 + 표를 body 순서대로 함께 처리
# ══════════════════════════════════════════════════════════════════════════════

def _iter_body(doc: Document):
    """body 자식 요소를 (tag, object) 순서대로 반환"""
    body = doc.element.body
    for child in body.iterchildren():
        local = child.tag.split("}")[-1]
        if local == "p":
            yield "p", Paragraph(child, doc), child
        elif local == "tbl":
            yield "tbl", Table(child, doc), child


def replace_section_with_placeholder(doc: Document, section_name: str, placeholder: str):
    """
    section_name 헤딩 이후의 단락·표를 모두 제거하고,
    헤딩 바로 다음 위치에 placeholder 단락 한 개를 삽입.
    동명 섹션이 여러 개 있으면 모두 처리한다.
    """
    from docx.oxml import OxmlElement

    body = doc.element.body

    def _make_placeholder_para():
        p_new = OxmlElement("w:p")
        r_new = OxmlElement("w:r")
        t_new = OxmlElement("w:t")
        t_new.text = placeholder
        t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_new.append(t_new)
        p_new.append(r_new)
        return p_new

    # 전체 body 요소를 리스트로 복사 (제거 중 반복 안정성 확보)
    all_elems = list(body.iterchildren())

    i = 0
    while i < len(all_elems):
        elem = all_elems[i]
        local = elem.tag.split("}")[-1]

        if local == "p":
            p = Paragraph(elem, doc)
            txt = p.text.strip()
            if p.style.name in HEADING_STYLES and txt == section_name:
                # 섹션 시작 — 다음 헤딩 또는 끝까지 수집
                heading_elem = elem
                to_remove = []
                j = i + 1
                while j < len(all_elems):
                    nxt = all_elems[j]
                    nxt_local = nxt.tag.split("}")[-1]
                    if nxt_local == "p":
                        np = Paragraph(nxt, doc)
                        if np.style.name in HEADING_STYLES and np.text.strip():
                            break   # 다음 헤딩 만남
                        to_remove.append(nxt)
                    elif nxt_local == "tbl":
                        to_remove.append(nxt)
                    j += 1

                # 제거 실행
                for el in to_remove:
                    body.remove(el)
                    # all_elems 리스트도 동기화
                    if el in all_elems:
                        all_elems.remove(el)

                # placeholder 단락 삽입
                p_new = _make_placeholder_para()
                heading_elem.addnext(p_new)
                # 삽입된 요소를 리스트에도 추가 (다음 순회에서 스킵)
                all_elems.insert(i + 1, p_new)
                i += 2   # heading + placeholder 넘기기
                continue
        i += 1


# ══════════════════════════════════════════════════════════════════════════════
# 유형별 치환 정의
# ══════════════════════════════════════════════════════════════════════════════

TEXT_REPLACEMENTS: dict[str, list[tuple[str, str]]] = {
    "A": [
        # 표지 공사명 (장-해 스타일)
        ("광주 도시철도 2호선 2단계 11공구 건설공사", "{{ project_name }}"),
        ("공사연장 간접비 산정 보고서", "공기연장 간접비 산정 보고서"),
        # 의뢰처 (표지 첫 줄)
        ("디엘이앤씨 주식회사", "{{ contractor }}"),
        # 제출문
        ("귀사로부터 의뢰받은 「{{ project_name }}」의 공사기간 연장에 따른 간접비 산정의 결과로 본 보고서를 제출합니다.",
         "귀사로부터 의뢰받은 「{{ project_name }}」의 공사기간 연장에 따른 간접비 산정의 결과로 본 보고서를 제출합니다."),
        # 날짜
        ("2025년 5월", "{{ report_year }}년  {{ report_month }}월"),
        # 요약문
        ("{{ contractor }}의「{{ project_name }}」현장 및 본사에서 제공한",
         "{{ contractor }}의「{{ project_name }}」현장 및 본사에서 제공한"),
    ],
    "B": [
        # 표지 수신
        ("롯데건설 주식회사 귀 중", "{{ contractor }} 귀 중"),
        # 표지 공사명
        ("인덕원~동탄 복선전철 제4공구 노반신설 기타공사", "{{ project_name }}"),
        # 계약자 (법인명 포함 → 단축명 순서로)
        ("롯데건설 주식회사", "{{ contractor }}"),
        ("롯데건설",          "{{ contractor }}"),
        # 날짜
        ("2026년 3월", "{{ report_year }}년  {{ report_month }}월"),
        # 원가계산서 제목 표
        ("인덕원~동탄 복선전철 제4공구 노반신설 기타공사 – 원가계산서",
         "{{ project_name }} – 원가계산서"),
        # 요약문 연결
        ("{{ contractor }}의「{{ project_name }}」 현장 및 본사에서 제공한",
         "{{ contractor }}의「{{ project_name }}」현장 및 본사에서 제공한"),
        # 장기계속 계약 단락의 현장 특정 계약일
        ("2023. 12. 21.", "{{ first_contract_date }}"),
        # 집계표 요약(Table 1) 하도급사명 → 자리표시자
        ("경동건설", "{{ subcon_name_1 }}"),
        ("도양기업", "{{ subcon_name_2 }}"),
        # 집계표(Table 2) 비고 열 요율 → 자리표시자
        ("(1)×4.18%", "(1)×{{ sangjae_rate_pct }}"),
        ("(1)×1.77%", "(1)×{{ goyong_rate_pct }}"),
        ("(3)×4.50%", "(3)×{{ admin_rate_pct }}"),
        ("(3+4)×0.00%", "(3+4)×{{ profit_rate_pct }}"),
    ],
    "C": [
        # 표지 공사명 (㈜ 형식)
        ("㈜한진 대전 SMART Mega-Hub 신축공사", "{{ project_name }}"),
        # 본문 공사명 ((주) 형식 — 동일 파일 내 두 가지 표기 혼용)
        ("(주)한진대전 SMART Mega-Hub 신축공사", "{{ project_name }}"),
        # 한진대전 단독 언급
        ("(주)한진대전", "{{ contractor_short }}"),
        # 원도급사 계약자
        ("삼성물산 주식회사", "{{ contractor }}"),
        ("삼성물산", "{{ contractor_short }}"),
        # 날짜
        ("2024년  07월", "{{ report_year }}년  {{ report_month }}월"),
    ],
}

# 현장마다 다른 섹션 → 전체를 자리표시자로 교체
# (분석 결과 0% 공통인 섹션들)
VARIABLE_SECTIONS: dict[str, list[tuple[str, str]]] = {
    "A": [
        ("공사의 개요",                         "{{ contract_overview }}"),
        ("공사 계약현황",                        "{{ contract_table }}"),
        ("청구의 사유 개요",                     "{{ cause_overview }}"),
        ("계약에 의한 청구권",                   "{{ claim_basis }}"),
        ("공기지연의 귀책사유 분석",              "{{ cause_analysis }}"),
        ("공기지연 귀책사유 관련 근거",            "{{ cause_grounds }}"),
        ("공기연장 기간의 산정",                  "{{ extension_section }}"),
        ("공기연장 일수의 산정방식",               "{{ extension_method }}"),
        ("공기연장 일수 산정",                    "{{ extension_table }}"),
        ("대상인원",                             "{{ labor_list }}"),
        # 급여내역: 설명 문장은 고정이나 표 데이터는 변수 → 통째로 교체
        ("급여내역",                             "{{ labor_table }}"),
        ("직접계상비목",                         "{{ expense_direct_table }}"),
        ("승률계상비목",                         "{{ expense_rate_table }}"),
        ("일반관리비",                           "{{ admin_table }}"),
        ("이윤",                                "{{ profit_table }}"),
        ("공기연장에 따른 간접비 집계표",          "{{ total_table }}"),
        # 하도급 있는 경우 (송도, 양산)
        ("원도급 계약 현황",                     "{{ prime_contract_table }}"),
        ("하도급 계약 현황",                     "{{ sub_contract_table }}"),
        ("공기연장에 따른 하도급 간접비 산정 결과", "{{ subcon_result }}"),
        ("원도급 공사 계약현황",                  "{{ prime_contract_table }}"),
        ("하도급 공사 계약현황",                  "{{ sub_contract_table }}"),
        ("원도급 계약 현황",                     "{{ prime_contract_table }}"),
        ("사정변경의 원칙",                       "{{ cause_principle }}"),
    ],
    "B": [
        ("용역 대상 공사의 개요 및 특성",          "{{ contract_overview }}"),
        ("공사 위치도",                           "{{ location_map }}"),
        ("계약당사자 현황",                       "{{ parties_table }}"),
        # 계약 현황·변경 현황 (원도급/하도급 모두)
        ("계약 현황",                            "{{ contract_table }}"),
        ("총공사 계약 현황",                      "{{ total_contract_table }}"),
        ("차수공사 계약 현황",                    "{{ phase_contract_table }}"),
        ("원도급사 계약 현황",                    "{{ prime_contract_table }}"),
        ("하도급사 계약 현황",                    "{{ sub_contract_table }}"),
        ("공사 기간 변경 계약 현황",               "{{ change_history_table }}"),
        ("원도급사 공사기간 변경 계약 현황",        "{{ prime_change_table }}"),
        ("하도급사 공사기간 변경 계약 현황",        "{{ sub_change_table }}"),
        # 청구권 검토 섹션 (계약의 성격 하위) — 한자 포함 법령 표 포함 → 자리표시자로 교체
        ("계약에 의한 청구권 검토",               "{{ contractual_basis }}"),
        ("관련 법령에 따른 청구권 검토",           "{{ legal_basis }}"),
        ("사정변경의 원칙",                       "{{ cause_principle }}"),
        # 귀책 분석 관련
        ("지연사유 발생 경위",                    "{{ cause_background }}"),
        ("공기지연 귀책 사유 검토",               "{{ cause_review }}"),
        ("절차적 요건에 대한 준수 여부",           "{{ procedure_check }}"),
        ("계약금액조정 신청 관련 문서",            "{{ adjustment_docs }}"),
        ("검토 의견",                            "{{ review_opinion }}"),
        ("계약에 따른 청구의 근거",               "{{ contractual_basis }}"),
        ("관련 법령에 따른 청구의 근거",           "{{ legal_basis }}"),
        # 산정 기간·일수
        ("공기연장 간접비 산정 대상 기간 검토",     "{{ target_period }}"),
        ("공기연장 기간의 산정",                  "{{ extension_section }}"),
        ("공기연장 일수의 산정방식",               "{{ extension_method }}"),
        ("공기연장 일수 산정",                    "{{ extension_table }}"),
        ("원도급사 계약의 공기연장 일수 산정",      "{{ prime_extension_table }}"),
        ("하도급사 계약의 공기연장 일수 산정",      "{{ sub_extension_table }}"),
        # 간접노무비 결과 (방법론 설명은 고정 → 결과 표만 교체)
        ("대상인원",                             "{{ labor_list }}"),
        ("급여내역",                             "{{ labor_table }}"),
        # 경비 결과 표
        ("직접계상비목",                         "{{ expense_direct_table }}"),
        ("승률계상비목",                         "{{ expense_rate_table }}"),
        # 일반관리비·이윤·집계
        ("일반관리비",                           "{{ admin_table }}"),
        ("이윤",                                "{{ profit_table }}"),
        ("공기연장 간접비 집계표",               "{{ total_table }}"),
        ("간접공사비 산정 결과",                 "{{ grand_result }}"),
        ("공기연장 간접비 원도급사 산정 결과",    "{{ prime_result }}"),
        ("공기연장 간접비 원도급사 집계표",       "{{ prime_total_table }}"),
        ("결론",                                "{{ conclusion }}"),
        # 하도급사 개별 섹션 (회사명이 섹션 제목인 경우)
        ("경동건설",   "{{ subcon_kyungdong }}"),
        ("도양기업",   "{{ subcon_doyang }}"),
        ("시재건설",   "{{ subcon_sijae }}"),
        ("주일건설",   "{{ subcon_juil }}"),
        ("디에이치건업","{{ subcon_dh }}"),
        ("원산건설",   "{{ subcon_wonsan }}"),
        ("하나전기",   "{{ subcon_hana }}"),
        ("환경이엔지", "{{ subcon_hwangkyung }}"),
        ("지준시스템", "{{ subcon_jijun }}"),
        ("한창건설",   "{{ subcon_hanchang }}"),
        ("케이제이건설산업", "{{ subcon_kj }}"),
        # 하도급 합계
        ("공기연장 간접비 하도급사 산정 결과",    "{{ subcon_result }}"),
        ("공기연장 간접비 하도급사 집계표",       "{{ subcon_total }}"),
        ("하도급 전체 집계표",                   "{{ subcon_grand_total }}"),
        # B 잔존 섹션
        ("국가계약법 적용 공사",                 "{{ national_contract_law }}"),
        ("계약문서(공사계약 일반조건 제3조)",     "{{ contract_documents }}"),
    ],
    "C": [
        ("공사의 개요",                                     "{{ contract_overview }}"),
        ("대상공사의 계약현황",                              "{{ contract_table }}"),
        ("청구의 사유 개요",                                 "{{ cause_overview }}"),
        ("사정변경의 원칙",                                  "{{ cause_principle }}"),
        ("계약에 의한 청구권",                               "{{ claim_basis }}"),
        ("공기지연의 귀책사유 분석",                         "{{ cause_analysis }}"),
        ("공기지연 귀책사유 관련 근거",                      "{{ cause_grounds }}"),
        ("공기지연 귀책사유 구분",                           "{{ cause_classification }}"),
        ("공기지연 일수의 산정방식",                         "{{ extension_method }}"),
        ("본건 공사 계약의 공기지연 일수 산정",              "{{ extension_table }}"),
        # 3장 방법론
        ("공기연장에 따른 직접비",                           "{{ direct_cost_method }}"),
        ("산업안전보건관리비 초과 계상 비용",                 "{{ safety_cost_method }}"),
        ("하도급사 공기연장 및 공정만회에 따른 추가공사비",   "{{ subcon_method }}"),
        # 4장 산정 결과
        ("대상인원",                                        "{{ labor_list }}"),
        ("급여내역",                                        "{{ labor_table }}"),
        ("직접계상비목",                                    "{{ expense_direct_table }}"),
        ("승률계상비목",                                    "{{ expense_rate_table }}"),
        ("일반관리비",                                      "{{ admin_table }}"),
        ("이윤",                                           "{{ profit_table }}"),
        ("공기연장에 따른 간접비 집계표",                    "{{ indirect_total_table }}"),
        ("추가공사비비 산정 결과",                           "{{ result_section }}"),
        ("추가공사비 총괄집계표",                            "{{ total_table }}"),
        # 잔존 섹션
        ("추가공사비 산정 방법",                            "{{ cost_estimation_method }}"),
        ("경비",                                           "{{ expense_section }}"),
    ],
}


# ══════════════════════════════════════════════════════════════════════════════
# FRONT 영역 요약 표 정리 (산정결과 / 원가계산서 요약)
# 모든 섹션 헤딩 이전 표에서 날짜·금액 셀을 Jinja2 변수로 교체
# ══════════════════════════════════════════════════════════════════════════════

import re as _re

_DATE_RE   = _re.compile(r"\d{4}[.년]\s*\d{1,2}[.월]")
_AMOUNT_RE = _re.compile(r"^[\d,]{5,}$")   # 숫자만 있는 셀
_PERIOD_RE = _re.compile(
    r"\d{4}\.\s*\d{1,2}\.\s*\d{1,2}\.?\s*[~～]\s*\d{4}\.\s*\d{1,2}\.\s*\d{1,2}"
)
_DAYS_RE   = _re.compile(r"^\d+일$")        # 단독 일수 (예: 159일)


def cleanup_front_tables(doc: Document):
    """
    첫 번째 헤딩이 나오기 전(FRONT 영역)의 표들에서
    날짜·기간·금액 셀을 적절한 Jinja2 변수로 교체합니다.
    """
    body = doc.element.body
    replacement_map = {
        "period":      "{{ extension_start }}~{{ extension_end }}({{ extension_days }}일)",
        "grand_total": "{{ grand_total }}",
        "date":        "{{ report_year }}. {{ report_month }}.",
    }

    for child in body.iterchildren():
        local = child.tag.split("}")[-1]
        if local == "p":
            p = Paragraph(child, doc)
            if p.style.name in HEADING_STYLES and p.text.strip():
                break   # 첫 헤딩 → FRONT 영역 종료
        elif local == "tbl":
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        txt = para.text.strip()
                        if not txt:
                            continue
                        # 이미 자리표시자인 경우 스킵
                        if "{{" in txt:
                            continue
                        # FRONT 표에서 단독 일수(예: 159일)는 항상 제거
                        # (산정기간 셀 두 번째 줄에 하드코딩된 일수)
                        if _DAYS_RE.match(txt):
                            for r in para.runs:
                                r.text = ""
                            continue
                        new_txt = None
                        if _PERIOD_RE.search(txt):
                            new_txt = replacement_map["period"]
                        elif _AMOUNT_RE.match(txt.replace(",", "")):
                            new_txt = replacement_map["grand_total"]
                        elif _DATE_RE.search(txt) and "시행" not in txt and "선고" not in txt:
                            new_txt = replacement_map["date"]

                        if new_txt and para.runs:
                            para.runs[0].text = new_txt
                            for r in para.runs[1:]:
                                r.text = ""
                        elif new_txt:
                            para.add_run(new_txt)


# ══════════════════════════════════════════════════════════════════════════════
# 이미지 제거 (레퍼런스 특정 로고·도면 이미지가 템플릿에 잔존하지 않도록)
# ══════════════════════════════════════════════════════════════════════════════

_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DNS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _remove_all_drawings(doc: Document) -> int:
    """
    문서 내 모든 <w:drawing> 요소를 제거하고 제거 개수를 반환한다.
    헤더/푸터의 이미지(로고 등)는 유지한다.
    """
    drawing_tag = f"{{{_WNS}}}drawing"
    count = 0

    def _remove_from(elem):
        nonlocal count
        for child in list(elem):
            if child.tag == drawing_tag:
                elem.remove(child)
                count += 1
            else:
                _remove_from(child)

    _remove_from(doc.element.body)
    return count


# ══════════════════════════════════════════════════════════════════════════════
# 메인 빌드
# ══════════════════════════════════════════════════════════════════════════════

def build(report_type: str):
    src   = STYLE_BASE[report_type]
    dst   = TMPL / f"template_{report_type}.docx"
    files_used = []
    for kw in TYPE_FILES[report_type]:
        matched = [f for f in REF.glob("*.docx") if kw in f.name]
        if matched:
            files_used.append(matched[0])

    print(f"\n[{report_type}] {len(files_used)}개 파일 참조 → {dst.name}")
    for f in files_used:
        print(f"       {f.name[:65]}")

    shutil.copy2(src, dst)
    doc = Document(dst)

    # 1) 변수 섹션 내용 제거 + 자리표시자 삽입 (표 포함)
    for sec_name, placeholder in VARIABLE_SECTIONS[report_type]:
        replace_section_with_placeholder(doc, sec_name, placeholder)
    print(f"  ✔ 변수 섹션 {len(VARIABLE_SECTIONS[report_type])}개 교체")

    # 2) FRONT 영역 요약 표 날짜·금액 교체
    cleanup_front_tables(doc)
    print(f"  ✔ FRONT 요약 표 정리")

    # 3) 공사명·계약자·날짜 텍스트 치환
    replace_in_doc(doc, TEXT_REPLACEMENTS[report_type])
    print(f"  ✔ 공사명/계약자/날짜 치환")

    # 4) 이미지 제거 (모든 유형 — 레퍼런스 특정 이미지가 잔존하지 않도록)
    removed_imgs = _remove_all_drawings(doc)
    if removed_imgs:
        print(f"  ✔ 이미지 {removed_imgs}개 제거")

    doc.save(dst)
    print(f"  저장: {dst}")


if __name__ == "__main__":
    for t in ("A", "B", "C"):
        build(t)
    print("\n템플릿 생성 완료 →", TMPL)
