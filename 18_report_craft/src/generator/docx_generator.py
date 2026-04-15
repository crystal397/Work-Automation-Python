"""
Word 문서 생성기
마크다운 텍스트 → .docx          (generate_docx)
템플릿 기반 생성                  (generate_docx_from_template)
python-docx 사용 (pip install python-docx)
"""

from __future__ import annotations
import copy
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 스타일 헬퍼 ───────────────────────────────────────────────────────────────

def _set_font(run, bold=False, size=None, color=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int):
    """H1~H6 → Word 제목 스타일"""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc: Document, text: str):
    """일반 단락 — **굵게**, *기울임* 처리"""
    p = doc.add_paragraph()
    # **bold** 와 *italic* 파싱
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p


def _set_cell_bg(cell, hex_color: str):
    """셀 배경색 설정 (hex: 'E8EDF2')"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_font(cell, bold=False, size=None):
    for para in cell.paragraphs:
        for run in para.runs:
            if bold:
                run.bold = True
            if size:
                run.font.size = Pt(size)


def _apply_table_style(table):
    """
    템플릿 표 스타일 매칭:
    얇은 검정 선(안팎 모두), 좁은 셀 여백, 100% 폭
    """
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    else:
        for child in list(tblPr):
            tblPr.remove(child)

    # 너비: 100% (페이지 맞춤)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # 테두리 (외곽 + 내부선 모두 얇은 검정)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:color"), "000000")
        b.set(qn("w:sz"),    "3")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # 셀 여백 (템플릿과 동일: 상하 28, 좌우 102 dxa)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("top", "28"), ("left", "102"), ("bottom", "28"), ("right", "102")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    val)
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)


def _add_table_from_md(doc: Document, rows: list[list[str]]):
    """마크다운 표 행 목록 → docx 표 (템플릿 스타일 매칭)"""
    if not rows:
        return
    # 구분선(|---|) 제거
    data = [r for r in rows if not re.match(r'^\|[-| ]+\|$', r[0] if r else "")]
    if not data:
        return

    col_count = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=col_count)
    _apply_table_style(table)

    for ri, row in enumerate(data):
        for ci, cell_text in enumerate(row):
            if ci < col_count:
                cell = table.cell(ri, ci)
                text = cell_text.strip().lstrip("*").rstrip("*")
                para = cell.paragraphs[0]
                para.clear()
                _add_paragraph_inline(para, text)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                # 헤더 행: 연회색 + 굵게 (템플릿 스타일)
                if ri == 0:
                    _set_cell_bg(cell, "D9D9D9")
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()  # 표 아래 공백


def _add_blockquote(doc: Document, text: str):
    """> 인용 → 들여쓰기 단락"""
    p = doc.add_paragraph(text.lstrip("> ").strip())
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── 마크다운 파서 ──────────────────────────────────────────────────────────────

def _parse_md_to_docx(doc: Document, md_text: str):
    """
    마크다운 텍스트를 순회하며 docx 요소로 변환.
    완전한 마크다운 파서가 아닌, 보고서에 필요한 요소만 처리:
    - # H1 ~ ### H3
    - | 표
    - > 인용
    - - 목록
    - --- 구분선
    - 일반 단락
    """
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── 구분선 ──
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph("─" * 50)
            i += 1
            continue

        # ── 제목 ──
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            _add_heading(doc, m.group(2).strip(), level=level)
            i += 1
            continue

        # ── 표 ── (연속된 | 행 수집)
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_lines.append(row)
                i += 1
            _add_table_from_md(doc, table_lines)
            continue

        # ── 인용 ──
        if line.startswith(">"):
            _add_blockquote(doc, line)
            i += 1
            continue

        # ── 목록 ──
        if re.match(r'^[-*]\s', line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r'^[-*]\s', "", line)
            _add_paragraph_inline(p, text)
            i += 1
            continue

        # ── 번호 목록 ──
        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r'^\d+\.\s', "", line)
            _add_paragraph_inline(p, text)
            i += 1
            continue

        # ── 빈 줄 ──
        if not line.strip():
            i += 1
            continue

        # ── 일반 단락 ──
        _add_paragraph(doc, line)
        i += 1


def _add_paragraph_inline(p, text: str):
    """기존 단락 객체에 인라인 마크다운(**bold**, *italic*) 추가"""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


# ── 문서 기본 설정 ────────────────────────────────────────────────────────────

def _add_page_number_field(run):
    """바닥글 페이지 번호 필드 삽입"""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_total_pages_field(run):
    """바닥글 전체 페이지 수 필드 삽입"""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _setup_header_footer(doc: Document, title: str = "공기연장 간접비 산정 보고서"):
    """머리글: 문서 제목 / 바닥글: 페이지 번호"""
    section = doc.sections[0]

    # ── 머리글 ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 머리글 아래 선
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── 바닥글 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run()
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _add_page_number_field(r1)
    r2 = fp.add_run(" / ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r3 = fp.add_run()
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _add_total_pages_field(r3)


def _setup_document(doc: Document):
    """여백, 기본 폰트 설정"""
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # 기본 폰트 한글 맑은 고딕
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    # 한글 폰트 설정 (East Asian)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)

    # 제목 스타일 한글 폰트 적용
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            h_style = doc.styles[style_name]
            h_style.font.name = "맑은 고딕"
            h_rPr = h_style.element.get_or_add_rPr()
            h_rFonts = OxmlElement("w:rFonts")
            h_rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            h_rPr.append(h_rFonts)
        except KeyError:
            pass


# ── 메인 ──────────────────────────────────────────────────────────────────────

def generate_docx(md_text: str, output_path: str | Path) -> Path:
    """
    마크다운 문자열 → output_path 에 .docx 저장
    반환값: 저장된 파일 경로
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_document(doc)
    _setup_header_footer(doc)
    _parse_md_to_docx(doc, md_text)

    doc.save(str(output_path))
    print(f"  Word 문서 저장: {output_path}")
    return output_path


# ── 템플릿 기반 Word 생성 ─────────────────────────────────────────────────────

_TMPL_DIR = Path(__file__).parent / "templates"

# {{ placeholder_name }} → 마크다운 섹션 헤딩 키워드
_PLACEHOLDER_MD_HEADING: dict[str, str] = {
    # 제1장: 개요
    "contract_overview":       "제2절 공사의 개요",
    "contract_table":          "제3절 계약현황",
    "prime_contract_table":    "1차수 공사 계약 현황",
    "total_contract_table":    "총공사 계약 현황",
    "sub_contract_table":      "하도급 계약 현황",
    "parties_table":           "제3절 계약당사자 현황",
    "change_history_table":    "1차수 공사 계약 현황",
    "prime_change_table":      "1차수 공사 계약 현황",
    "sub_change_table":        "하도급 계약 현황",
    # 제3장: 귀책 분석
    "cause_overview":          "제1절 공기지연의 귀책사유 분석",
    "cause_analysis":          "제1절 공기지연의 귀책사유 분석",
    "cause_background":        "1.1 공기지연 관련 공문 이력",
    "cause_review":            "1.2 귀책사유 분류",
    "cause_grounds":           "1.2 귀책사유 분류",
    "cause_classification":    "1.2 귀책사유 분류",
    "claim_basis":             "제2절 청구의 근거",
    "contractual_basis":       "2.1 계약에 의한 청구권",
    "legal_basis":             "2.2 법령에 의한 청구권",
    "cause_principle":         "2.3 사정변경의 원칙",
    "procedure_check":         "2.4 청구의 근거 검토 (종합)",
    "adjustment_docs":         "2.4 청구의 근거 검토 (종합)",
    "review_opinion":          "2.4 청구의 근거 검토 (종합)",
    "national_contract_law":   "제1절 적용 법령 체계",
    "contract_documents":      "계약문서 구성",
    # 제4장: 공기연장 기간
    "target_period":           "제1절 공기연장 기간의 산정",
    "extension_section":       "제1절 공기연장 기간의 산정",
    "extension_method":        "1.1 공기연장 일수의 산정방식",
    "extension_table":         "1.2 공기연장 일수 산정",
    "prime_extension_table":   "1.2 공기연장 일수 산정",
    "sub_extension_table":     "1.2 공기연장 일수 산정",
    # 제4장: 산정 결과
    "total_table":             "3.1 집계표",
    "prime_total_table":       "3.1 집계표",
    "prime_result":            "3.1 집계표",
    "grand_result":            "3.1 집계표",
    "indirect_total_table":    "3.1 집계표",
    "labor_list":              "① 대상인원 현황",
    "labor_table":             "② 급여 산정내역",
    "expense_direct_table":    "가. 직접계상비목",
    "expense_rate_table":      "나. 승률계상비목",
    "expense_section":         "3.3 경비 상세",
    "admin_table":             "3.4 일반관리비 및 이윤",
    "profit_table":            "3.4 일반관리비 및 이윤",
    # 유형 C
    "direct_cost_method":      "제2절 간접비 산정방식",
    "safety_cost_method":      "제2절 간접비 산정방식",
    "subcon_method":           "제4절 하도급사 간접비 산정 결과",
    "cost_estimation_method":  "제2절 간접비 산정방식",
    "result_section":          "결론",
    # 결론
    "conclusion":              "결론",
}

# 하도급사별 플레이스홀더 (첫 번째에만 전체 하도급 섹션 삽입, 나머지 제거)
_SUBCON_PLACEHOLDERS = {
    "subcon_kyungdong", "subcon_doyang", "subcon_sijae", "subcon_juil",
    "subcon_dh", "subcon_wonsan", "subcon_hana", "subcon_hwangkyung",
    "subcon_jijun", "subcon_hanchang", "subcon_kj",
    "subcon_result", "subcon_total", "subcon_grand_total",
}


def _extract_md_sections(md_text: str) -> dict[str, str]:
    """마크다운을 헤딩 기준으로 {헤딩 텍스트: 이후 내용(직접)} 으로 분리."""
    sections: dict[str, str] = {}
    current_heading: str | None = None
    current_lines: list[str] = []

    for line in md_text.splitlines():
        m = re.match(r'^#{1,6}\s+(.*)', line)
        if m:
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = m.group(1).strip()
            current_lines = []
        else:
            if current_heading is not None:
                current_lines.append(line)

    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()

    return sections


def _extract_section_block(md_text: str, heading_keyword: str) -> str:
    """
    heading_keyword 를 포함하는 헤딩부터 동일·상위 레벨 헤딩 전까지의
    전체 내용(하위 섹션 포함)을 반환.
    부분 매칭 지원: 키워드가 헤딩 텍스트에 포함되면 매칭.
    """
    lines = md_text.splitlines()

    start_idx = None
    heading_level = 0
    for i, line in enumerate(lines):
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m and heading_keyword in m.group(2).strip():
            start_idx = i + 1
            heading_level = len(m.group(1))
            break

    if start_idx is None:
        return ""

    content_lines: list[str] = []
    for line in lines[start_idx:]:
        m = re.match(r'^(#{1,6})\s+', line)
        if m and len(m.group(1)) <= heading_level:
            break
        content_lines.append(line)

    return '\n'.join(content_lines).strip()


def _replace_all_text(doc: Document, src: str, dst: str):
    """문서 전체(본문·표·머리글·바닥글)에서 src → dst 치환."""
    def _do(para):
        full = "".join(r.text for r in para.runs)
        if src in full:
            new_text = full.replace(src, dst)
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""

    for p in doc.paragraphs:
        _do(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _do(p)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if not part:
                continue
            for p in part.paragraphs:
                _do(p)
            for tbl in getattr(part, "tables", []):
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _do(p)


def _insert_md_at_placeholder(doc: Document, placeholder_text: str, md_content: str) -> bool:
    """
    body 내 placeholder_text 단락을 찾아 md_content(마크다운)로 교체.
    반환값: 하나 이상 교체 여부
    """
    from docx.text.paragraph import Paragraph as _Para

    body = doc.element.body
    all_elems = list(body.iterchildren())
    replaced = False
    i = 0

    while i < len(all_elems):
        elem = all_elems[i]
        if elem.tag.split("}")[-1] != "p":
            i += 1
            continue

        p = _Para(elem, doc)
        if p.text.strip() != placeholder_text:
            i += 1
            continue

        # 임시 문서에서 마크다운 변환
        temp_doc = Document()
        _is_missing = (
            not md_content
            or not md_content.strip()
            or (md_content.strip().startswith("[") and md_content.strip().endswith("]"))
        )
        if not _is_missing:
            _parse_md_to_docx(temp_doc, md_content)
        else:
            notice = md_content.strip() if (md_content and md_content.strip()) else f"[{placeholder_text} — 작성 필요]"
            p_miss = temp_doc.add_paragraph()
            run_miss = p_miss.add_run(notice)
            run_miss.bold = True
            run_miss.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # 새 요소 수집 (sectPr 제외)
        new_elems = [
            copy.deepcopy(child)
            for child in temp_doc.element.body.iterchildren()
            if child.tag.split("}")[-1] != "sectPr"
        ]

        # placeholder 다음에 삽입
        ref = elem
        for ne in new_elems:
            ref.addnext(ne)
            ref = ne

        # placeholder 제거 후 리스트 동기화
        body.remove(elem)
        all_elems.pop(i)
        for j, ne in enumerate(new_elems):
            all_elems.insert(i + j, ne)
        i += len(new_elems)
        replaced = True

    return replaced


def generate_docx_from_template(
    data: dict,
    calc: dict,
    report_type: str,
    output_path: str | Path,
    md_text: str,
) -> Path:
    """
    유형별 워드 템플릿 + 분析 결과 → .docx 생성

    1) 유형별 템플릿(template_A/B/C.docx) 복사
    2) {{ project_name }} 등 단순 텍스트 치환
    3) {{ labor_table }} 등 섹션 플레이스홀더를 마크다운 섹션 내용으로 교체
    4) 템플릿이 없으면 generate_docx() 폴백
    """
    import shutil

    output_path = Path(output_path)
    template_path = _TMPL_DIR / f"template_{report_type}.docx"

    if not template_path.exists():
        print(f"  [경고] 템플릿 없음 ({template_path.name}), 마크다운 직접 변환")
        return generate_docx(md_text, output_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))

    # ── 단순 텍스트 치환 ──────────────────────────────────────────────────────
    c     = data.get("contract", {})
    ext   = data.get("extension", {})
    rates = data.get("rates", {})
    subcons = data.get("subcontractors", [])
    from datetime import date as _today_cls
    today = _today_cls.today()
    final = calc.get("final_rounded", 0)

    def _pct_str(v) -> str:
        """float → 'X.XX%' (소수점 둘째 자리)"""
        if v is None:
            return "확인 필요"
        return f"{v * 100:.2f}%"

    for src, dst in [
        ("{{ project_name }}",    c.get("name", "[공사명]")),
        ("{{ contractor }}",      c.get("contractor", "[계약상대자]")),
        ("{{ contractor_short }}", c.get("contractor", "[계약상대자]")),
        ("{{ report_year }}",     str(today.year)),
        ("{{ report_month }}",    f"{today.month:02d}"),
        ("{{ grand_total }}",     f"{final:,}"),
        ("{{ extension_start }}", ext.get("start_date", "")),
        ("{{ extension_end }}",   ext.get("end_date", "")),
        ("{{ extension_days }}",  str(ext.get("total_days", ""))),
        # 장기계속 계약 단락의 최초 계약일
        ("{{ first_contract_date }}", c.get("start_date", "")),
        # 집계표 하도급사 이름
        ("{{ subcon_name_1 }}", subcons[0]["name"] if len(subcons) > 0 else ""),
        ("{{ subcon_name_2 }}", subcons[1]["name"] if len(subcons) > 1 else ""),
        # 집계표 비고 열 요율
        ("{{ sangjae_rate_pct }}", _pct_str(rates.get("industrial_accident"))),
        ("{{ goyong_rate_pct }}",  _pct_str(rates.get("employment"))),
        ("{{ admin_rate_pct }}",   _pct_str(rates.get("general_admin"))),
        ("{{ profit_rate_pct }}",  _pct_str(rates.get("profit"))),
    ]:
        _replace_all_text(doc, src, dst)

    # ── 섹션 플레이스홀더 교체 ────────────────────────────────────────────────
    md_sections = _extract_md_sections(md_text)
    subcon_md   = md_sections.get("제4절 하도급사 간접비 산정 결과", "")
    subcon_used = False

    # body 탐색하여 {{ ... }} 플레이스홀더 수집 (DOM 수정 전 완전히 탐색)
    from docx.text.paragraph import Paragraph as _Para
    pending: list[tuple[str, str]] = []

    for elem in list(doc.element.body.iterchildren()):
        if elem.tag.split("}")[-1] != "p":
            continue
        p   = _Para(elem, doc)
        txt = p.text.strip()
        if not (txt.startswith("{{") and txt.endswith("}}")):
            continue
        m = re.match(r'\{\{\s*(\w+)\s*\}\}', txt)
        if not m:
            continue
        pname = m.group(1)

        # 하도급사 전용
        if pname in _SUBCON_PLACEHOLDERS:
            if not subcon_used and subcon_md:
                pending.append((txt, subcon_md))
                subcon_used = True
            else:
                pending.append((txt, ""))
            continue

        # 일반 섹션
        md_heading = _PLACEHOLDER_MD_HEADING.get(pname)
        if md_heading:
            # 1) 직접 매칭
            content = md_sections.get(md_heading, "")
            # 2) 직접 내용이 없으면 하위 섹션 포함 블록 추출
            if not content:
                content = _extract_section_block(md_text, md_heading)
            pending.append((txt, content))
        else:
            pending.append((txt, f"[{pname} — 수동 작성 필요]"))   # [] 감싸면 빨간 볼드로 처리됨

    # 일괄 교체
    for placeholder_text, content in pending:
        _insert_md_at_placeholder(doc, placeholder_text, content)

    doc.save(str(output_path))
    print(f"  Word 문서 저장 (템플릿 기반): {output_path}")
    return output_path
